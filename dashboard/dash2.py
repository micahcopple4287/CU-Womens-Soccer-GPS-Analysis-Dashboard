# app.py
# ------------------------------------------------------------------------------
# CU WSOC PERFORMANCE DASHBOARD (Catapult + Game Stats + Schedules) 2023–2025
#
# Run:
#   pip install streamlit plotly scikit-learn pandas numpy openpyxl
#   streamlit run app.py
# ------------------------------------------------------------------------------

import re
from pathlib import Path
from typing import List, Dict, Tuple

import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio

# Optional explainability / reporting deps (kept offline)
try:
    import shap  # type: ignore
    SHAP_AVAILABLE = True
except Exception:
    SHAP_AVAILABLE = False

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

_DARK_TEMPLATE = go.layout.Template(
    layout=go.Layout(
        template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Inter, sans-serif"),
        colorway=["#CFB87C", "#22c55e", "#f59e0b", "#ef4444", "#565A5C", "#06b6d4", "#ec4899", "#14b8a6"],
    )
)
pio.templates["cu_dark"] = _DARK_TEMPLATE
pio.templates.default = "plotly_dark+cu_dark"

from sklearn.ensemble import (
    IsolationForest,
    RandomForestClassifier,
    GradientBoostingClassifier,
    ExtraTreesClassifier,
    AdaBoostClassifier,
    RandomForestRegressor,
    GradientBoostingRegressor,
)
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import cross_val_score, StratifiedKFold
from sklearn.metrics import (
    accuracy_score,
    roc_auc_score,
    mean_absolute_error,
    r2_score,
    confusion_matrix,
    roc_curve,
    auc,
    classification_report,
)
from sklearn.linear_model import LogisticRegression, Ridge
from sklearn.neural_network import MLPClassifier, MLPRegressor
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import GaussianNB
from sklearn.svm import SVC
from sklearn.pipeline import Pipeline
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# RAG uses TF-IDF for fast keyword search (no heavy model downloads needed)
SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from xgboost import XGBClassifier as _xgb_cls  # type: ignore
except ImportError:
    _xgb_cls = None


# ----------------------------
# CU BOULDER VISUAL IDENTITY COLORS
# ----------------------------
CU_GOLD = "#CFB87C"
CU_BLACK = "#000000"
CU_DARK_GRAY = "#565A5C"
CU_LIGHT_GRAY = "#A2A4A3"
CU_CHART_COLORS = [CU_GOLD, "#565A5C", "#A2A4A3", "#d4a843", "#333333", "#888888"]

# Column-name prettifier: snake_case → Title Case
_PRETTY_NAMES = {
    "player_id": "Player ID", "total_player_load": "Player Load",
    "player_load_per_min_est": "Load/Min", "total_distance": "Distance (mi)",
    "maximum_velocity": "Max Velocity (mph)", "max_acceleration": "Max Accel (m/s²)",
    "max_deceleration": "Max Decel (m/s²)", "total_acceleration_load": "Accel Load",
    "explosive_efforts": "Explosive Efforts", "hsd_m": "High-Speed Dist (m)",
    "sprint_m": "Sprint Dist (m)", "acwr_ewma_7_28": "ACWR",
    "readiness": "Readiness", "fatigue_idx": "Fatigue Index",
    "wellness_total": "Wellness Score", "mental_score": "Mental",
    "physical_score": "Physical", "sleep_score": "Sleep",
    "soreness_score": "Soreness", "injury_status_actual": "Injury Status",
    "injury_status": "Injury Status", "anomaly_score": "Anomaly Score",
    "anomaly_flag": "Flagged", "vmax_zone": "Speed Zone",
    "speed_band": "Speed Band", "vmax_pct_pb": "% of Personal Best",
    "vmax_personal_best": "Personal Best (mph)", "position_name": "Position",
    "date": "Date", "year": "Year", "session_classification": "Session Type",
    "days_in_status": "Days in Status", "injury_date": "Injury Date",
    "status_start": "Status Start", "wellness_date": "Wellness Date",
    "pos_group": "Position Group", "flag_label": "Status",
    "match_hs_flag": "Hit HS Threshold", "match_sprint_flag": "Hit Sprint Threshold",
}

def pretty(col_name: str) -> str:
    """Convert a snake_case column name to a display-friendly label."""
    return _PRETTY_NAMES.get(col_name, col_name.replace("_", " ").title())

def pretty_cols(df: pd.DataFrame) -> pd.DataFrame:
    """Rename dataframe columns to display-friendly names."""
    return df.rename(columns={c: pretty(c) for c in df.columns})

# ----------------------------
# PAGE CONFIG (safe for older Streamlit)
# ----------------------------
if hasattr(st, "set_page_config"):
    st.set_page_config(page_title="CU WSOC Performance (2023–2025)", page_icon="⚽", layout="wide")

# Simple “pro” styling
st.markdown(
    """
    <style>
      .block-container {padding-top: 1.2rem; padding-bottom: 2rem;}
      h1, h2, h3 {letter-spacing: 0.2px;}
      [data-testid="stMetricValue"] {font-size: 1.6rem;}
      .stTabs [data-baseweb="tab-list"] {gap: 6px;}
      .stTabs [data-baseweb="tab"] {border-radius: 12px; padding: 8px 14px;}
      /* --- PREMIUM DARK THEME --- */
      @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
      html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
      .block-container { max-width: 1400px; }
      div[data-testid="stMetric"] {
          background: linear-gradient(135deg, #1e1e2f 0%, #2a2a40 100%);
          border: 1px solid rgba(255,255,255,0.08);
          border-radius: 16px; padding: 18px 22px;
          box-shadow: 0 4px 24px rgba(0,0,0,0.25);
      }
      div[data-testid="stMetric"] label { color: #9ca3af; font-size: 0.82rem; font-weight: 500; }
      div[data-testid="stMetric"] [data-testid="stMetricValue"] {
          font-size: 1.8rem; font-weight: 700; color: #f1f5f9;
      }
      .stTabs [data-baseweb="tab-list"] {
          background: #141422; border-radius: 14px; padding: 4px; gap: 4px;
      }
      .stTabs [data-baseweb="tab"] {
          border-radius: 10px; padding: 8px 16px; font-weight: 500; font-size: 0.85rem;
          color: #9ca3af; transition: all 0.2s;
      }
      .stTabs [aria-selected="true"] {
          background: linear-gradient(135deg, #CFB87C, #a8944f) !important;
          color: #000 !important; font-weight: 600;
      }
      .stDataFrame { border-radius: 12px; overflow: hidden; }
      section[data-testid="stSidebar"] {
          background: linear-gradient(180deg, #0f0f1a 0%, #1a1a2e 100%);
      }
      .js-plotly-plot .plotly .main-svg { border-radius: 12px; }
      h1 { letter-spacing: -0.5px; }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """<div style="background: linear-gradient(135deg, #000000 0%, #1a1a1a 50%, #2a2a2a 100%);
    border-radius: 16px; padding: 28px 36px; margin-bottom: 20px;
    border: 2px solid #CFB87C;">
    <h1 style="margin:0; color:#CFB87C; font-size:2.2rem; letter-spacing:-0.5px;">
    ⚽ CU WSOC — Performance Intelligence Dashboard</h1>
    <p style="margin:6px 0 0 0; color:#A2A4A3; font-size:0.95rem;">
    Catapult GPS · Game Stats · Schedule · Injury · Wellness · ML Model Lab · Coach Assistant</p>
    </div>""",
    unsafe_allow_html=True,
)

# ==============================================================================
# PATHS — relative to project root (portable across machines)
# ==============================================================================
_PROJECT_ROOT = Path(__file__).resolve().parent

DEFAULT_CATAPULT = _PROJECT_ROOT / "outputs" / "clean_session_level_with_acwr.csv"

SCHED_2023 = _PROJECT_ROOT / "Data" / "2023" / "2023 Schedule Data" / "2023 Schedule.xlsx"
SCHED_2024 = _PROJECT_ROOT / "Data" / "2024" / "2024 Schedule Data" / "2024 Schedule.xlsx"
SCHED_2025 = _PROJECT_ROOT / "Data" / "2025" / "2025 Schedule Data" / "2025 Schedule.xlsx"

GAME_2023_DIR = _PROJECT_ROOT / "Data" / "2023" / "2023 Game Stat Data"
GAME_2024_DIR = _PROJECT_ROOT / "Data" / "2024" / "2024 Clean Game Stat Data"
GAME_2025_DIR = _PROJECT_ROOT / "Data" / "2025" / "2025 Clean Game Stat Data"

INJ_2023 = _PROJECT_ROOT / "Data" / "2023" / "2023 Injuries.csv"
INJ_2024 = _PROJECT_ROOT / "Data" / "2024" / "2024 Injuries.csv"
INJ_2025 = _PROJECT_ROOT / "Data" / "2025" / "2025 Injuries.csv"
WELL_2024 = _PROJECT_ROOT / "Data" / "2024" / "2024 Wellness.csv"
WELL_2025 = _PROJECT_ROOT / "Data" / "2025" / "2025 Wellness.csv"


# ==============================================================================
# HELPERS
# ==============================================================================
def normalize_opponent_token(s: str) -> str:
    if pd.isna(s):
        return ""
    x = str(s).lower().strip()
    x = re.sub(r"[^a-z0-9]+", "_", x)
    x = re.sub(r"_+", "_", x).strip("_")
    return x


@st.cache_data(show_spinner=False)
def load_catapult_csv(path: str) -> pd.DataFrame:
    df = pd.read_csv(path, low_memory=False)

    if "date" in df.columns:
        df["date"] = pd.to_datetime(df["date"], errors="coerce")

    numeric_cols = [
        "total_player_load", "total_distance", "maximum_velocity",
        "max_acceleration", "max_deceleration", "explosive_efforts",
        "total_acceleration_load", "hsd_m", "sprint_m",
        "period_duration_min", "player_load_per_min_est",
        "ewma7", "ewma28", "acwr_ewma_7_28", "baseline_chronic", "pct_vs_baseline",
        "pct_of_hist_proxy", "max_vel_hist_proxy",
        # if you created extra engineered columns, they’ll just coerce if present:
        "distance_per_min", "accel_load_per_min", "accel_to_distance", "load_to_distance",
        "total_player_load_pct_of_p14", "total_distance_pct_of_p14",
        "maximum_velocity_pct_of_p14", "total_acceleration_load_pct_of_p14",
    ]
    for c in numeric_cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    for c in ["player_id", "team_name", "position_name", "session_classification", "workload_col_used"]:
        if c in df.columns:
            df[c] = df[c].astype(str)

    df = df[df["date"].notna()].copy()
    df["year"] = df["date"].dt.year

    # ── UNIT CORRECTIONS (per CU WSOC Data Dictionary) ──
    # maximum_velocity  = MPH  (NOT m/s). CU top ≈ 19-20 mph.
    # total_distance    = Miles (NOT meters).
    # Cap velocity: >= 22 mph is GPS error → NaN
    if "maximum_velocity" in df.columns:
        df.loc[df["maximum_velocity"] >= 22.0, "maximum_velocity"] = np.nan
    if "total_distance" in df.columns:
        df.loc[df["total_distance"] >= 10.0, "total_distance"] = np.nan

    # 3-sigma outlier filtering on key metrics
    for _col in ["total_player_load", "total_distance", "maximum_velocity", "total_acceleration_load"]:
        if _col in df.columns:
            _vals = pd.to_numeric(df[_col], errors="coerce")
            _mean, _std = _vals.mean(), _vals.std()
            if _std > 0:
                df.loc[_vals > _mean + 3 * _std, _col] = np.nan

    return df


# ----------------------------
# ROBUST SCHEDULE LOADER (fixes KeyError: 'Date' / missing date column)
# ----------------------------
def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [re.sub(r"\s+", " ", str(c)).strip().lower() for c in out.columns]
    return out


def find_date_column(df: pd.DataFrame):
    # prefer exact-ish matches, then substrings
    candidates = [
        "date", "match date", "game date", "gamedate", "matchdate",
        "day", "game day", "datetime", "date/time", "date time"
    ]
    cols = list(df.columns)
    # exact / normalized exact
    for c in candidates:
        if c in cols:
            return c
    # substring
    for col in cols:
        for c in candidates:
            if c in col:
                return col
    return None


def find_opponent_column(df: pd.DataFrame):
    candidates = ["opponent", "opp", "vs", "team", "opponent team"]
    cols = list(df.columns)
    for c in candidates:
        if c in cols:
            return c
    for col in cols:
        for c in candidates:
            if c in col:
                return col
    return None


def find_ha_column(df: pd.DataFrame):
    candidates = ["home/away", "home away", "homeaway", "location", "site", "venue", "at"]
    cols = list(df.columns)
    for c in candidates:
        if c in cols:
            return c
    for col in cols:
        for c in candidates:
            if c in col:
                return col
    return None


def find_result_column(df: pd.DataFrame):
    candidates = ["result", "w/l", "outcome", "score"]
    cols = list(df.columns)
    for c in candidates:
        if c in cols:
            return c
    for col in cols:
        for c in candidates:
            if c in col:
                return col
    return None


def parse_schedule_date_any(x, year: int):
    """
    Handles:
    - Excel datetime values / serials
    - strings like 'Aug 10 (Sat)'
    - strings like '8/10' or '08/10/2025'
    """
    if pd.isna(x):
        return pd.NaT

    # already Timestamp
    if isinstance(x, pd.Timestamp):
        return pd.to_datetime(x, errors="coerce")

    # Excel serial (often float/int)
    if isinstance(x, (int, float, np.integer, np.floating)) and not pd.isna(x):
        # Try excel serial origin
        try:
            ts = pd.to_datetime("1899-12-30") + pd.to_timedelta(float(x), unit="D")
            # If it lands in a reasonable season year range, accept
            if 2000 <= ts.year <= 2100:
                return ts
        except Exception:
            pass

    s = str(x).strip()
    if not s:
        return pd.NaT

    # remove "(Sat)" etc
    s2 = re.sub(r"\(.*?\)", "", s).strip()

    ts = pd.to_datetime(s2, errors="coerce")
    if pd.notna(ts):
        # if year missing and parser defaults to 1900/1970, force provided year
        if ts.year in (1900, 1970):
            return pd.to_datetime(f"{s2} {year}", errors="coerce")
        return ts

    return pd.to_datetime(f"{s2} {year}", errors="coerce")


@st.cache_data(show_spinner=False)
def load_schedule_xlsx(xlsx_path: str, year: int) -> pd.DataFrame:
    """
    Bulletproof schedule loader:
    - scans ALL sheets
    - tries multiple header rows (0..15)
    - finds date/opponent columns by flexible matching
    """
    p = Path(str(xlsx_path))
    xl = pd.ExcelFile(p)

    best = None
    best_sheet = None
    best_hdr = None
    best_score = -1

    for sheet in xl.sheet_names:
        raw = pd.read_excel(p, sheet_name=sheet, header=None)

        max_hdr = min(16, raw.shape[0])
        for hdr in range(0, max_hdr):
            df = pd.read_excel(p, sheet_name=sheet, header=hdr)
            if df is None or df.empty:
                continue
            df = normalize_columns(df)

            dcol = find_date_column(df)
            ocol = find_opponent_column(df)
            if dcol is None or ocol is None:
                continue

            tmp = df.copy()
            tmp[dcol] = tmp[dcol].apply(lambda v: parse_schedule_date_any(v, year))
            score = tmp[dcol].notna().sum()
            if score > best_score:
                best_score = score
                best = tmp
                best_sheet = sheet
                best_hdr = hdr

    if best is None or best_score <= 0:
        raise KeyError(
            f"Schedule file missing date column or header not detected: {xlsx_path}\n"
            f"Sheets found: {xl.sheet_names}"
        )

    df = best.copy()

    dcol = find_date_column(df)
    ocol = find_opponent_column(df)
    hcol = find_ha_column(df)
    rcol = find_result_column(df)

    # standard names
    if dcol != "date":
        df = df.rename(columns={dcol: "date"})
    if ocol != "opponent":
        df = df.rename(columns={ocol: "opponent"})
    if hcol and hcol != "home_away":
        df = df.rename(columns={hcol: "home_away"})
    if rcol and rcol != "result":
        df = df.rename(columns={rcol: "result"})

    df = df.dropna(subset=["date"]).copy()
    df["match_date"] = pd.to_datetime(df["date"], errors="coerce")
    df = df.dropna(subset=["match_date"]).copy()

    df["year"] = year
    df["Opponent_clean"] = (
        df["opponent"].astype(str)
        .str.replace(r"^#\d+\s+", "", regex=True)
        .str.strip()
    )

    if "home_away" in df.columns:
        df["venue"] = df["home_away"].astype(str).str.strip()
    else:
        df["venue"] = "Unknown"

    if "result" not in df.columns:
        df["result"] = np.nan

    df["schedule_sheet_used"] = best_sheet
    df["schedule_header_row_used"] = best_hdr

    keep = ["match_date", "year", "Opponent_clean", "venue", "result", "schedule_sheet_used", "schedule_header_row_used"]
    df = df[keep].sort_values("match_date").drop_duplicates(subset=["match_date", "year"], keep="first")
    return df


@st.cache_data(show_spinner=False)
def load_all_schedules(s23: Path, s24: Path, s25: Path) -> pd.DataFrame:
    out = []
    if s23.exists():
        out.append(load_schedule_xlsx(str(s23), 2023))
    if s24.exists():
        out.append(load_schedule_xlsx(str(s24), 2024))
    if s25.exists():
        out.append(load_schedule_xlsx(str(s25), 2025))
    if len(out) == 0:
        return pd.DataFrame(columns=["match_date", "year", "Opponent_clean", "venue", "result"])
    sched = pd.concat(out, ignore_index=True)
    sched = sched.sort_values(["match_date", "year"]).drop_duplicates(subset=["match_date", "year"], keep="first")
    return sched


@st.cache_data(show_spinner=False)
def load_game_stats_folder(folder: str, year: int) -> pd.DataFrame:
    """
    Reads all CSVs inside the folder, adds:
      - file_name
      - opponent_token from file name
      - year
    Your game stat files look like:
      Player ID, Type, No, Pos, MIN, G, A, Sh, SOG, GA, Saves
    """
    folder = Path(folder)
    if not folder.exists():
        return pd.DataFrame()

    files = sorted(list(folder.glob("*.csv")))
    if len(files) == 0:
        return pd.DataFrame()

    frames = []
    for f in files:
        try:
            df = pd.read_csv(f)
        except Exception:
            df = pd.read_csv(f, sep=",", engine="python")

        df.columns = [str(c).strip() for c in df.columns]

        if "Player ID" in df.columns:
            df = df.rename(columns={"Player ID": "player_id"})
        elif "player_id" not in df.columns:
            continue

        df["player_id"] = df["player_id"].astype(str).str.strip()
        df["file_name"] = f.name

        # Opponent token from file name: "2024_baylor_game_stat.csv" -> "baylor"
        base = f.stem
        base = re.sub(rf"^{year}_", "", base)
        base = re.sub(r"_game.*$", "", base)
        df["opponent_token"] = normalize_opponent_token(base)
        df["year"] = year
        frames.append(df)

    if len(frames) == 0:
        return pd.DataFrame()

    return pd.concat(frames, ignore_index=True)


def safe_numeric(df, cols):
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    return df


def schedule_to_token_map(schedule_df: pd.DataFrame) -> pd.DataFrame:
    s = schedule_df.copy()
    s["opponent_token"] = s["Opponent_clean"].apply(normalize_opponent_token)
    return s


@st.cache_data(show_spinner=False)
def build_game_match_table(schedule_df: pd.DataFrame, g23: pd.DataFrame, g24: pd.DataFrame, g25: pd.DataFrame) -> pd.DataFrame:
    frames = []
    for g in [g23, g24, g25]:
        if g is not None and len(g) > 0:
            frames.append(g)
    if len(frames) == 0:
        return pd.DataFrame()

    games = pd.concat(frames, ignore_index=True)
    sched = schedule_to_token_map(schedule_df)

    merged = games.merge(
        sched[["year", "match_date", "Opponent_clean", "venue", "result", "opponent_token"]],
        how="left",
        on=["year", "opponent_token"]
    )
    merged = merged.sort_values(["year", "file_name", "player_id"])
    return merged


# ----------------------------
# INJURY + WELLNESS LOADERS
# ----------------------------
@st.cache_data(show_spinner=False)
def load_injury_csv(path: str, year: int) -> pd.DataFrame:
    p = Path(path)
    if not p.exists():
        return pd.DataFrame()
    df = pd.read_csv(p)
    df.columns = [str(c).strip() for c in df.columns]
    rename = {}
    for c in df.columns:
        cl = c.lower()
        if "player" in cl and "id" in cl:
            rename[c] = "player_id"
        elif cl == "date":
            rename[c] = "report_date"
        elif "date of injury" in cl or "onset" in cl:
            rename[c] = "injury_date"
        elif cl == "status":
            rename[c] = "injury_status"
        elif "start of status" in cl:
            rename[c] = "status_start"
        elif "days in status" in cl:
            rename[c] = "days_in_status"
    df = df.rename(columns=rename)
    if "player_id" not in df.columns:
        return pd.DataFrame()
    df["player_id"] = df["player_id"].astype(str).str.strip()
    for dc in ["report_date", "injury_date", "status_start"]:
        if dc in df.columns:
            df[dc] = pd.to_datetime(df[dc], errors="coerce")
    if "days_in_status" in df.columns:
        df["days_in_status"] = pd.to_numeric(df["days_in_status"], errors="coerce")
    if "injury_status" in df.columns:
        df["injury_status"] = df["injury_status"].astype(str).str.strip()
    df["year"] = year
    return df


@st.cache_data(show_spinner=False)
def load_all_injuries(i23, i24, i25) -> pd.DataFrame:
    frames = []
    for p, y in [(i23, 2023), (i24, 2024), (i25, 2025)]:
        if Path(p).exists():
            tmp = load_injury_csv(str(p), y)
            if len(tmp) > 0:
                frames.append(tmp)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


@st.cache_data(show_spinner=False)
def load_wellness_csv(path: str, year: int) -> pd.DataFrame:
    p = Path(path)
    if not p.exists():
        return pd.DataFrame()
    df = pd.read_csv(p)
    df.columns = [str(c).strip() for c in df.columns]
    rename = {}
    for c in df.columns:
        cl = c.lower()
        if "player" in cl and "id" in cl:
            rename[c] = "player_id"
        elif cl == "date":
            rename[c] = "wellness_date"
        elif "mental" in cl:
            rename[c] = "mental_score"
        elif "physical" in cl:
            rename[c] = "physical_score"
        elif "sleep" in cl:
            rename[c] = "sleep_score"
        elif "soreness" in cl:
            rename[c] = "soreness_score"
        elif "overall" in cl:
            rename[c] = "wellness_total"
    df = df.rename(columns=rename)
    if "player_id" not in df.columns:
        return pd.DataFrame()
    df["player_id"] = df["player_id"].astype(str).str.strip()
    if "wellness_date" in df.columns:
        df["wellness_date"] = pd.to_datetime(df["wellness_date"], errors="coerce")
    for sc in ["mental_score", "physical_score", "sleep_score", "soreness_score", "wellness_total"]:
        if sc in df.columns:
            df[sc] = pd.to_numeric(df[sc], errors="coerce")
    df["year"] = year
    return df


@st.cache_data(show_spinner=False)
def load_all_wellness(w24, w25) -> pd.DataFrame:
    frames = []
    for p, y in [(w24, 2024), (w25, 2025)]:
        if Path(p).exists():
            tmp = load_wellness_csv(str(p), y)
            if len(tmp) > 0:
                frames.append(tmp)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def merge_injury_status(df, injuries_df):
    out = df.copy()
    out["injury_status_actual"] = "Available"
    if injuries_df.empty or "status_start" not in injuries_df.columns:
        return out
    inj = injuries_df[["player_id", "status_start", "injury_status"]].copy()
    inj = inj.dropna(subset=["status_start"]).sort_values("status_start")
    out = out.sort_values("date")
    merged = pd.merge_asof(
        out[["player_id", "date"]].reset_index(),
        inj.rename(columns={"status_start": "date"}),
        on="date", by="player_id", direction="backward"
    )
    status_map = merged.set_index("index")["injury_status"]
    out["injury_status_actual"] = status_map.reindex(out.index).fillna("Available")
    full_go = out["injury_status_actual"].str.lower().isin(["full go", "available", "nan", ""])
    out.loc[full_go, "injury_status_actual"] = "Available"
    return out


def merge_wellness(df, wellness_df):
    out = df.copy()
    score_cols = ["mental_score", "physical_score", "sleep_score", "soreness_score", "wellness_total"]
    for sc in score_cols:
        out[sc] = np.nan
    if wellness_df.empty or "wellness_date" not in wellness_df.columns:
        return out
    well = wellness_df.copy()
    well["wellness_date"] = pd.to_datetime(well["wellness_date"], errors="coerce")
    well = well.dropna(subset=["wellness_date"]).sort_values("wellness_date")
    keep = ["player_id", "wellness_date"] + [c for c in score_cols if c in well.columns]
    well = well[keep].drop_duplicates(subset=["player_id", "wellness_date"], keep="last")
    out = out.sort_values("date")
    merged = pd.merge_asof(
        out[["player_id", "date"]].reset_index(),
        well.rename(columns={"wellness_date": "date"}),
        on="date", by="player_id", direction="backward",
        tolerance=pd.Timedelta(days=3)
    )
    for sc in [c for c in score_cols if c in merged.columns]:
        out[sc] = merged.set_index("index")[sc].reindex(out.index).values
    return out


# ----------------------------
# METRIC INFO TOOLTIPS
# ----------------------------
METRIC_INFO = {
    "total_player_load": (
        "Total Player Load",
        "Accumulated body load from the Catapult accelerometer (triaxial). "
        "It combines all accelerations, decelerations, and changes of direction into one number.\n\n"
        "**Formula:** Sum of instantaneous accelerometer magnitudes across the session.\n\n"
        "**Example:** A player with a load of 450 had a moderate-intensity session. "
        "Squad average is typically ~300–500. Values >600 indicate a very demanding session."
    ),
    "maximum_velocity": (
        "Maximum Velocity (mph)",
        "The peak speed recorded by the GPS unit during the session. Reported in **miles per hour (MPH)**.\n\n"
        "**GPS cap:** Values ≥ 22 mph are removed as GPS errors. CU WSOC typical top speeds are 16–20 mph.\n\n"
        "**Example:** A forward recording 18.5 mph is performing well. "
        "A drop >10% from personal best signals potential fatigue."
    ),
    "acwr": (
        "ACWR (Acute:Chronic Workload Ratio)",
        "Compares recent load (7-day EWMA) to long-term load (28-day EWMA). "
        "Helps identify load spikes that increase injury risk.\n\n"
        "**Formula:** ACWR = EWMA₇ / EWMA₂₈\n\n"
        "**Zones:**\n"
        "- < 0.8 → Under-loaded (losing fitness)\n"
        "- 0.8–1.3 → Sweet spot (optimal)\n"
        "- > 1.3 → Danger zone (injury risk ↑)\n\n"
        "**Example:** ACWR = 1.1 means recent load is 10% above the chronic baseline — ideal."
    ),
    "readiness": (
        "Readiness Score (0–100)",
        "A composite score estimating how prepared a player is for training/competition. "
        "Combines up to 5 components, each scored 0–100, then averaged.\n\n"
        "**Components:**\n"
        "1. **ACWR zone** — 0.8–1.3 = high score, outside = penalized\n"
        "2. **Velocity freshness** — recent avg speed ÷ personal best × 100\n"
        "3. **Anomaly rate** — fewer anomalies in last 14 sessions = higher score\n"
        "4. **Load stability** — lower CV (coefficient of variation) of last 7 sessions = higher score\n"
        "5. **Wellness** — latest wellness total ÷ 20 × 100\n\n"
        "**Example:** A player with ACWR=1.05, speed at 90% PB, no anomalies, stable load, "
        "wellness 16/20 → readiness ≈ 82/100."
    ),
    "fatigue_idx": (
        "Fatigue Index",
        "Ratio of recent average load to longer-term average. Detects acute fatigue build-up.\n\n"
        "**Formula:** (7-session rolling mean load) ÷ (28-session rolling mean load)\n\n"
        "**Interpretation:**\n"
        "- ~1.0 → normal load balance\n"
        "- > 1.5 → significant acute fatigue accumulation\n"
        "- < 0.7 → recent load is well below chronic level (detraining risk)\n\n"
        "**Example:** Fatigue index = 1.3 means the last 7 sessions averaged 30% more load than "
        "the 28-session average — monitor for recovery needs."
    ),
    "anomaly_score": (
        "Anomaly Detection",
        "An Isolation Forest model flags sessions that are unusual across multiple metrics simultaneously "
        "(load, speed, distance, acceleration, ACWR).\n\n"
        "**How it works:** The algorithm builds random decision trees. Sessions that are isolated quickly "
        "(i.e., are far from normal patterns) get high anomaly scores.\n\n"
        "**Sensitivity slider:** Controls what % of sessions are flagged (default ~4%).\n\n"
        "**Example:** A session with load=700, speed=12mph, ACWR=1.8 would likely be flagged — "
        "high load combined with low speed and spiked ACWR is an unusual combination."
    ),
    "wellness": (
        "Wellness Score (0–20)",
        "Self-reported wellness consisting of 4 sub-scores, each rated 1–5.\n\n"
        "**Sub-scores:**\n"
        "- Mental score (1–5)\n"
        "- Physical score (1–5)\n"
        "- Sleep score (1–5)\n"
        "- Soreness score (1–5)\n\n"
        "**Formula:** Wellness Total = Mental + Physical + Sleep + Soreness\n\n"
        "**Example:** A player reporting Mental=4, Physical=3, Sleep=4, Soreness=3 → Total = 14/20. "
        "Scores below 12/20 warrant a check-in."
    ),
    "monotony_strain": (
        "Monotony & Strain (Foster, 1998)",
        "Classic load monitoring metrics introduced by Foster et al. (1998) in *Monitoring Training "
        "in Athletes with Reference to Overtraining Syndrome*, Exercise Immunology Review.\n\n"
        "**Monotony** = Mean daily load ÷ SD of daily loads (within one week).\n"
        "- Monotony > 2.0 indicates overly uniform training with insufficient variation between "
        "hard and easy days. This has been linked to increased illness risk.\n"
        "- Monotony is capped at 10.0 in this dashboard; weeks with fewer than 3 training days "
        "are excluded since SD is unreliable.\n\n"
        "**Strain** = Weekly total load × Monotony.\n"
        "- High strain weeks (especially > P90) are associated with illness and soft-tissue "
        "injury in the following 7–14 days (Foster, 1998; Gabbett, 2010).\n\n"
        "**Practical application:** Vary session intensity across the week. Include at least "
        "one high-intensity and one recovery session to keep monotony below 2.0.\n\n"
        "**References:**\n"
        "- Foster C (1998). *Exerc Immunol Rev*, 4, 40–51.\n"
        "- Gabbett TJ (2010). *J Sports Sci*, 28(12), 1255–1264."
    ),
    "speed_bands": (
        "Speed Bands (CU Data Dictionary)",
        "Absolute velocity bands defined by the CU WSOC Catapult Data Dictionary.\n\n"
        "- **Band 1:** 0–9.6 mph (walking/jogging)\n"
        "- **Band 2:** 9.6–12 mph (running)\n"
        "- **Band 3:** 12–14.4 mph (high-speed running)\n"
        "- **Band 4:** 14.4+ mph (sprinting)\n\n"
        "**Example:** ~97% of training distance typically falls in Bands 1+2. "
        "Only ~3% is at or above 12 mph. Forwards tend to have slightly more Band 4 exposure."
    ),
    "md_labels": (
        "Matchday Window Labels",
        "Each training session is labeled relative to the nearest match date.\n\n"
        "- **MD** = Match Day\n"
        "- **MD-1** = 1 day before match (should be low volume, activation only)\n"
        "- **MD-2** = 2 days before (typically highest intensity day)\n"
        "- **MD-3/MD-4** = Primary loading days\n"
        "- **MD+1** = Day after match (active recovery)\n\n"
        "**Example:** If a match is on Saturday, Friday's session is labeled MD-1, "
        "Thursday is MD-2, and Sunday is MD+1."
    ),
    "injury_status": (
        "Injury Status Categories",
        "Each player's participation level based on their injury/health status.\n\n"
        "- **Full Go** — No restrictions, fully cleared\n"
        "- **As Tolerated** — Can participate but with modifications at athlete's discretion\n"
        "- **Limited** — Restricted participation, specific activities only\n"
        "- **Out** — Cannot participate in any training/competition\n\n"
        "**Example:** A player transitioning Out → Limited → As Tolerated → Full Go "
        "follows a graded return-to-play protocol."
    ),
    "player_load_per_min": (
        "Player Load per Minute (Intensity)",
        "Training intensity metric — how much load is generated per minute of activity.\n\n"
        "**Formula:** Total Player Load ÷ Session Duration (minutes)\n\n"
        "**Why it matters:** Two sessions can have the same total load, but different intensities. "
        "A 30-min session with load 300 (10/min) is much more intense than a 90-min session "
        "with load 300 (3.3/min).\n\n"
        "**Example:** Load/min of 8+ indicates high intensity. Below 4 is typically warm-up/recovery."
    ),
    "risk_board": (
        "Player Risk Board (ML Predictions)",
        "Machine learning predicts the probability that each player will be unavailable "
        "for their next session, based on current workload, wellness, injury status, and ACWR.\n\n"
        "**Traffic-light system:**\n"
        "- 🔴 **RED** (≥60% risk) — Reduce load 20–30%, boost recovery\n"
        "- 🟡 **AMBER** (35–60% risk) — Hold or reduce load 10–15%, monitor closely\n"
        "- 🟢 **GREEN** (<35% risk) — Continue planned load\n\n"
        "**Example:** A player currently 'Limited' with ACWR=1.5 and low wellness (10/20) "
        "would likely be flagged RED."
    ),
    "radar_chart": (
        "Player Radar Chart",
        "Compares a player's latest session metrics to the squad distribution. "
        "Each axis is normalized 0–100 using the squad's 10th and 90th percentiles.\n\n"
        "**Formula per axis:** Score = 100 × (player value − P10) ÷ (P90 − P10), clamped to 0–100.\n\n"
        "**Example:** If squad load P10=200 and P90=600, a player with load 500 scores "
        "(500−200)÷(600−200)×100 = 75/100 on the Load axis."
    ),
    "forecast": (
        "Next-Session Forecast",
        "Predicts a player's next session metric using lagged features from their recent history.\n\n"
        "**Features used:** Last N session values, 3-session rolling mean/std, 7-session rolling mean, "
        "and percentage change.\n\n"
        "**Models tested:** Ridge Regression, Random Forest, Gradient Boosting, and MLP Neural Net. "
        "The best model (lowest MAE) is selected automatically.\n\n"
        "**Example:** If a player's last 5 loads were [400, 420, 380, 450, 410], "
        "the model might predict ~415 for the next session. "
        "If actual is 550, that's flagged as a spike."
    ),
    "match_readiness_ratio": (
        "Training vs Match-Day Readiness Ratio",
        "Based on the **training-competition match principle** (Gabbett, 2016): athletes should be "
        "exposed to loads in training that approximate the physical demands of competition. Under-exposure "
        "leaves athletes physically unprepared for worst-case game scenarios.\n\n"
        "**Formula:** Ratio = Player's Avg Training Load ÷ Player's Avg Match-Day Load\n\n"
        "**Threshold:**\n"
        "- < 0.5 → **Under-prepared** — training load is less than half of match-day demands. "
        "These athletes have not been sufficiently exposed to game-intensity efforts and may face "
        "higher injury risk when match demands exceed their preparation (Gabbett, BJSM 2016).\n"
        "- ≥ 0.5 → Adequate preparation relative to match demands.\n\n"
        "**Note:** This ratio is not used to flag 'over-trained' players because 2nd/3rd string "
        "athletes will naturally have lower game loads than practice loads, which would produce "
        "artificially high ratios that do not indicate over-training.\n\n"
        "**References:**\n"
        "- Gabbett TJ (2016). The training-injury prevention paradox. *BJSM*, 50(5), 273–280.\n"
        "- Foster C (1998). Monitoring training in athletes. *Exerc Immunol Rev*, 4, 40–51."
    ),
    "position_groups": (
        "Position Group Classification",
        "Players are grouped by their Catapult position tag into 4 main groups.\n\n"
        "- **GK** — Goalkeepers (goal, keeper)\n"
        "- **DEF** — Defenders (back, cb, rb, lb, fb)\n"
        "- **MID** — Midfielders (cm, cdm, cam, dm, am)\n"
        "- **FWD** — Forwards (forward, wing, striker, att, st, cf, lw, rw)\n\n"
        "**Why it matters:** Each position has different physical demands. "
        "Forwards typically have higher sprint exposure; defenders have higher deceleration load."
    ),
}


def metric_tooltip(key: str):
    """Render a small ℹ️ expander with the metric explanation."""
    if key in METRIC_INFO:
        title, body = METRIC_INFO[key]
        with st.expander(f"ℹ️ How is **{title}** calculated?", expanded=False):
            st.markdown(body)


# ----------------------------
# ANOMALIES
# ----------------------------
def add_anomaly_scores(df: pd.DataFrame, features: list, contamination: float) -> pd.DataFrame:
    out = df.copy()
    out["anomaly_score"] = np.nan
    out["anomaly_flag"] = False

    use = [c for c in features if c in out.columns]
    if len(use) < 4:
        return out

    X = out[use].copy()
    mask = X.notna().all(axis=1)
    if mask.sum() < 50:
        return out

    scaler = StandardScaler()
    Xs = scaler.fit_transform(X.loc[mask].values)

    iso = IsolationForest(n_estimators=400, contamination=contamination, random_state=42)
    iso.fit(Xs)

    scores = -iso.decision_function(Xs)
    out.loc[mask, "anomaly_score"] = scores
    thr = np.nanquantile(out["anomaly_score"], 1 - contamination)
    out["anomaly_flag"] = out["anomaly_score"] >= thr
    return out


# ----------------------------
# READINESS / FATIGUE (coach-facing indices)
# ----------------------------
def compute_fatigue_index(df: pd.DataFrame) -> pd.DataFrame:
    """Rolling fatigue index = 7-session rolling mean / 28-session rolling mean (per player)."""
    out = df.copy()
    out["fatigue_idx"] = np.nan
    if out.empty or "player_id" not in out.columns or "date" not in out.columns or "total_player_load" not in out.columns:
        return out
    out = out.sort_values(["player_id", "date"])
    g = out.groupby("player_id", sort=False)["total_player_load"]
    r7 = g.transform(lambda x: x.fillna(0).rolling(7, min_periods=3).mean())
    r28 = g.transform(lambda x: x.fillna(0).rolling(28, min_periods=7).mean())
    out["fatigue_idx"] = (r7 / r28.replace(0, np.nan)).values
    return out


def compute_readiness(df: pd.DataFrame) -> pd.DataFrame:
    """
    Composite 0-100 readiness score per player, using what is available in the dataset:
    - ACWR zone (optimal 0.8–1.3)
    - Velocity freshness (recent vmax vs personal best)
    - Anomaly rate (lower is better)
    - Load stability (lower CV is better)
    - Wellness (if available; higher is better)
    """
    out = df.copy()
    out["readiness"] = np.nan
    if out.empty or "player_id" not in out.columns or "date" not in out.columns:
        return out

    out = out.sort_values(["player_id", "date"])
    squad_mean_load = float(out["total_player_load"].mean()) if "total_player_load" in out.columns else np.nan

    for pid, p in out.groupby("player_id", sort=False):
        p = p.sort_values("date")
        if len(p) < 5:
            continue

        scores = []

        # ACWR component
        if "acwr_ewma_7_28" in p.columns:
            acwr = p["acwr_ewma_7_28"].dropna()
            if len(acwr) > 0:
                latest = float(acwr.iloc[-1])
                if 0.8 <= latest <= 1.3:
                    s = 90 + 10 * (1 - abs(latest - 1.05) / 0.5)
                elif latest < 0.8:
                    s = max(40, 80 * (latest / 0.8))
                else:
                    s = max(10, 90 - 40 * (latest - 1.3))
                scores.append(float(np.clip(s, 0, 100)))

        # Velocity freshness
        if "maximum_velocity" in p.columns:
            vel = pd.to_numeric(p["maximum_velocity"], errors="coerce").dropna()
            if len(vel) >= 3:
                pb = float(vel.max())
                recent = float(vel.iloc[-3:].mean())
                if np.isfinite(pb) and pb > 0:
                    scores.append(float(np.clip(100 * (recent / pb), 0, 100)))

        # Anomaly rate (last 14 sessions)
        if "anomaly_flag" in p.columns:
            recent_anom = p["anomaly_flag"].iloc[-14:]
            rate = float(recent_anom.mean()) if len(recent_anom) else 0.0
            scores.append(float(np.clip(100 * (1 - rate * 3), 0, 100)))

        # Load stability (CV of last 7 sessions)
        if "total_player_load" in p.columns:
            recent_load = pd.to_numeric(p["total_player_load"], errors="coerce").dropna().iloc[-7:]
            if len(recent_load) >= 3 and float(recent_load.mean()) > 0:
                cv = float(recent_load.std() / recent_load.mean())
                scores.append(float(np.clip(100 * (1 - cv), 0, 100)))
            elif np.isfinite(squad_mean_load) and np.isfinite(float(p["total_player_load"].iloc[-1])) and squad_mean_load > 0:
                # fallback: encourage staying near squad mean
                s = 100 * (1 - abs(float(p["total_player_load"].iloc[-1]) - squad_mean_load) / (2 * squad_mean_load))
                scores.append(float(np.clip(s, 0, 100)))

        # Wellness component (out of 20 -> 0..100)
        if "wellness_total" in p.columns:
            w = pd.to_numeric(p["wellness_total"], errors="coerce").dropna()
            if len(w) > 0:
                scores.append(float(np.clip(100 * (float(w.iloc[-1]) / 20.0), 0, 100)))

        if scores:
            out.loc[p.index, "readiness"] = float(np.mean(scores))

    return out


def compute_velocity_zones(df: pd.DataFrame) -> pd.DataFrame:
    """Assign velocity zones using TWO systems:
    1) Absolute MPH bands from CU WSOC Data Dictionary:
       Band 1: 0–9.6 MPH, Band 2: 9.6–12 MPH, Band 3: 12–14.38 MPH, Band 4: 14.38+ MPH
    2) % of personal best (match-based position-specific readiness):
       Z1 <70%, Z2 70-80%, Z3 80-90%, Z4 90-95%, Z5 95-100%
    """
    out = df.copy()
    out["vmax_personal_best"] = np.nan
    out["vmax_pct_pb"] = np.nan
    out["vmax_zone"] = "UNK"
    out["speed_band"] = "UNK"
    if out.empty or "player_id" not in out.columns or "maximum_velocity" not in out.columns:
        return out

    vel = pd.to_numeric(out["maximum_velocity"], errors="coerce")

    # Absolute MPH bands (Data Dictionary)
    out.loc[vel.notna() & (vel < 9.6), "speed_band"] = "Band 1 (0–9.6 mph)"
    out.loc[vel.notna() & (vel >= 9.6) & (vel < 12.0), "speed_band"] = "Band 2 (9.6–12 mph)"
    out.loc[vel.notna() & (vel >= 12.0) & (vel < 14.38), "speed_band"] = "Band 3 (12–14.4 mph)"
    out.loc[vel.notna() & (vel >= 14.38), "speed_band"] = "Band 4 (14.4+ mph)"

    # % of personal best zones (match readiness)
    pb = vel.groupby(out["player_id"]).transform("max")
    out["vmax_personal_best"] = pb
    pct = 100 * (vel / pb.replace(0, np.nan))
    out["vmax_pct_pb"] = pct

    out.loc[pct.notna() & (pct < 70), "vmax_zone"] = "Z1 (<70%)"
    out.loc[pct.notna() & (pct >= 70) & (pct < 80), "vmax_zone"] = "Z2 (70–80%)"
    out.loc[pct.notna() & (pct >= 80) & (pct < 90), "vmax_zone"] = "Z3 (80–90%)"
    out.loc[pct.notna() & (pct >= 90) & (pct < 95), "vmax_zone"] = "Z4 (90–95%)"
    out.loc[pct.notna() & (pct >= 95), "vmax_zone"] = "Z5 (95–100%)"

    # Match-based position-specific thresholds (from Capstone: 67 validated matches)
    _match_thresh = {"DEF": (17.7, 18.2), "FWD": (17.7, 19.0), "MID": (17.8, 18.4), "GK": (15.0, 15.9)}
    out["match_hs_flag"] = False
    out["match_sprint_flag"] = False
    if "position_name" in out.columns:
        def _pg(p):
            pl = str(p).lower()
            if any(k in pl for k in ["gk", "keeper", "goal"]): return "GK"
            if any(k in pl for k in ["def", "back", "cb", "rb", "lb", "fb"]): return "DEF"
            if any(k in pl for k in ["mid", "cm", "cdm", "cam", "dm", "am"]): return "MID"
            if any(k in pl for k in ["fwd", "forward", "striker", "wing", "att", "st", "cf", "lw", "rw"]): return "FWD"
            return "OTHER"
        _grps = out["position_name"].apply(_pg)
        for grp, (hs, sp) in _match_thresh.items():
            mask = (_grps == grp) & vel.notna()
            out.loc[mask & (vel >= hs), "match_hs_flag"] = True
            out.loc[mask & (vel >= sp), "match_sprint_flag"] = True

    return out


def summarize_player_drivers(p: pd.DataFrame) -> list:
    """
    Return a ranked list of coach-facing "what's driving this player's status today?" items.
    Uses only information already computed/merged into the training dataframe.
    """
    if p is None or len(p) == 0:
        return []
    r = p.sort_values("date").iloc[-1]

    drivers = []

    def _add(score: float, title: str, detail: str):
        drivers.append({"score": float(score), "title": str(title), "detail": str(detail)})

    # 1) Injury status (highest priority)
    st_val = r.get("injury_status_actual", None)
    if pd.notna(st_val):
        st_s = str(st_val).strip()
        if st_s and st_s.lower() not in ["available", "nan", ""]:
            base = 100.0 if st_s.lower() == "out" else 85.0 if st_s.lower() == "limited" else 75.0
            _add(base, "Injury / availability status", f"Current status: **{st_s}** (from injury log)")

    # 2) ACWR
    acwr = pd.to_numeric(r.get("acwr_ewma_7_28", np.nan), errors="coerce")
    if np.isfinite(acwr):
        if acwr > 1.3:
            _add(80.0, "ACWR spike risk", f"ACWR **{acwr:.2f}** (> 1.30 red zone)")
        elif acwr < 0.8:
            _add(55.0, "Possible under-loading", f"ACWR **{acwr:.2f}** (< 0.80 under-load zone)")

    # 3) Fatigue index
    fat = pd.to_numeric(r.get("fatigue_idx", np.nan), errors="coerce")
    if np.isfinite(fat):
        if fat > 1.5:
            _add(70.0, "High fatigue index", f"Fatigue index **{fat:.2f}** (> 1.50)")
        elif fat > 1.2:
            _add(40.0, "Rising fatigue", f"Fatigue index **{fat:.2f}** (> 1.20)")

    # 4) Wellness + sub-scores (if present)
    wt = pd.to_numeric(r.get("wellness_total", np.nan), errors="coerce")
    if np.isfinite(wt):
        if wt < 12:
            _add(65.0, "Low wellness", f"Wellness **{wt:.0f}/20** (below 12/20)")
        elif wt < 14:
            _add(35.0, "Moderate wellness", f"Wellness **{wt:.0f}/20**")

    sl = pd.to_numeric(r.get("sleep_score", np.nan), errors="coerce")
    if np.isfinite(sl) and sl < 3:
        _add(45.0, "Sleep concern", f"Sleep **{sl:.0f}/5** (below 3/5)")

    ms = pd.to_numeric(r.get("mental_score", np.nan), errors="coerce")
    if np.isfinite(ms) and ms < 3:
        _add(40.0, "Mental score concern", f"Mental **{ms:.0f}/5** (below 3/5)")

    # 5) Anomaly flag
    af = r.get("anomaly_flag", False)
    try:
        is_flag = bool(af)
    except Exception:
        is_flag = False
    if is_flag:
        ascore = pd.to_numeric(r.get("anomaly_score", np.nan), errors="coerce")
        if np.isfinite(ascore):
            _add(50.0, "Unusual session pattern", f"Flagged anomaly (score **{ascore:.3f}**)")
        else:
            _add(50.0, "Unusual session pattern", "Flagged anomaly session")

    # 6) Velocity freshness (recent 3 vs PB)
    if "maximum_velocity" in p.columns:
        vel = pd.to_numeric(p["maximum_velocity"], errors="coerce").dropna()
        if len(vel) >= 5:
            pb = float(vel.max())
            recent = float(vel.iloc[-3:].mean())
            if np.isfinite(pb) and pb > 0 and np.isfinite(recent):
                pct = 100.0 * recent / pb
                if pct < 90:
                    _add(35.0, "Velocity freshness dip", f"Recent speed ~**{pct:.0f}%** of PB (last 3 sessions)")

    # Rank high → low and dedupe by title
    drivers = sorted(drivers, key=lambda x: x["score"], reverse=True)
    seen = set()
    out = []
    for d in drivers:
        if d["title"] in seen:
            continue
        seen.add(d["title"])
        out.append(d)
    return out


def explain_risk_row(row) -> str:
    """
    Build a short, coach-facing reason string for why a player is RED / AMBER / GREEN
    on the Coach ML Player Risk Board. Uses the actual merged values on the row.
    """
    try:
        flag = str(row.get("coach_flag", "")).upper()
    except Exception:
        flag = ""

    reasons = []

    st_val = row.get("injury_status_actual", None)
    if pd.notna(st_val):
        st_s = str(st_val).strip()
        if st_s and st_s.lower() not in ["available", "nan", ""]:
            reasons.append(f"status **{st_s}**")

    acwr = pd.to_numeric(row.get("acwr_ewma_7_28", np.nan), errors="coerce")
    if np.isfinite(acwr):
        if acwr > 1.3:
            reasons.append(f"ACWR spike **{acwr:.2f}** (>1.30)")
        elif acwr < 0.8:
            reasons.append(f"ACWR under-load **{acwr:.2f}** (<0.80)")

    fat = pd.to_numeric(row.get("fatigue_idx", np.nan), errors="coerce")
    if np.isfinite(fat):
        if fat > 1.5:
            reasons.append(f"high fatigue **{fat:.2f}** (>1.50)")
        elif fat > 1.2:
            reasons.append(f"rising fatigue **{fat:.2f}**")

    wt = pd.to_numeric(row.get("wellness_total", np.nan), errors="coerce")
    if np.isfinite(wt):
        if wt < 12:
            reasons.append(f"low wellness **{wt:.0f}/20**")
        elif wt < 14:
            reasons.append(f"moderate wellness **{wt:.0f}/20**")

    sl = pd.to_numeric(row.get("sleep_score", np.nan), errors="coerce")
    if np.isfinite(sl) and sl < 3:
        reasons.append(f"sleep low **{sl:.0f}/5**")
    ms = pd.to_numeric(row.get("mental_score", np.nan), errors="coerce")
    if np.isfinite(ms) and ms < 3:
        reasons.append(f"mental low **{ms:.0f}/5**")

    try:
        if bool(row.get("anomaly_flag", False)):
            reasons.append("flagged as unusual session")
    except Exception:
        pass

    prob = pd.to_numeric(row.get("risk_next_non_available", np.nan), errors="coerce")
    prob_txt = f"prob **{prob * 100:.0f}%**" if np.isfinite(prob) else ""

    if flag == "RED":
        head = "🔴 RED"
        if not reasons:
            reasons = ["model sees combined risk signals"]
        return f"{head} — {', '.join(reasons)}" + (f" · {prob_txt}" if prob_txt else "")
    if flag == "AMBER":
        head = "🟡 AMBER"
        if not reasons:
            reasons = ["moderate risk from recent load / wellness trend"]
        return f"{head} — {', '.join(reasons)}" + (f" · {prob_txt}" if prob_txt else "")
    head = "🟢 GREEN"
    return f"{head} — available, stable load & wellness" + (f" · {prob_txt}" if prob_txt else "")


def coach_guide(title: str, bullets: list):
    """Render a short, consistent 'Coach Guide' box at the top of a tab."""
    with st.expander(f"🧑‍🏫 Coach Guide — {title}", expanded=False):
        for b in bullets:
            st.markdown(f"- {b}")


def compute_weekly_monotony_strain(df: pd.DataFrame) -> pd.DataFrame:
    """Weekly team monotony/strain using daily mean loads (classic load monitoring)."""
    if df.empty or "date" not in df.columns or "total_player_load" not in df.columns:
        return pd.DataFrame()
    tmp = df.copy()
    tmp["week"] = pd.to_datetime(tmp["date"]).dt.isocalendar().week.astype(int)
    tmp["year"] = pd.to_datetime(tmp["date"]).dt.year.astype(int)
    daily = tmp.groupby(["year", "week", "date"], as_index=False)["total_player_load"].mean().rename(columns={"total_player_load": "mean_daily"})
    weekly = daily.groupby(["year", "week"], as_index=False).agg(
        week_load=("mean_daily", "sum"),
        mean_daily=("mean_daily", "mean"),
        sd_daily=("mean_daily", "std"),
        days=("date", "nunique"),
    )
    weekly["monotony"] = weekly["mean_daily"] / weekly["sd_daily"].replace(0, np.nan)
    # Cap monotony at 10 to prevent numerical blow-ups from near-zero SD
    # (e.g., weeks with only 1-2 sessions or identical loads produce SD ≈ 0)
    weekly["monotony"] = weekly["monotony"].clip(upper=10.0)
    weekly.loc[weekly["days"] < 3, "monotony"] = np.nan
    weekly["strain"] = weekly["week_load"] * weekly["monotony"]
    return weekly


# ----------------------------
# COACH ML: AVAILABILITY RISK
# ----------------------------
@st.cache_data(show_spinner=False)
def build_availability_risk_dataset(df: pd.DataFrame) -> Tuple[pd.DataFrame, List[str]]:
    """Build next-session availability labels and model features.

    Uses LAGGED injury history (1–3 sessions back) and rolling unavailability
    counts instead of current status to avoid data leakage, while still
    capturing the strongest predictor in sports medicine: previous injury.
    """
    if df.empty:
        return pd.DataFrame(), []

    work = df.copy().sort_values(["player_id", "date"])
    work["injury_status_actual"] = work.get("injury_status_actual", "Available").fillna("Available")
    work["next_status"] = work.groupby("player_id")["injury_status_actual"].shift(-1)
    work["next_date"] = work.groupby("player_id")["date"].shift(-1)
    work["days_to_next"] = (work["next_date"] - work["date"]).dt.days
    work = work[work["next_status"].notna()].copy()
    if work.empty:
        return pd.DataFrame(), []

    work = work[(work["days_to_next"].fillna(99) >= 0) & (work["days_to_next"].fillna(99) <= 14)].copy()
    if work.empty:
        return pd.DataFrame(), []

    work["target_not_available"] = (work["next_status"].astype(str).str.lower() != "available").astype(int)

    g = work.groupby("player_id", sort=False)

    # ── INJURY HISTORY features (lagged — no current-session leakage) ──
    unavail = (work["injury_status_actual"].str.lower() != "available").astype(int)
    work["unavail_lag1"] = g.apply(lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(1)).droplevel(0)
    work["unavail_lag2"] = g.apply(lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(2)).droplevel(0)
    work["unavail_lag3"] = g.apply(lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(3)).droplevel(0)
    work["unavail_roll7"] = g.apply(
        lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(1).rolling(7, min_periods=1).sum()
    ).droplevel(0)
    work["unavail_roll14"] = g.apply(
        lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(1).rolling(14, min_periods=1).sum()
    ).droplevel(0)
    work["unavail_pct28"] = g.apply(
        lambda x: x.assign(_u=unavail.loc[x.index])["_u"].shift(1).rolling(28, min_periods=3).mean()
    ).droplevel(0)

    # ── WORKLOAD trajectory features ──
    for raw_col in ["total_player_load", "total_distance", "total_acceleration_load"]:
        if raw_col not in work.columns:
            continue
        work[f"{raw_col}_lag1"] = g[raw_col].shift(1)
        work[f"{raw_col}_lag2"] = g[raw_col].shift(2)
        work[f"{raw_col}_roll3_mean"] = g[raw_col].transform(
            lambda x: x.shift(1).rolling(3, min_periods=1).mean()
        )
        work[f"{raw_col}_roll7_mean"] = g[raw_col].transform(
            lambda x: x.shift(1).rolling(7, min_periods=2).mean()
        )
        work[f"{raw_col}_roll3_std"] = g[raw_col].transform(
            lambda x: x.shift(1).rolling(3, min_periods=2).std()
        )
    if "maximum_velocity" in work.columns:
        work["max_vel_lag1"] = g["maximum_velocity"].shift(1)
        work["max_vel_roll3_mean"] = g["maximum_velocity"].transform(
            lambda x: x.shift(1).rolling(3, min_periods=1).mean()
        )
    if "acwr_ewma_7_28" in work.columns:
        work["acwr_lag1"] = g["acwr_ewma_7_28"].shift(1)
        work["acwr_delta"] = work["acwr_ewma_7_28"] - work["acwr_lag1"]
    if "fatigue_idx" in work.columns:
        work["fatigue_lag1"] = g["fatigue_idx"].shift(1)
    if "readiness" in work.columns:
        work["readiness_lag1"] = g["readiness"].shift(1)

    # ── Feature lists ──
    injury_history_features = [
        "unavail_lag1", "unavail_lag2", "unavail_lag3",
        "unavail_roll7", "unavail_roll14", "unavail_pct28",
    ]
    load_wellness_features = [
        "total_player_load", "player_load_per_min_est", "total_distance",
        "maximum_velocity", "total_acceleration_load", "acwr_ewma_7_28",
        "anomaly_score", "fatigue_idx", "readiness",
        "mental_score", "physical_score", "sleep_score",
        "soreness_score", "wellness_total",
    ]
    lag_features = [
        "total_player_load_lag1", "total_player_load_lag2",
        "total_player_load_roll3_mean", "total_player_load_roll7_mean",
        "total_player_load_roll3_std",
        "total_distance_lag1", "total_distance_lag2",
        "total_distance_roll3_mean", "total_distance_roll7_mean",
        "total_distance_roll3_std",
        "total_acceleration_load_lag1", "total_acceleration_load_lag2",
        "total_acceleration_load_roll3_mean", "total_acceleration_load_roll7_mean",
        "total_acceleration_load_roll3_std",
        "max_vel_lag1", "max_vel_roll3_mean",
        "acwr_lag1", "acwr_delta",
        "fatigue_lag1", "readiness_lag1",
    ]

    feature_cols = [c for c in (injury_history_features + load_wellness_features + lag_features) if c in work.columns]

    if len(feature_cols) < 5:
        return pd.DataFrame(), []

    keep_cols = ["player_id", "date", "injury_status_actual", "next_status", "target_not_available"] + feature_cols
    return work[keep_cols].copy(), feature_cols


# ----------------------------
# LOAD FORECASTING (regression)
# ----------------------------
@st.cache_data(show_spinner=False)
def build_forecast_dataset(df: pd.DataFrame, target_col: str, lags: int = 5) -> Tuple[pd.DataFrame, List[str]]:
    """Create lagged features for forecasting a numeric metric one session ahead."""
    if df.empty or target_col not in df.columns:
        return pd.DataFrame(), []
    work = df[["player_id", "date", target_col]].dropna().copy()
    work = work.sort_values(["player_id", "date"])
    for i in range(1, lags + 1):
        work[f"lag_{i}"] = work.groupby("player_id")[target_col].shift(i)
    work["roll_mean_3"] = work.groupby("player_id")[target_col].transform(lambda x: x.shift(1).rolling(3, min_periods=1).mean())
    work["roll_std_3"] = work.groupby("player_id")[target_col].transform(lambda x: x.shift(1).rolling(3, min_periods=1).std())
    work["roll_mean_7"] = work.groupby("player_id")[target_col].transform(lambda x: x.shift(1).rolling(7, min_periods=1).mean())
    work["pct_change_1"] = work.groupby("player_id")[target_col].pct_change().shift(1)
    work["pct_change_1"] = work["pct_change_1"].replace([np.inf, -np.inf], np.nan)
    work = work.dropna().copy()
    feat_cols = [c for c in work.columns if c not in ["player_id", "date", target_col]]
    return work, feat_cols


def run_forecast_benchmark(X_train, y_train, X_test, y_test):
    """Train multiple regressors and return leaderboard + best model."""
    models = {
        "Ridge": Pipeline([("scaler", StandardScaler()), ("model", Ridge(alpha=1.0))]),
        "Random Forest": RandomForestRegressor(n_estimators=300, max_depth=8, random_state=42, n_jobs=-1),
        "Gradient Boosting": GradientBoostingRegressor(n_estimators=200, max_depth=5, random_state=42),
        "MLP Regressor": Pipeline([("scaler", StandardScaler()), ("model", MLPRegressor(hidden_layer_sizes=(64, 32), max_iter=500, random_state=42))]),
    }
    rows = []
    best_name, best_score, best_mdl = None, np.inf, None
    for name, mdl in models.items():
        try:
            mdl.fit(X_train, y_train)
            preds = mdl.predict(X_test)
            mae = mean_absolute_error(y_test, preds)
            r2 = r2_score(y_test, preds)
            rows.append({"model": name, "MAE": round(mae, 3), "R2": round(r2, 4)})
            if mae < best_score:
                best_score = mae
                best_name = name
                best_mdl = mdl
        except Exception:
            rows.append({"model": name, "MAE": np.nan, "R2": np.nan})
    return pd.DataFrame(rows).sort_values("MAE"), best_name, best_mdl


def rolling_backtest_availability(
    ml_df: pd.DataFrame,
    feat_cols: List[str],
    model_factory,
    train_days: int = 60,
    test_days: int = 7,
    step_days: int = 7,
) -> pd.DataFrame:
    """Rolling-window backtest: train on last N days, predict next M days, slide forward."""
    if ml_df.empty or "date" not in ml_df.columns:
        return pd.DataFrame()
    data = ml_df.sort_values("date").copy()
    data["date"] = pd.to_datetime(data["date"], errors="coerce")
    data = data.dropna(subset=["date"])
    if data.empty:
        return pd.DataFrame()

    X_all = data[feat_cols].copy()
    med = X_all.median(numeric_only=True)
    X_all = X_all.fillna(med)
    y_all = data["target_not_available"].astype(int).copy()

    start = data["date"].min() + pd.Timedelta(days=train_days)
    end = data["date"].max() - pd.Timedelta(days=test_days)
    if start >= end:
        return pd.DataFrame()

    rows = []
    t = start
    while t <= end:
        train_start = t - pd.Timedelta(days=train_days)
        train_end = t
        test_end = t + pd.Timedelta(days=test_days)

        train_mask = (data["date"] >= train_start) & (data["date"] < train_end)
        test_mask = (data["date"] >= train_end) & (data["date"] < test_end)

        if train_mask.sum() < 50 or test_mask.sum() < 10:
            t = t + pd.Timedelta(days=step_days)
            continue

        X_tr, y_tr = X_all.loc[train_mask], y_all.loc[train_mask]
        X_te, y_te = X_all.loc[test_mask], y_all.loc[test_mask]

        if y_tr.nunique() < 2 or y_te.nunique() < 2:
            t = t + pd.Timedelta(days=step_days)
            continue

        mdl = model_factory()
        try:
            mdl.fit(X_tr, y_tr)
            p = mdl.predict_proba(X_te)[:, 1]
            auc = float(roc_auc_score(y_te, p))
            acc = float(accuracy_score(y_te, (p >= 0.5).astype(int)))
            rows.append({
                "window_end": train_end,
                "train_n": int(len(y_tr)),
                "test_n": int(len(y_te)),
                "test_pos": int(y_te.sum()),
                "roc_auc": auc,
                "accuracy": acc,
            })
        except Exception:
            pass

        t = t + pd.Timedelta(days=step_days)

    return pd.DataFrame(rows)


@st.cache_resource(show_spinner="Training ML models…")
def _train_availability_models(_data_hash: str, X_train, y_train, X_test, y_test, feat_cols):
    """Cached model training — only re-runs when data changes."""
    from sklearn.base import clone as _clone
    n_neg = int((y_train == 0).sum())
    n_pos = max(1, int((y_train == 1).sum()))
    spw = n_neg / n_pos

    model_defs = {
        "Logistic Regression": Pipeline([
            ("scaler", StandardScaler()),
            ("model", LogisticRegression(max_iter=2000, class_weight="balanced", random_state=42))
        ]),
        "Random Forest": RandomForestClassifier(
            n_estimators=500, max_depth=8, random_state=42,
            class_weight="balanced_subsample", min_samples_leaf=5
        ),
        "Extra Trees": ExtraTreesClassifier(
            n_estimators=500, max_depth=8, random_state=42,
            class_weight="balanced_subsample", min_samples_leaf=5
        ),
        "Gradient Boosting": GradientBoostingClassifier(
            random_state=42, n_estimators=400, max_depth=4,
            learning_rate=0.05, subsample=0.8, min_samples_leaf=10,
        ),
        **({"XGBoost": _xgb_cls(
            eval_metric="logloss", random_state=42, n_estimators=400,
            max_depth=4, learning_rate=0.05, scale_pos_weight=spw,
            subsample=0.8, colsample_bytree=0.8,
        )} if _xgb_cls is not None else {}),
        "AdaBoost": AdaBoostClassifier(n_estimators=200, random_state=42),
        "SVM (RBF)": Pipeline([
            ("scaler", StandardScaler()),
            ("model", SVC(probability=True, class_weight="balanced", random_state=42))
        ]),
        "KNN": Pipeline([
            ("scaler", StandardScaler()),
            ("model", KNeighborsClassifier(n_neighbors=7, weights="distance"))
        ]),
        "Naive Bayes": GaussianNB(),
        "Neural Net (MLP)": Pipeline([
            ("scaler", StandardScaler()),
            ("model", MLPClassifier(
                hidden_layer_sizes=(64, 32), alpha=1e-3,
                learning_rate_init=1e-3, max_iter=800, random_state=42
            ))
        ]),
    }
    min_class_train = int(y_train.value_counts().min()) if len(y_train) > 0 else 1
    n_cv = min(5, len(y_train), min_class_train)
    cv = StratifiedKFold(n_splits=n_cv, shuffle=True, random_state=42) if n_cv >= 2 else None

    leaderboard = []
    best_name = None
    best_cv = -np.inf
    errors = []
    trained_models = {}

    from sklearn.utils.class_weight import compute_sample_weight
    sw_train = compute_sample_weight("balanced", y_train)
    sw_models = {"Gradient Boosting", "AdaBoost"}

    for name, mdl in model_defs.items():
        cv_auc = np.nan
        hold_auc = np.nan
        hold_acc = np.nan
        try:
            if cv is not None:
                scores = cross_val_score(mdl, X_train, y_train, cv=cv, scoring="roc_auc")
                cv_auc = float(np.nanmean(scores))
            if name in sw_models:
                mdl.fit(X_train, y_train, sample_weight=sw_train)
            else:
                mdl.fit(X_train, y_train)
            p_test = mdl.predict_proba(X_test)[:, 1]
            y_hat = (p_test >= 0.50).astype(int)
            hold_acc = float(accuracy_score(y_test, y_hat))
            if y_test.nunique() > 1:
                hold_auc = float(roc_auc_score(y_test, p_test))
            trained_models[name] = mdl
        except Exception as e:
            errors.append(f"{name}: {e}")

        leaderboard.append({
            "model": name, "cv_roc_auc": cv_auc,
            "holdout_roc_auc": hold_auc, "holdout_accuracy": hold_acc,
        })
        score_for_pick = cv_auc if np.isfinite(cv_auc) else hold_auc
        if np.isfinite(score_for_pick) and score_for_pick > best_cv:
            best_cv = score_for_pick
            best_name = name

    return leaderboard, best_name, trained_models, errors


# ----------------------------
# MATCHDAY / MD-1 LOGIC (unique)
# ----------------------------
def label_matchday_windows(training_df: pd.DataFrame, schedule_df: pd.DataFrame) -> pd.DataFrame:
    df = training_df.copy()
    df["date"] = pd.to_datetime(df["date"], errors="coerce")

    sched = schedule_df.copy()
    if "match_date" not in sched.columns:
        return df.assign(md_label="UNK")

    match_dates = sched["match_date"].dropna().sort_values().unique()
    if len(match_dates) == 0:
        return df.assign(md_label="UNK")

    td = df["date"].dropna().unique()
    md_map = {}
    for d in td:
        diffs = np.array([(pd.Timestamp(m) - pd.Timestamp(d)).days for m in match_dates], dtype=float)
        idx = np.argmin(np.abs(diffs))
        delta = int(diffs[idx])

        if abs(delta) <= 7:
            if delta == 0:
                lab = "MD"
            elif delta < 0:
                lab = f"MD+{abs(delta)}"
            else:
                lab = f"MD-{delta}"
        else:
            lab = "Non-MD Window"

        md_map[pd.Timestamp(d)] = lab

    df["md_label"] = df["date"].map(lambda x: md_map.get(pd.Timestamp(x), "UNK"))
    return df


# ==============================================================================
# RAG SYSTEM FOR OFFLINE AI INSIGHTS
# ==============================================================================
class OfflineRAGSystem:
    """Lightweight offline Q&A system for sports performance analytics.
    Pre-computes TF-IDF once and caches everything to disk.
    Subsequent loads take <0.3 sec — no fitting, no model downloads."""

    _CACHE_DIR  = Path(__file__).resolve().parent / ".rag_cache"
    _KB_FILE    = _CACHE_DIR / "knowledge_base.pkl"
    _VEC_FILE   = _CACHE_DIR / "tfidf_vectorizer.pkl"
    _MAT_FILE   = _CACHE_DIR / "tfidf_matrix.npz"
    _HASH_FILE  = _CACHE_DIR / "data_hash.txt"

    def __init__(self):
        self.knowledge_base: List[Dict] = []
        self.vectorizer = TfidfVectorizer(max_features=2000, stop_words='english', ngram_range=(1, 2))
        self.tfidf_matrix = None
        self._last_df = pd.DataFrame()
        self._last_game_all = pd.DataFrame()
        self._last_team_match = pd.DataFrame()
        self._CACHE_DIR.mkdir(exist_ok=True)

    # ── helpers ──
    @staticmethod
    def _doc(text: str, dtype: str, **meta) -> Dict:
        return {"text": text, "type": dtype, "metadata": meta}

    @staticmethod
    def _data_hash(df_train: pd.DataFrame) -> str:
        """Stable fingerprint: row count + date range + player count.
        Does NOT depend on filter-specific aggregates so sidebar
        changes won't trigger an expensive rebuild."""
        import hashlib
        sig = (f"{len(df_train)}"
               f"_{df_train['date'].min()}_{df_train['date'].max()}"
               f"_{df_train['player_id'].nunique()}"
               f"_{sorted(df_train.columns.tolist())}")
        return hashlib.md5(sig.encode()).hexdigest()

    # ── Disk persistence (KB + fitted vectorizer + sparse matrix) ──
    def _save_cache(self, data_hash: str):
        import pickle
        from scipy import sparse as sp
        try:
            with open(self._KB_FILE, "wb") as f:
                pickle.dump(self.knowledge_base, f, protocol=pickle.HIGHEST_PROTOCOL)
            with open(self._VEC_FILE, "wb") as f:
                pickle.dump(self.vectorizer, f, protocol=pickle.HIGHEST_PROTOCOL)
            if self.tfidf_matrix is not None:
                sp.save_npz(str(self._MAT_FILE), self.tfidf_matrix)
            self._HASH_FILE.write_text(data_hash, encoding="utf-8")
        except Exception:
            pass

    def _load_cache(self, data_hash: str) -> bool:
        """Load KB + pre-fitted vectorizer + sparse matrix from disk.
        Skips fit_transform entirely — instant startup."""
        import pickle
        from scipy import sparse as sp
        try:
            if not self._KB_FILE.exists() or not self._HASH_FILE.exists():
                return False
            if self._HASH_FILE.read_text(encoding="utf-8").strip() != data_hash:
                return False
            with open(self._KB_FILE, "rb") as f:
                self.knowledge_base = pickle.load(f)
            if self._VEC_FILE.exists() and self._MAT_FILE.exists():
                with open(self._VEC_FILE, "rb") as f:
                    self.vectorizer = pickle.load(f)
                self.tfidf_matrix = sp.load_npz(str(self._MAT_FILE))
            else:
                self._fit_tfidf()
            return len(self.knowledge_base) > 0
        except Exception:
            return False

    def _fit_tfidf(self):
        """Fit TF-IDF from scratch (only on first build or cache miss)."""
        texts = [d["text"] for d in self.knowledge_base]
        if texts:
            try:
                self.tfidf_matrix = self.vectorizer.fit_transform(texts)
            except Exception:
                pass
    
    def create_knowledge_base(self, df_train: pd.DataFrame, game_all: pd.DataFrame,
                              team_match: pd.DataFrame, player_match: pd.DataFrame,
                              force_rebuild: bool = False):
        self._last_df = df_train
        self._last_game_all = game_all
        self._last_team_match = team_match
        h = self._data_hash(df_train) if not df_train.empty else ""
        if not force_rebuild and h and self._load_cache(h):
            return
        docs: List[Dict] = []

        # ── 1. Player-level summaries (ALL players) ──
        if not df_train.empty:
            latest = df_train.sort_values("date").groupby("player_id").tail(1)
            for _, r in latest.iterrows():
                pid = r["player_id"]
                p = df_train[df_train["player_id"] == pid]
                parts = [f"Player {pid}: {len(p)} sessions"]
                meta: Dict = {"player_id": pid, "sessions": len(p)}
                if "total_player_load" in p.columns:
                    v = float(p["total_player_load"].mean())
                    parts.append(f"avg load {v:.0f}")
                    meta["avg_load"] = round(v, 1)
                if "maximum_velocity" in p.columns:
                    v = float(p["maximum_velocity"].max())
                    parts.append(f"top speed {v:.1f} mph")
                    meta["max_velocity"] = round(v, 2)
                if "acwr_ewma_7_28" in p.columns:
                    v = float(p["acwr_ewma_7_28"].dropna().iloc[-1]) if p["acwr_ewma_7_28"].notna().any() else np.nan
                    if np.isfinite(v):
                        parts.append(f"latest ACWR {v:.2f}")
                        meta["acwr"] = round(v, 2)
                if "readiness" in p.columns and p["readiness"].notna().any():
                    v = float(p["readiness"].dropna().iloc[-1])
                    parts.append(f"readiness {v:.0f}/100")
                    meta["readiness"] = round(v, 1)
                if "fatigue_idx" in p.columns and p["fatigue_idx"].notna().any():
                    v = float(p["fatigue_idx"].dropna().iloc[-1])
                    parts.append(f"fatigue idx {v:.2f}")
                    meta["fatigue_idx"] = round(v, 2)
                if "injury_status_actual" in r and pd.notna(r.get("injury_status_actual")):
                    parts.append(f"status {r['injury_status_actual']}")
                    meta["injury_status"] = str(r["injury_status_actual"])
                if "wellness_total" in p.columns and p["wellness_total"].notna().any():
                    v = float(p["wellness_total"].dropna().iloc[-1])
                    parts.append(f"wellness {v:.1f}/20")
                    meta["wellness"] = round(v, 1)
                if "anomaly_flag" in p.columns:
                    n = int(p["anomaly_flag"].sum())
                    parts.append(f"{n} anomaly sessions")
                    meta["anomaly_count"] = n
                if "position_name" in p.columns:
                    meta["position"] = str(p["position_name"].iloc[-1])
                docs.append(self._doc(", ".join(parts) + ".", "player_summary", **meta))

        # ── 2. Team daily summaries (top-20% load days + bottom-20%) ──
        if not df_train.empty and "total_player_load" in df_train.columns:
            td = df_train.groupby("date").agg(
                avg_load=("total_player_load", "mean"),
                max_vel=("maximum_velocity", "max") if "maximum_velocity" in df_train.columns else ("total_player_load", "count"),
                n_anom=("anomaly_flag", "sum") if "anomaly_flag" in df_train.columns else ("total_player_load", "count"),
                n_players=("player_id", "nunique"),
            ).reset_index()
            q80 = td["avg_load"].quantile(0.80)
            q20 = td["avg_load"].quantile(0.20)
            for _, row in td[td["avg_load"] >= q80].iterrows():
                docs.append(self._doc(
                    f"High-load day {row['date'].strftime('%m/%d/%Y')}: team avg load {row['avg_load']:.0f}, {row['n_players']} players, {row['n_anom']} anomalies.",
                    "team_high_load", date=str(row["date"]), avg_load=round(float(row["avg_load"]), 1)))
            for _, row in td[td["avg_load"] <= q20].iterrows():
                docs.append(self._doc(
                    f"Low-load / recovery day {row['date'].strftime('%m/%d/%Y')}: team avg load {row['avg_load']:.0f}.",
                    "team_low_load", date=str(row["date"]), avg_load=round(float(row["avg_load"]), 1)))

        # ── 3. Injury progression documents ──
        if not df_train.empty and "injury_status_actual" in df_train.columns:
            inj = df_train[~df_train["injury_status_actual"].astype(str).str.lower().isin(["available", "nan", ""])].copy()
            if not inj.empty:
                for pid, grp in inj.groupby("player_id"):
                    statuses = grp.sort_values("date")["injury_status_actual"].tolist()
                    first_d = grp["date"].min().strftime("%m/%d/%Y")
                    last_d = grp["date"].max().strftime("%m/%d/%Y")
                    docs.append(self._doc(
                        f"Player {pid} was non-available from {first_d} to {last_d}. Status history: {' → '.join(statuses[:10])}.",
                        "injury_timeline", player_id=pid, first_date=first_d, last_date=last_d))

        # ── 4. Microcycle / day-of-week norms ──
        if not df_train.empty and "total_player_load" in df_train.columns:
            _d = df_train.copy()
            _d["day_name"] = pd.to_datetime(_d["date"]).dt.day_name()
            dow = _d.groupby("day_name")["total_player_load"].agg(["mean", "std"]).reset_index()
            for _, r in dow.iterrows():
                docs.append(self._doc(
                    f"Typical {r['day_name']} training load: mean {r['mean']:.0f} ± {r['std']:.0f}.",
                    "microcycle_norm", day=r["day_name"], mean=round(float(r["mean"]), 1), std=round(float(r["std"]), 1)))

        # ── 5. Velocity zone summaries ──
        if not df_train.empty and "vmax_zone" in df_train.columns:
            vz = df_train["vmax_zone"].value_counts()
            for zone, cnt in vz.items():
                pct = 100 * cnt / len(df_train)
                docs.append(self._doc(
                    f"Velocity zone {zone} accounts for {pct:.1f}% of all sessions ({cnt} sessions).",
                    "velocity_zone", zone=zone, pct=round(pct, 1), count=int(cnt)))

        # ── 6. Game stats ──
        _stat_names = {"G": "goals", "A": "assists", "Sh": "shots taken", "SOG": "shots on goal"}
        if game_all is not None and not game_all.empty:
            n_matches = game_all["match_date"].nunique() if "match_date" in game_all.columns else game_all.groupby(["player_id"]).ngroups
            for col in ["G", "A", "Sh", "SOG"]:
                if col in game_all.columns:
                    top = game_all.groupby("player_id")[col].sum().sort_values(ascending=False).head(5)
                    for pid, val in top.items():
                        if val > 0:
                            label = _stat_names.get(col, col)
                            docs.append(self._doc(
                                f"Player {pid} recorded {int(val)} {label} in total across the full season (not a single game).",
                                "game_top_performer", player_id=pid, stat=label, value=int(val)))

        # ── 7. Opponent analysis ──
        if not team_match.empty and "Opponent_clean" in team_match.columns:
            for _, r in team_match.groupby("Opponent_clean").agg(
                team_load=("team_load", "mean"), mean_acwr=("mean_acwr", "mean")
            ).reset_index().iterrows():
                docs.append(self._doc(
                    f"Against {r['Opponent_clean']}: avg team load {r['team_load']:.0f}, avg ACWR {r['mean_acwr']:.2f}.",
                    "opponent_analysis", opponent=r["Opponent_clean"],
                    avg_load=round(float(r["team_load"]), 1), avg_acwr=round(float(r["mean_acwr"]), 2)))

        # ── 8. Expert rules (comprehensive sports science knowledge base) ──
        _rules = [
            ("ACWR above 1.3 indicates elevated injury risk. Above 1.5 is a danger zone requiring immediate load reduction.", "acwr_injury_risk"),
            ("ACWR below 0.8 indicates chronic under-loading. The athlete may lose fitness and be under-prepared for competition.", "acwr_underload"),
            ("The ACWR sweet spot is 0.8–1.3. Keeping athletes in this range minimizes injury risk while maintaining fitness.", "acwr_sweet_spot"),
            ("Load should increase by no more than 10% per week (the 10% rule) to avoid overuse injuries.", "load_progression"),
            ("Weekly monotony above 2.0 indicates overly uniform training. Vary session intensities to reduce illness risk.", "monotony_high"),
            ("Weekly strain (load × monotony) spikes are associated with illness and soft-tissue injury in the following week.", "strain_spike"),
            ("Maximum velocity is a key indicator of neuromuscular readiness. A drop >10% from personal best warrants investigation.", "velocity_readiness"),
            ("High-speed running (≥70% vmax) and sprinting (≥90% vmax) exposure in training should approach game demands.", "velocity_zones"),
            ("If training never reaches game-day intensity, athletes are under-prepared. If it always exceeds it, they are over-exposed.", "match_readiness"),
            ("MD-1 (day before match) should be low volume, moderate intensity — activation, not fatigue.", "md1_recovery"),
            ("MD-2 is typically the highest-intensity training day in the microcycle.", "md2_intensity"),
            ("MD-3 and MD-4 are the primary loading days. Volume should be highest here.", "md34_loading"),
            ("MD+1 (day after match) should focus on active recovery: low load, light movement.", "md_plus1"),
            ("A readiness score below 60/100 should trigger a coach–sports scientist conversation about load management.", "readiness_low"),
            ("Fatigue index (7-day load / 28-day avg) above 1.5 indicates acute fatigue accumulation.", "fatigue_high"),
            ("Soreness score trending down over multiple days combined with rising load is a warning sign.", "soreness_trend"),
            ("Sleep quality below 3/5 for multiple consecutive days impairs recovery and increases injury risk.", "sleep_quality"),
            ("Wellness scores should be monitored Monday morning (post-weekend) and Friday (pre-match) as minimum.", "wellness_timing"),
            ("Players returning from injury (Out → Limited → As Tolerated) should follow a graded return-to-play protocol.", "return_to_play"),
            ("Deceleration load is often overlooked but is strongly associated with hamstring strain risk.", "decel_hamstring"),
            ("Explosive efforts are a key metric for soccer: they correlate with match actions like tackles, sprints, and jumps.", "explosive_efforts"),
            ("Position-specific loading matters: forwards typically have higher sprint exposure, defenders higher deceleration load.", "position_loading"),
            ("Goalkeeper training load profiles differ substantially from outfield players. Monitor them separately.", "gk_specific"),
            ("Comparing training P90 to match-day P90 reveals whether athletes are physically prepared for worst-case game scenarios.", "p90_comparison"),
        ]
        for text, rule in _rules:
            docs.append(self._doc(text, "expert_rule", rule=rule))

        self.knowledge_base = docs
        self._fit_tfidf()
        if h:
            self._save_cache(h)
    
    def query(self, question: str, top_k: int = 5) -> List[Dict]:
        """Fast TF-IDF keyword search over the knowledge base."""
        if not self.knowledge_base:
            return []
        similarities = self._tfidf_query(question)
        top_indices = np.argsort(similarities)[::-1][:top_k]
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.05:
                results.append({
                    'text': self.knowledge_base[idx]['text'],
                    'type': self.knowledge_base[idx]['type'],
                    'metadata': self.knowledge_base[idx].get('metadata', {}),
                    'similarity': float(similarities[idx]),
                })
        return results
    
    def _tfidf_query(self, question: str) -> np.ndarray:
        """Fallback TF-IDF query"""
        if self.tfidf_matrix is None:
            return np.zeros(len(self.knowledge_base))
        
        try:
            query_vec = self.vectorizer.transform([question])
            similarities = cosine_similarity(query_vec, self.tfidf_matrix)[0]
            return similarities
        except Exception:
            return np.zeros(len(self.knowledge_base))
    
    # ── Smart answer engine (template-based NLG) ──
    @staticmethod
    def _classify_question(q: str) -> str:
        q = q.lower()
        kw = {
            "match_performance": [
                "match", "game", "perform", "best game", "worst game", "won", "lost",
                "win", "loss", "opponent", "result", "score", "goal", "assist",
                "shot", "played", "season record", "fixture", "compete"],
            "injury": [
                "injur", "out", "limited", "return", "hurt", "status", "availability",
                "unavailable", "tolerated", "sideline", "rehab", "medical", "cleared",
                "fit", "unfit", "broken", "strain", "sprain", "concuss", "muscle"],
            "load": [
                "load", "overload", "training load", "volume", "intensity", "periodiz",
                "overtraining", "underload", "heavy", "manage load", "reduce load",
                "increase load", "taper", "ramp", "how much", "distance", "meters",
                "total distance"],
            "fatigue": [
                "fatigue", "tired", "rest", "recovery", "readiness", "fresh", "exhausted",
                "worn out", "need rest", "overworked", "burnout", "recup", "fatigued",
                "ready to play", "ready to train"],
            "speed": [
                "speed", "velocity", "sprint", "fast", "vmax", "zone", "high-speed",
                "fastest", "quick", "acceleration", "deceleration", "top speed",
                "how fast", "mph", "running"],
            "wellness": [
                "wellness", "soreness", "sleep", "mental", "physical", "mood",
                "feel", "feeling", "sore", "pain", "energy", "stress", "wellbeing",
                "how are", "how is"],
            "acwr": [
                "acwr", "acute", "chronic", "ratio", "workload ratio", "a:c",
                "sweet spot", "danger zone", "injury risk", "risk zone"],
            "player": [
                "player", "who", "best player", "top player", "worst", "compare",
                "which player", "strongest", "weakest", "standout", "overview",
                "squad", "roster", "team summary"],
            "week": [
                "week", "this week", "next week", "monday", "tuesday", "wednesday",
                "thursday", "friday", "saturday", "sunday", "day", "schedule",
                "plan", "today", "tomorrow", "microcycle", "session plan",
                "what should", "recommend"],
            "position": [
                "forward", "midfielder", "defender", "goalkeeper", "position", "gk",
                "def", "mid", "fwd", "striker", "winger", "fullback", "center back",
                "attacking", "defensive", "holding"],
        }
        scores = {}
        for cat, words in kw.items():
            scores[cat] = sum(1 for w in words if w in q)
        best = max(scores, key=scores.get)
        return best if scores[best] > 0 else "general"

    def smart_answer(self, question: str, df: pd.DataFrame,
                     game_all: pd.DataFrame, team_match: pd.DataFrame) -> str:
        """Generate an intelligent, conversational answer by analyzing the actual data."""
        cat = self._classify_question(question)
        lines: List[str] = []

        # ── MATCH PERFORMANCE ──
        if cat == "match_performance":
            q_lower = question.lower()
            asking_wins = any(w in q_lower for w in ["won", "win", "wins", "victories", "beat"])
            asking_losses = any(w in q_lower for w in ["lost", "lose", "loss", "losses", "defeat"])
            asking_best = any(w in q_lower for w in ["best", "well", "great", "top"])
            asking_worst = any(w in q_lower for w in ["worst", "bad", "poor", "struggle"])

            if team_match is not None and not team_match.empty and "Opponent_clean" in team_match.columns:
                tm = team_match.copy()
                has_result = "result" in tm.columns and tm["result"].notna().any() and (tm["result"].str.strip() != "").any()

                if has_result:
                    tm["_res"] = tm["result"].fillna("").str.strip().str.lower()
                    wins = tm[tm["_res"].str.startswith("w")]
                    losses = tm[tm["_res"].str.startswith("l")]
                    draws = tm[tm["_res"].str.match(r"^[dt]")]
                    total = len(wins) + len(losses) + len(draws)
                    lines.append(f"Overall record: **{len(wins)}W – {len(losses)}L – {len(draws)}D** across {total} matches.")
                    lines.append("")

                    def _fmt_match(row):
                        opp = row.get("Opponent_clean", "?")
                        res = str(row.get("result", "")).strip()
                        dt = row.get("date", "")
                        dt_s = dt.strftime("%m/%d/%Y") if hasattr(dt, "strftime") else str(dt)[:10]
                        load = row.get("team_load", 0)
                        acwr = row.get("mean_acwr", 0)
                        n = int(row.get("players", 0))
                        parts = [f"**vs {opp}** ({dt_s}) — {res}"]
                        if load: parts.append(f"team load {load:,.0f}")
                        if acwr: parts.append(f"ACWR {acwr:.2f}")
                        if n: parts.append(f"{n} players")
                        return "- " + ", ".join(parts) + "."

                    if asking_wins or (asking_best and not asking_worst):
                        if len(wins) == 0:
                            lines.append("No wins recorded in the current data window.")
                        else:
                            lines.append(f"**All {len(wins)} wins:**")
                            for _, r in wins.sort_values("date").iterrows():
                                lines.append(_fmt_match(r))
                            lines.append("")
                            best_win = wins.sort_values("team_load", ascending=False).head(1).iloc[0]
                            lines.append(f"The **biggest performance win** was vs **{best_win['Opponent_clean']}** "
                                         f"({str(best_win.get('result','')).strip()}) with a team load of {best_win['team_load']:,.0f}. "
                                         f"High team load in a win means the squad was fully engaged and competitive.")
                    elif asking_losses or asking_worst:
                        if len(losses) == 0:
                            lines.append("No losses recorded — impressive!")
                        else:
                            lines.append(f"**All {len(losses)} losses:**")
                            for _, r in losses.sort_values("date").iterrows():
                                lines.append(_fmt_match(r))
                    else:
                        lines.append("**Recent matches:**")
                        for _, r in tm.sort_values("date", ascending=False).head(10).iterrows():
                            lines.append(_fmt_match(r))
                else:
                    lines.append("Match results (W/L/D) are not available in the schedule data. "
                                 "I can still show team load per opponent:")
                    lines.append("")
                    for _, r in tm.sort_values("team_load", ascending=False).head(8).iterrows():
                        opp = r.get("Opponent_clean", "?")
                        load = r.get("team_load", 0)
                        lines.append(f"- **vs {opp}** — team load {load:,.0f}")
            else:
                lines.append("No match-level data available with current filters.")

            if game_all is not None and not game_all.empty:
                _stat_labels = {"G": "goals", "A": "assists", "Sh": "shots", "SOG": "shots on goal"}
                stats_lines = []
                for col, label in _stat_labels.items():
                    if col in game_all.columns:
                        top = game_all.groupby("player_id")[col].sum().sort_values(ascending=False).head(3)
                        top = top[top > 0]
                        if len(top):
                            stats_lines.append(f"- **{label.title()}**: " + ", ".join(
                                f"**{pid}** ({int(v)} total across all matches)" for pid, v in top.items()))
                if stats_lines:
                    lines.append("")
                    lines.append("**Season leaders:**")
                    lines.extend(stats_lines)

        # ── INJURY ──
        elif cat == "injury":
            if "injury_status_actual" in df.columns:
                latest = df.sort_values("date").groupby("player_id")["injury_status_actual"].last()
                non_avail = latest[~latest.astype(str).str.lower().isin(["available", "nan", ""])]
                avail_count = len(latest) - len(non_avail)
                lines.append(f"Currently **{avail_count}** players are fully available and **{len(non_avail)}** are not.")
                if len(non_avail):
                    lines.append("")
                    for status in ["Out", "Limited", "As Tolerated"]:
                        players = non_avail[non_avail.astype(str).str.lower() == status.lower()]
                        if len(players):
                            names = ", ".join(f"**{p}**" for p in players.index[:8])
                            lines.append(f"- **{status}** ({len(players)}): {names}")
                    lines.append("")
                    lines.append("Players who are 'Limited' or 'As Tolerated' should follow a graded return-to-play "
                                 "protocol. Monitor their load carefully and check wellness scores before clearing them for full training.")
                else:
                    lines.append("Great news — the entire squad is currently available!")

        # ── LOAD ──
        elif cat == "load":
            if "total_player_load" in df.columns:
                avg_l = df["total_player_load"].mean()
                last7 = df[df["date"] >= df["date"].max() - pd.Timedelta(days=7)]
                avg_l7 = last7["total_player_load"].mean() if len(last7) else avg_l
                change = 100 * (avg_l7 - avg_l) / avg_l if avg_l > 0 else 0
                lines.append(f"The squad's overall average training load is **{avg_l:.0f}**.")
                if abs(change) > 5:
                    direction = "higher" if change > 0 else "lower"
                    lines.append(f"Over the **last 7 days**, average load was **{avg_l7:.0f}** — "
                                 f"that's **{abs(change):.0f}% {direction}** than the overall average.")
                else:
                    lines.append(f"The last 7 days ({avg_l7:.0f}) are in line with the season average — load is stable.")
                top_loaded = last7.groupby("player_id")["total_player_load"].mean().sort_values(ascending=False).head(5)
                if len(top_loaded):
                    lines.append("")
                    lines.append("**Highest loaded players this week:**")
                    for pid, val in top_loaded.items():
                        ratio = val / avg_l if avg_l > 0 else 1
                        flag = " ⚠️" if ratio > 1.3 else ""
                        lines.append(f"- **{pid}** — avg load {val:.0f} ({ratio:.1f}× squad avg){flag}")
                lines.append("")
                lines.append("Remember: load should increase by no more than 10% per week. "
                             "Players above 1.3× the squad average may need a lighter session or recovery day.")

        # ── FATIGUE / READINESS ──
        elif cat == "fatigue":
            parts = []
            if "readiness" in df.columns and df["readiness"].notna().any():
                last_r = df.sort_values("date").groupby("player_id")["readiness"].last()
                low_r = last_r[last_r < 60].sort_values()
                high_r = last_r[last_r >= 80].sort_values(ascending=False)
                parts.append(f"**Readiness scores** range from {last_r.min():.0f} to {last_r.max():.0f} (squad avg {last_r.mean():.0f}/100).")
                if len(low_r):
                    names = ", ".join(f"**{p}** ({v:.0f})" for p, v in low_r.head(5).items())
                    parts.append(f"")
                    parts.append(f"🟠 **Need attention** (readiness below 60): {names}.")
                    parts.append("These players should be monitored closely. Check their sleep and soreness scores "
                                 "and consider a lighter session or active recovery day.")
                if len(high_r):
                    names = ", ".join(f"**{p}** ({v:.0f})" for p, v in high_r.head(5).items())
                    parts.append(f"")
                    parts.append(f"🟢 **Ready to go** (readiness 80+): {names}.")
                    parts.append("These players are fresh and can handle higher-intensity work.")
                lines.extend(parts)
            if "fatigue_idx" in df.columns and df["fatigue_idx"].notna().any():
                last_f = df.sort_values("date").groupby("player_id")["fatigue_idx"].last()
                high_f = last_f[last_f > 1.5].sort_values(ascending=False)
                if len(high_f):
                    lines.append("")
                    names = ", ".join(f"**{p}** ({v:.2f})" for p, v in high_f.head(5).items())
                    lines.append(f"🔥 **High fatigue index** (above 1.5): {names}.")
                    lines.append("A fatigue index above 1.5 means the 7-day load is significantly higher than the 28-day "
                                 "trend — these players have been accumulating load and need recovery time.")

        # ── SPEED / VELOCITY ──
        elif cat == "speed":
            if "maximum_velocity" in df.columns:
                avg_v = df["maximum_velocity"].mean()
                top_v = df.groupby("player_id")["maximum_velocity"].max().sort_values(ascending=False).head(5)
                lines.append(f"Squad average top speed is **{avg_v:.1f} mph**.")
                lines.append("")
                lines.append("**Fastest players** (highest recorded top speed):")
                for pid, v in top_v.items():
                    lines.append(f"- **{pid}** — {v:.1f} mph")
            if "vmax_zone" in df.columns:
                vz = df["vmax_zone"].value_counts().sort_index()
                lines.append("")
                lines.append("**Speed zone distribution:**")
                for zone, cnt in vz.items():
                    pct = 100 * cnt / len(df)
                    lines.append(f"- {zone}: {pct:.0f}% of sessions")
                lines.append("")
                z4z5 = df[df["vmax_zone"].astype(str).str.startswith(("Z4", "Z5"))]
                pct_high = 100 * len(z4z5) / max(len(df), 1)
                lines.append(f"High-speed exposure (Z4+Z5) is **{pct_high:.0f}%** of sessions. "
                             f"{'This is good for match readiness.' if pct_high >= 15 else 'Consider adding more high-speed work to prepare for game demands.'}")

        # ── WELLNESS ──
        elif cat == "wellness":
            if "wellness_total" in df.columns and df["wellness_total"].notna().any():
                last_w = df.sort_values("date").groupby("player_id")["wellness_total"].last()
                avg_w = last_w.mean()
                low_w = last_w[last_w < 12].sort_values()
                lines.append(f"Squad average wellness is **{avg_w:.1f}/20**.")
                if len(low_w):
                    names = ", ".join(f"**{p}** ({v:.0f}/20)" for p, v in low_w.head(5).items())
                    lines.append(f"")
                    lines.append(f"😴 **Low wellness** (below 12/20): {names}.")
                for sub in ["sleep_score", "soreness_score", "mental_score"]:
                    if sub in df.columns and df[sub].notna().any():
                        last_sub = df.sort_values("date").groupby("player_id")[sub].last()
                        low_sub = last_sub[last_sub <= 2].sort_values()
                        label = sub.replace("_score", "").replace("_", " ").title()
                        if len(low_sub):
                            names = ", ".join(f"**{p}** ({v:.0f}/5)" for p, v in low_sub.head(5).items())
                            lines.append(f"- Low **{label}**: {names}")
                if len(low_w):
                    lines.append("")
                    lines.append("Players with low wellness should be checked in with individually. "
                                 "If sleep is the issue, consider adjusting training time. If soreness, adjust load.")
                else:
                    lines.append("The squad is reporting good wellness scores overall.")

        # ── ACWR ──
        elif cat == "acwr":
            if "acwr_ewma_7_28" in df.columns:
                last_a = df.sort_values("date").groupby("player_id")["acwr_ewma_7_28"].last().dropna()
                avg_a = last_a.mean()
                danger = last_a[last_a > 1.3].sort_values(ascending=False)
                under = last_a[last_a < 0.8].sort_values()
                sweet = last_a[(last_a >= 0.8) & (last_a <= 1.3)]
                lines.append(f"Squad average ACWR is **{avg_a:.2f}**. "
                             f"**{len(sweet)}** players are in the sweet spot (0.8–1.3), "
                             f"**{len(danger)}** are in the danger zone (>1.3), "
                             f"and **{len(under)}** are under-loaded (<0.8).")
                if len(danger):
                    lines.append("")
                    lines.append("🔴 **Danger zone** (ACWR > 1.3 — increased injury risk):")
                    for pid, v in danger.head(5).items():
                        lines.append(f"- **{pid}** — ACWR {v:.2f}")
                    lines.append("These players' recent load has spiked relative to their long-term trend. "
                                 "Reduce training volume 20–30% this week and monitor how they respond.")
                if len(under):
                    lines.append("")
                    lines.append("🟡 **Under-loaded** (ACWR < 0.8 — may be losing fitness):")
                    for pid, v in under.head(5).items():
                        lines.append(f"- **{pid}** — ACWR {v:.2f}")
                    lines.append("These players aren't doing enough relative to their baseline. "
                                 "Gradually increase their load to avoid being under-prepared for match demands.")

        # ── PLAYER-specific ──
        elif cat == "player":
            if "total_player_load" in df.columns:
                latest = df.sort_values("date").groupby("player_id").tail(1)
                cols_to_show = [c for c in ["player_id", "readiness", "fatigue_idx", "acwr_ewma_7_28",
                                            "total_player_load", "wellness_total", "injury_status_actual"] if c in latest.columns]
                top = latest.sort_values("total_player_load", ascending=False).head(5)
                lines.append("**Top 5 players by recent training load:**")
                for _, r in top.iterrows():
                    parts = [f"**{r['player_id']}**"]
                    if "total_player_load" in r: parts.append(f"load {r['total_player_load']:.0f}")
                    if "readiness" in r and pd.notna(r.get("readiness")): parts.append(f"readiness {r['readiness']:.0f}")
                    if "acwr_ewma_7_28" in r and pd.notna(r.get("acwr_ewma_7_28")): parts.append(f"ACWR {r['acwr_ewma_7_28']:.2f}")
                    lines.append(f"- {' · '.join(parts)}")

        # ── WEEKLY ──
        elif cat == "week":
            if "total_player_load" in df.columns:
                df_c = df.copy()
                df_c["day_name"] = pd.to_datetime(df_c["date"]).dt.day_name()
                dow = df_c.groupby("day_name")["total_player_load"].agg(["mean", "std"]).reindex(
                    ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]).dropna()
                lines.append("**Typical weekly training pattern:**")
                for day, row in dow.iterrows():
                    bar = "█" * int(row["mean"] / max(dow["mean"].max(), 1) * 15)
                    lines.append(f"- **{day}**: avg load {row['mean']:.0f} ± {row['std']:.0f} {bar}")
                heaviest = dow["mean"].idxmax()
                lightest = dow["mean"].idxmin()
                lines.append(f"")
                lines.append(f"Heaviest day is typically **{heaviest}** and lightest is **{lightest}**.")

        # ── POSITION ──
        elif cat == "position":
            if "position_name" in df.columns:
                _pos_map = {"gk": "GK", "goalkeeper": "GK", "def": "DEF", "defender": "DEF", "cb": "DEF",
                            "mid": "MID", "midfielder": "MID", "fwd": "FWD", "forward": "FWD", "striker": "FWD"}
                df_c = df.copy()
                df_c["_pg"] = df_c["position_name"].str.lower().str.strip().map(
                    lambda x: next((v for k, v in _pos_map.items() if k in str(x)), "OTHER"))
                grp = df_c.groupby("_pg").agg(
                    players=("player_id", "nunique"),
                    avg_load=("total_player_load", "mean") if "total_player_load" in df_c.columns else ("player_id", "count"),
                ).reset_index()
                lines.append("**Position group breakdown:**")
                for _, r in grp.iterrows():
                    lines.append(f"- **{r['_pg']}**: {r['players']} players, avg load {r.get('avg_load', 0):.0f}")

        # ── GENERAL / FALLBACK — use RAG retrieval ──
        if not lines or cat == "general":
            results = self.query(question, top_k=8)
            if results:
                lines.append(f"Here's what I found related to your question:")
                lines.append("")
                for r in results[:5]:
                    lines.append(f"- {r['text']}")
                expert = [r for r in results if r["type"] == "expert_rule"]
                if expert:
                    lines.append("")
                    lines.append("**Relevant sports science principles:**")
                    for r in expert[:3]:
                        lines.append(f"- {r['text']}")
            else:
                lines.append("I couldn't find specific data for that question. Try asking about:")
                lines.append("- Player availability or injuries")
                lines.append("- Training load and ACWR")
                lines.append("- Speed and velocity zones")
                lines.append("- Wellness and recovery")
                lines.append("- Match performance")

        lines.append("")
        lines.append("---")
        lines.append("*Generated from your CU Boulder data.*")
        return "\n".join(lines)

    def answer_for_coach(self, question: str, top_k: int = 8) -> Tuple[str, List[Dict]]:
        """Main entry: data-driven smart_answer first (instant), then optional supporting docs."""
        smart = None
        if not self._last_df.empty:
            try:
                smart = self.smart_answer(question, self._last_df, self._last_game_all, self._last_team_match)
            except Exception:
                smart = None

        results = self.query(question, top_k=top_k) if self.knowledge_base else []

        if smart and len(smart.strip()) > 80:
            return smart, results

        if not results:
            return (
                "I don't have enough information to answer this. "
                "Try asking about injuries, load, readiness, match performance, velocity, or wellness.",
                [],
            )
        lines = ["Here's what I found:", ""]
        for r in results[:6]:
            lines.append(f"- {r['text']}")
        expert = [r for r in results if r["type"] == "expert_rule"]
        if expert:
            lines.append("")
            lines.append("**Keep in mind:**")
            for r in expert[:2]:
                lines.append(f"- {r['text']}")
        lines.append("")
        lines.append("---")
        lines.append("*Generated from your CU data.*")
        return "\n".join(lines), results

    @staticmethod
    def follow_up_suggestions(question: str) -> List[str]:
        """Return 2-3 contextual follow-up questions based on the topic."""
        cat = OfflineRAGSystem._classify_question(question)
        _suggestions = {
            "load":              ["Which players need a lighter session?",
                                  "What does the weekly load pattern look like?",
                                  "Who has high ACWR right now?"],
            "acwr":              ["Which players are at injury risk?",
                                  "How should we manage load this week?",
                                  "What are the readiness scores?"],
            "speed":             ["How does speed compare by position?",
                                  "Are players reaching game-speed in training?",
                                  "Who is the fastest forward?"],
            "fatigue":           ["Which players have high ACWR?",
                                  "How is the team's wellness?",
                                  "Who should rest tomorrow?"],
            "injury":            ["How should we manage returning players' load?",
                                  "What are the team's readiness scores?",
                                  "Who is at high ACWR?"],
            "wellness":          ["Which players have low readiness?",
                                  "Who is reporting high soreness?",
                                  "What does the load look like this week?"],
            "match_performance": ["How does load compare across opponents?",
                                  "Who are the top scorers?",
                                  "What's the injury status?"],
            "player":            ["What are the ACWR values across the squad?",
                                  "Who has the highest load this week?",
                                  "Which players have low wellness?"],
            "week":              ["What does the microcycle look like?",
                                  "Who has high fatigue?",
                                  "How should we plan match-day minus 1?"],
            "position":          ["Which forwards have the highest speed?",
                                  "Compare defender vs midfielder load",
                                  "How is goalkeeper load different?"],
            "general":           ["Who has high ACWR right now?",
                                  "What's the injury status?",
                                  "How is the team's wellness?"],
        }
        return _suggestions.get(cat, _suggestions["general"])

    def generate_insights(self, df_train: pd.DataFrame) -> List[str]:
        """Auto-generate a comprehensive list of data-driven alerts & insights."""
        alerts: List[str] = []
        if df_train.empty:
            return alerts

        # 1 — Load spikes
        if "total_player_load" in df_train.columns:
            avg_l = df_train["total_player_load"].mean()
            max_l = df_train["total_player_load"].max()
            if max_l > avg_l * 2:
                alerts.append(f"⚠️ **Load spike detected** — peak session load ({max_l:.0f}) is {max_l/avg_l:.1f}× the squad average ({avg_l:.0f}). Review periodization.")

        # 2 — ACWR red zone
        if "acwr_ewma_7_28" in df_train.columns:
            high = df_train[df_train["acwr_ewma_7_28"] > 1.3]
            if len(high):
                n_pl = high["player_id"].nunique()
                alerts.append(f"🔴 **Injury risk (ACWR > 1.3)** — {len(high)} sessions across {n_pl} player(s). Prioritise load management for those athletes.")
            low = df_train[df_train["acwr_ewma_7_28"] < 0.8]
            if len(low) > len(df_train) * 0.15:
                alerts.append(f"🟡 **Under-loading** — {100*len(low)/len(df_train):.0f}% of sessions have ACWR < 0.8. Some athletes may be losing fitness.")

        # 3 — Readiness
        if "readiness" in df_train.columns and df_train["readiness"].notna().any():
            last = df_train.sort_values("date").groupby("player_id")["readiness"].last()
            low_r = last[last < 60]
            if len(low_r):
                names = ", ".join(str(x) for x in low_r.index[:5])
                alerts.append(f"🟠 **Low readiness (< 60)** — {len(low_r)} player(s): {names}{'…' if len(low_r)>5 else ''}. Consider lighter training or wellness check.")

        # 4 — Fatigue
        if "fatigue_idx" in df_train.columns and df_train["fatigue_idx"].notna().any():
            last_f = df_train.sort_values("date").groupby("player_id")["fatigue_idx"].last()
            high_f = last_f[last_f > 1.5]
            if len(high_f):
                names = ", ".join(str(x) for x in high_f.index[:5])
                alerts.append(f"🔥 **High fatigue index (> 1.5)** — {len(high_f)} player(s): {names}{'…' if len(high_f)>5 else ''}. Risk of cumulative overload.")

        # 5 — Anomalies
        if "anomaly_flag" in df_train.columns:
            rate = df_train["anomaly_flag"].mean()
            if rate > 0.10:
                alerts.append(f"📊 **Anomaly rate {rate*100:.0f}%** — higher than the 10% threshold. Check for data quality issues or unexpected training patterns.")

        # 6 — Velocity drop
        if "maximum_velocity" in df_train.columns:
            rec = df_train[df_train["date"] >= df_train["date"].max() - pd.Timedelta(days=14)]
            if len(rec) > 10:
                rec_avg = rec["maximum_velocity"].mean()
                all_avg = df_train["maximum_velocity"].mean()
                if rec_avg < all_avg * 0.90:
                    alerts.append(f"🏃 **Velocity dip** — last-14-day avg ({rec_avg:.1f} mph) is {100*(1-rec_avg/all_avg):.0f}% below squad norm ({all_avg:.1f} mph). Possible squad fatigue.")

        # 7 — Injury status
        if "injury_status_actual" in df_train.columns:
            last_status = df_train.sort_values("date").groupby("player_id")["injury_status_actual"].last()
            non_avail = last_status[~last_status.astype(str).str.lower().isin(["available", "nan", ""])]
            if len(non_avail):
                for st_val in ["Out", "Limited", "As Tolerated"]:
                    cnt = (non_avail.astype(str).str.lower() == st_val.lower()).sum()
                    if cnt:
                        alerts.append(f"🏥 **{cnt} player(s) currently '{st_val}'** — monitor return-to-play timelines.")

        # 8 — Wellness
        if "wellness_total" in df_train.columns and df_train["wellness_total"].notna().any():
            last_w = df_train.sort_values("date").groupby("player_id")["wellness_total"].last()
            low_w = last_w[last_w < 12]
            if len(low_w):
                alerts.append(f"😴 **Low wellness (< 12/20)** — {len(low_w)} player(s). Check sleep quality and soreness sub-scores.")

        # 9 — Weekly monotony
        if "total_player_load" in df_train.columns:
            wk = df_train.groupby(df_train["date"].dt.isocalendar().week.astype(int)).agg(
                mean_load=("total_player_load", "mean"), std_load=("total_player_load", "std")).dropna()
            wk["monotony"] = wk["mean_load"] / wk["std_load"].replace(0, np.nan)
            high_m = wk[wk["monotony"] > 2.0]
            if len(high_m):
                alerts.append(f"📈 **High monotony (> 2.0) in {len(high_m)} week(s)** — vary session intensity to reduce illness risk.")

        if not alerts:
            alerts.append("✅ **All clear** — no significant alerts from the current data window.")
        return alerts

# Initialize RAG system
@st.cache_resource(show_spinner=False)
def get_rag_system() -> OfflineRAGSystem:
    return OfflineRAGSystem()

# ==============================================================================
# SIDEBAR — INPUTS
# ==============================================================================
st.sidebar.header("Display / Theme")
theme_mode = st.sidebar.radio("Theme mode", ["Dark", "Medium", "Light"], index=0)
try:
    if theme_mode == "Light":
        pio.templates.default = "plotly_white"
    else:
        # keep our dark custom template for Dark + Medium
        pio.templates.default = "plotly_dark+cu_dark"
except Exception:
    pass

st.sidebar.header("Paths (edit if needed)")

catapult_path = st.sidebar.text_input("Catapult CSV", value=str(DEFAULT_CATAPULT))
sched_2023_path = st.sidebar.text_input("2023 Schedule", value=str(SCHED_2023))
sched_2024_path = st.sidebar.text_input("2024 Schedule", value=str(SCHED_2024))
sched_2025_path = st.sidebar.text_input("2025 Schedule", value=str(SCHED_2025))

g23_dir = st.sidebar.text_input("2023 Game Stat Folder", value=str(GAME_2023_DIR))
g24_dir = st.sidebar.text_input("2024 Game Stat Folder", value=str(GAME_2024_DIR))
g25_dir = st.sidebar.text_input("2025 Game Stat Folder", value=str(GAME_2025_DIR))

st.sidebar.subheader("Injury & Wellness Files")
inj_2023_path = st.sidebar.text_input("2023 Injuries CSV", value=str(INJ_2023))
inj_2024_path = st.sidebar.text_input("2024 Injuries CSV", value=str(INJ_2024))
inj_2025_path = st.sidebar.text_input("2025 Injuries CSV", value=str(INJ_2025))
well_2024_path = st.sidebar.text_input("2024 Wellness CSV", value=str(WELL_2024))
well_2025_path = st.sidebar.text_input("2025 Wellness CSV", value=str(WELL_2025))

if not Path(catapult_path).exists():
    st.error(f"Catapult CSV not found:\n{catapult_path}")
    st.stop()

# Load core data
df_train_raw = load_catapult_csv(catapult_path)

# Load schedules (robust)
try:
    sched_all = load_all_schedules(Path(sched_2023_path), Path(sched_2024_path), Path(sched_2025_path))
except Exception as e:
    st.error("Schedule load failed. This usually means the header row / sheet / Date+Opponent names differ.")
    st.exception(e)
    st.stop()

# Quick schedule debug panel (helps you instantly confirm it found the right sheet)
with st.expander("🔎 Schedule loader debug (what sheet/header was used?)", expanded=False):
    if len(sched_all) == 0:
        st.write("No schedules loaded.")
    else:
        _sched_display = sched_all[["year", "match_date", "Opponent_clean", "venue", "result", "schedule_sheet_used", "schedule_header_row_used"]].sort_values(["year", "match_date"]).head(30).copy()
        if "match_date" in _sched_display.columns:
            _sched_display["match_date"] = pd.to_datetime(_sched_display["match_date"], errors="coerce").dt.strftime("%m/%d/%Y")
        st.dataframe(pretty_cols(_sched_display), use_container_width=True)

# Load game stats folders + join to schedule
game23 = load_game_stats_folder(g23_dir, 2023)
game24 = load_game_stats_folder(g24_dir, 2024)
game25 = load_game_stats_folder(g25_dir, 2025)
game_all = build_game_match_table(sched_all, game23, game24, game25)
game_all = safe_numeric(game_all, ["MIN", "G", "A", "Sh", "SOG", "GA", "Saves"])

injuries_all = load_all_injuries(inj_2023_path, inj_2024_path, inj_2025_path)
wellness_all = load_all_wellness(well_2024_path, well_2025_path)

# Show warnings when injury/wellness data is missing
_missing_data_msgs = []
for _label, _fpath in [("2023 Injuries", inj_2023_path), ("2024 Injuries", inj_2024_path),
                        ("2025 Injuries", inj_2025_path), ("2024 Wellness", well_2024_path),
                        ("2025 Wellness", well_2025_path)]:
    if not Path(_fpath).exists():
        _missing_data_msgs.append(f"**{_label}**: `{_fpath}`")
if _missing_data_msgs:
    st.warning("⚠️ The following data files were **not found** — Health & Wellness and Coach ML tabs may show incomplete data. "
               "Update the paths in the sidebar or place the files in the expected locations.\n\n" +
               "\n".join(_missing_data_msgs))
if injuries_all.empty:
    st.info("ℹ️ No injury data loaded. Injury-related features will default to 'Available'.")
if wellness_all.empty:
    st.info("ℹ️ No wellness data loaded. Wellness scores and readiness components using wellness will be unavailable.")

# Join training -> schedule by date
df_train = df_train_raw.merge(
    sched_all[["match_date", "year", "Opponent_clean", "venue", "result"]],
    how="left",
    left_on=["date", "year"],
    right_on=["match_date", "year"]
)
df_train["Opponent_clean"] = df_train["Opponent_clean"].fillna("Unknown")
df_train["venue"] = df_train["venue"].fillna("Unknown")
df_train["result"] = df_train["result"].fillna("")

# Add MD window labels
df_train = label_matchday_windows(df_train, sched_all)

with st.spinner("Merging injury status..."):
    df_train = merge_injury_status(df_train, injuries_all)
with st.spinner("Merging wellness data..."):
    df_train = merge_wellness(df_train, wellness_all)

# ==============================================================================
# FILTERS
# ==============================================================================
st.sidebar.header("Filters")
years = sorted(df_train["year"].dropna().unique().tolist())
year_sel = st.sidebar.multiselect("Years", years, default=years)

df_train = df_train[df_train["year"].isin(year_sel)].copy()
if len(df_train) == 0:
    st.warning("No training rows after year filtering.")
    st.stop()

min_d, max_d = df_train["date"].min(), df_train["date"].max()
date_range = st.sidebar.date_input("Date range", value=(min_d.date(), max_d.date()))
if isinstance(date_range, (list, tuple)) and len(date_range) == 2:
    df_train = df_train[(df_train["date"] >= pd.to_datetime(date_range[0])) & (df_train["date"] <= pd.to_datetime(date_range[1]))].copy()
elif isinstance(date_range, (list, tuple)) and len(date_range) == 1:
    df_train = df_train[df_train["date"] >= pd.to_datetime(date_range[0])].copy()

pos_vals = sorted(df_train["position_name"].dropna().unique().tolist()) if "position_name" in df_train.columns else []
pos_sel = st.sidebar.multiselect("Positions", pos_vals, default=pos_vals)
if pos_vals:
    df_train = df_train[df_train["position_name"].isin(pos_sel)].copy()

player_vals = sorted(df_train["player_id"].dropna().unique().tolist())
player_sel = st.sidebar.selectbox("Player", ["(All Players)"] + player_vals)

contam = st.sidebar.slider("Anomaly sensitivity", 0.01, 0.10, 0.04, 0.01)

anom_features = [c for c in [
    "total_player_load", "player_load_per_min_est", "total_distance",
    "maximum_velocity", "total_acceleration_load",
    "max_acceleration", "max_deceleration", "explosive_efforts",
    "acwr_ewma_7_28"
] if c in df_train.columns]

df_train = add_anomaly_scores(df_train, anom_features, contamination=contam)

# Add coach-facing indices (ported from dash.py best features)
with st.spinner("Computing readiness / fatigue / velocity zones..."):
    df_train = compute_velocity_zones(df_train)
    df_train = compute_fatigue_index(df_train)
    df_train = compute_readiness(df_train)
weekly_team = compute_weekly_monotony_strain(df_train)

# Velocity flags removed from sidebar per coach request

# ==============================================================================
# AGGREGATIONS
# ==============================================================================
team_match = (
    df_train.groupby(["date", "year", "Opponent_clean", "venue", "result"], as_index=False)
            .agg(
                team_load=("total_player_load", "sum") if "total_player_load" in df_train.columns else ("player_id", "count"),
                team_hsd=("hsd_m", "sum") if "hsd_m" in df_train.columns else ("player_id", "count"),
                team_sprint=("sprint_m", "sum") if "sprint_m" in df_train.columns else ("player_id", "count"),
                team_accel=("total_acceleration_load", "sum") if "total_acceleration_load" in df_train.columns else ("player_id", "count"),
                mean_acwr=("acwr_ewma_7_28", "mean") if "acwr_ewma_7_28" in df_train.columns else ("player_id", "count"),
                anomaly_sessions=("anomaly_flag", "sum"),
                players=("player_id", "nunique")
            )
)

player_match = (
    df_train.groupby(["player_id", "date", "year", "Opponent_clean", "venue", "result", "position_name"], as_index=False)
            .agg(
                load=("total_player_load", "sum") if "total_player_load" in df_train.columns else ("player_id", "count"),
                hsd=("hsd_m", "sum") if "hsd_m" in df_train.columns else ("player_id", "count"),
                sprint=("sprint_m", "sum") if "sprint_m" in df_train.columns else ("player_id", "count"),
                accel=("total_acceleration_load", "sum") if "total_acceleration_load" in df_train.columns else ("player_id", "count"),
                vmax=("maximum_velocity", "max") if "maximum_velocity" in df_train.columns else ("player_id", "count"),
                acwr=("acwr_ewma_7_28", "mean") if "acwr_ewma_7_28" in df_train.columns else ("player_id", "count"),
                anom=("anomaly_flag", "max")
            )
)

# ==============================================================================
# TABS
# ==============================================================================
(tab_exec, tab_team, tab_player, tab_match, tab_game, tab_md, tab_vs,
 tab_inj, tab_pos, tab_anom, tab_forecast, tab_ml, tab_rag) = st.tabs([
    "🏠 Dashboard",
    "📊 Team",
    "🧍 Player",
    "🏟 Match Center",
    "📄 Game Stats",
    "📅 Microcycle",
    "🆚 Opponent",
    "🏥 Health & Wellness",
    "👥 Position Groups",
    "⚠️ Anomalies",
    "📈 Forecast",
    "🧠 Coach ML",
    "💬 Chat with Data",
])

# ==============================================================================
# TAB: EXECUTIVE DASHBOARD
# ==============================================================================
with tab_exec:
    st.header("🏠 Dashboard Overview")
    coach_guide("Dashboard", [
        "One-glance team health: **injury board**, **wellness**, **load KPIs**, **smart alerts**.",
        "Start here every morning — spend 2–3 minutes before making session decisions.",
        "Red or amber tiles = jump into the relevant tab (Health & Wellness, Coach ML) to dig deeper.",
    ])
    st.markdown("---")

    n_players = df_train["player_id"].nunique()
    n_training_days = df_train["date"].nunique()
    avg_load = df_train["total_player_load"].mean()
    max_vel = df_train["maximum_velocity"].max() if "maximum_velocity" in df_train.columns else 0
    avg_acwr = df_train["acwr_ewma_7_28"].mean() if "acwr_ewma_7_28" in df_train.columns else 0

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Players", f"{n_players}")
    c2.metric("Training Days", f"{n_training_days}")
    c3.metric("Avg Load", f"{avg_load:,.0f}")
    c4.metric("Max Velocity", f"{max_vel:.1f} mph")
    c5.metric("Avg ACWR", f"{avg_acwr:.2f}")

    _info_c1, _info_c2, _info_c3 = st.columns(3)
    with _info_c1:
        metric_tooltip("total_player_load")
    with _info_c2:
        metric_tooltip("maximum_velocity")
    with _info_c3:
        metric_tooltip("acwr")

    st.markdown("---")

    inj_summary_cols = st.columns(4)
    _inj_col = "injury_status_actual" if "injury_status_actual" in df_train.columns else ("injury_status" if "injury_status" in df_train.columns else None)
    if _inj_col:
        latest_per_player = df_train[df_train[_inj_col].notna() & (df_train[_inj_col].astype(str).str.strip() != "")].sort_values("date").groupby("player_id").tail(1)
        n_avail = (latest_per_player[_inj_col] == "Available").sum()
        n_out = (latest_per_player[_inj_col] == "Out").sum()
        n_limited = (latest_per_player[_inj_col] == "Limited").sum()
        n_astol = (latest_per_player[_inj_col] == "As Tolerated").sum()
        inj_summary_cols[0].metric("✅ Available", n_avail)
        inj_summary_cols[1].metric("🔴 Out", n_out)
        inj_summary_cols[2].metric("🟡 Limited", n_limited)
        inj_summary_cols[3].metric("🟠 As Tolerated", n_astol)
    else:
        for c_ in inj_summary_cols:
            c_.info("No injury data")

    # Wellness summary tiles
    _well_cols = st.columns(5)
    if "wellness_total" in df_train.columns and df_train["wellness_total"].notna().sum() > 0:
        _latest_well = df_train.dropna(subset=["wellness_total"]).sort_values("date").groupby("player_id").tail(1)
        _avg_well = _latest_well["wellness_total"].mean()
        _avg_phys = _latest_well["physical_score"].mean() if "physical_score" in _latest_well.columns else 0
        _avg_ment = _latest_well["mental_score"].mean() if "mental_score" in _latest_well.columns else 0
        _avg_sleep = _latest_well["sleep_score"].mean() if "sleep_score" in _latest_well.columns else 0
        _avg_sore = _latest_well["soreness_score"].mean() if "soreness_score" in _latest_well.columns else 0
        _well_cols[0].metric("🧠 Wellness (Avg)", f"{_avg_well:.1f}/20")
        _well_cols[1].metric("Physical", f"{_avg_phys:.1f}/5")
        _well_cols[2].metric("Mental", f"{_avg_ment:.1f}/5")
        _well_cols[3].metric("Sleep", f"{_avg_sleep:.1f}/5")
        _well_cols[4].metric("Soreness", f"{_avg_sore:.1f}/5")
    else:
        for c_ in _well_cols:
            c_.info("No wellness data")

    st.markdown("---")

    if "total_player_load" in df_train.columns:
        team_daily = df_train.groupby("date").agg(
            avg_load=("total_player_load", "mean"),
            max_load=("total_player_load", "max"),
            n_players=("player_id", "nunique"),
        ).reset_index()
        fig_squad = go.Figure()
        fig_squad.add_trace(go.Scatter(
            x=team_daily["date"], y=team_daily["avg_load"], mode="lines",
            name="Avg Load", line=dict(color=CU_GOLD, width=3),
            fill="tozeroy", fillcolor="rgba(207,184,124,0.1)",
        ))
        fig_squad.add_trace(go.Scatter(
            x=team_daily["date"], y=team_daily["max_load"], mode="lines",
            name="Max Load", line=dict(color="#ef4444", width=1, dash="dot"),
        ))
        fig_squad.update_layout(
            title="Squad Load Trend",
            height=380, margin=dict(l=40, r=20, t=50, b=30),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
            xaxis=dict(tickformat="%m/%d/%Y"),
        )
        fig_squad.add_hline(y=avg_load, line_dash="dot", line_color="#CFB87C",
                             annotation_text=f"Avg ({avg_load:,.0f})")
        st.plotly_chart(fig_squad, use_container_width=True)

    col_l, col_r = st.columns(2)
    with col_l:
        if "acwr_ewma_7_28" in df_train.columns:
            acwr_vals = df_train.dropna(subset=["acwr_ewma_7_28"])
            n_safe = ((acwr_vals["acwr_ewma_7_28"] >= 0.8) & (acwr_vals["acwr_ewma_7_28"] <= 1.3)).sum()
            n_total = len(acwr_vals)
            pct_safe = n_safe / n_total * 100 if n_total > 0 else 0
            fig_gauge = go.Figure(go.Indicator(
                mode="gauge+number",
                value=pct_safe,
                title={"text": "% Sessions in ACWR Sweet Spot (0.8-1.3)"},
                gauge=dict(
                    axis=dict(range=[0, 100]),
                    bar=dict(color=CU_GOLD),
                    steps=[
                        dict(range=[0, 50], color="#1e1e2f"),
                        dict(range=[50, 75], color="#2a2a40"),
                        dict(range=[75, 100], color="#1a2e1a"),
                    ],
                    threshold=dict(line=dict(color="#22c55e", width=4), thickness=0.75, value=75),
                ),
                number=dict(suffix="%"),
            ))
            fig_gauge.update_layout(height=300, margin=dict(l=20, r=20, t=60, b=20))
            st.plotly_chart(fig_gauge, use_container_width=True)

    with col_r:
        if "wellness_total" in df_train.columns:
            well_avg = df_train.dropna(subset=["wellness_total"])
            if not well_avg.empty:
                daily_well = well_avg.groupby("date")["wellness_total"].mean().reset_index()
                fig_well_t = go.Figure(go.Scatter(
                    x=daily_well["date"], y=daily_well["wellness_total"],
                    mode="lines+markers", marker=dict(size=4, color="#22c55e"),
                    line=dict(color="#22c55e", width=2),
                    fill="tozeroy", fillcolor="rgba(34,197,94,0.08)",
                ))
                fig_well_t.update_layout(
                    title="Squad Avg Wellness Over Time", template="plotly_dark",
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    height=300, margin=dict(l=40, r=20, t=50, b=30),
                    yaxis_title="Wellness (out of 20)",
                    xaxis=dict(tickformat="%m/%d/%Y"),
                )
                st.plotly_chart(fig_well_t, use_container_width=True)
            else:
                st.info("No wellness data in selected range.")
        else:
            st.info("No wellness column found.")

    # ── PLAYER RISK MATRIX (unique feature) ──
    if "readiness" in df_train.columns and "fatigue_idx" in df_train.columns:
        st.markdown("---")
        st.markdown("### 🎯 Player Risk Matrix")
        st.caption("Each bubble is a player (latest session). **Top-left = high risk** (high fatigue, low readiness). Bubble size = training load.")
        _rm_c1, _rm_c2 = st.columns(2)
        with _rm_c1:
            metric_tooltip("readiness")
        with _rm_c2:
            metric_tooltip("fatigue_idx")

        _latest_risk = df_train.sort_values("date").groupby("player_id").tail(1).copy()
        _latest_risk = _latest_risk.dropna(subset=["readiness", "fatigue_idx"])
        if "total_player_load" in _latest_risk.columns:
            _latest_risk["total_player_load"] = _latest_risk["total_player_load"].fillna(_latest_risk["total_player_load"].median())
        if len(_latest_risk) > 3:
            _status_col = "injury_status_actual" if "injury_status_actual" in _latest_risk.columns else None
            fig_risk = px.scatter(
                _latest_risk, x="readiness", y="fatigue_idx",
                size="total_player_load" if "total_player_load" in _latest_risk.columns else None,
                color=_status_col,
                color_discrete_map={"Available": "#22c55e", "Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY} if _status_col else None,
                hover_data=["player_id"],
                title="Readiness vs Fatigue — Where Is Each Player?",
                labels={"readiness": pretty("readiness"), "fatigue_idx": pretty("fatigue_idx"),
                         "total_player_load": pretty("total_player_load")},
            )
            fig_risk.add_hrect(y0=1.5, y1=_latest_risk["fatigue_idx"].max() * 1.1, fillcolor="rgba(239,68,68,0.05)", line_width=0)
            fig_risk.add_vrect(x0=0, x1=50, fillcolor="rgba(239,68,68,0.05)", line_width=0)
            fig_risk.add_annotation(x=25, y=_latest_risk["fatigue_idx"].max() * 0.9, text="⚠️ HIGH RISK ZONE",
                                     showarrow=False, font=dict(color="#ef4444", size=12))
            fig_risk.add_annotation(x=85, y=0.8, text="✅ GOOD ZONE",
                                     showarrow=False, font=dict(color="#22c55e", size=12))
            fig_risk.update_layout(
                xaxis_title="Readiness (0–100)", yaxis_title="Fatigue Index",
                height=450, margin=dict(l=40, r=20, t=50, b=30),
            )
            st.plotly_chart(fig_risk, use_container_width=True)

    # ── SQUAD LOAD HEATMAP (last 14 team session dates) ──
    st.markdown("---")
    st.markdown("### 🔥 Player Load — Last 14 Sessions")
    st.caption("Each row is a player, each column is a recent team session date. Color intensity shows load level. Red = spike, green = low.")

    if "total_player_load" in df_train.columns:
        _recent_dates = sorted(df_train["date"].dropna().unique())[-14:]
        _spark_data = df_train[df_train["date"].isin(_recent_dates)].copy()
        _spark_data = _spark_data[["player_id", "date", "total_player_load"]].copy()
        _spark_data = _spark_data.sort_values("date")
        _spark_data["date_str"] = _spark_data["date"].dt.strftime("%m/%d")

        _date_order = _spark_data.drop_duplicates("date").sort_values("date")["date_str"].tolist()

        _heat_pivot = _spark_data.pivot_table(
            index="player_id", columns="date_str", values="total_player_load", aggfunc="mean"
        )
        _heat_pivot = _heat_pivot.reindex(columns=_date_order)

        _heat_pivot = _heat_pivot.dropna(how="all")

        if not _heat_pivot.empty:
            _last_col = _heat_pivot.columns[-1] if len(_heat_pivot.columns) > 0 else None
            _player_order = _heat_pivot[_last_col].sort_values(ascending=False, na_position="last").index.tolist() if _last_col else _heat_pivot.index.tolist()
            _heat_pivot = _heat_pivot.reindex(_player_order)

            _player_labels = []
            for pid in _heat_pivot.index:
                short = str(pid)
                if short.startswith("PID_"):
                    short = "Player " + short[4:8].upper()
                _player_labels.append(short)

            fig_heat = px.imshow(
                _heat_pivot.values,
                x=_heat_pivot.columns.tolist(),
                y=_player_labels,
                color_continuous_scale="YlOrRd",
                aspect="auto",
                title="Player Load Heatmap (last 14 team sessions)",
                labels=dict(x="Session Date", y="Player", color="Load"),
            )
            n_players = len(_heat_pivot)
            fig_heat.update_layout(
                height=max(400, 28 * n_players + 120),
                margin=dict(l=10, r=10, t=50, b=40),
                yaxis=dict(tickfont=dict(size=max(8, min(12, 400 // max(n_players, 1))))),
            )
            fig_heat.update_traces(
                text=np.where(np.isnan(_heat_pivot.values), "", _heat_pivot.values.round(0).astype(int).astype(str)),
                texttemplate="%{text}",
                textfont_size=max(7, min(10, 350 // max(n_players, 1))),
            )
            with st.expander("Show player load heatmap", expanded=True):
                st.plotly_chart(fig_heat, use_container_width=True)

    # ── MATCH-READINESS QUICK CHECK ──
    if "match_hs_flag" in df_train.columns and "match_sprint_flag" in df_train.columns:
        st.markdown("---")
        st.markdown("### 🎯 Match-Readiness Quick Check")
        st.caption("How many players have reached their position's match-based high-speed and sprint thresholds in the last 14 days?")
        _recent = df_train[df_train["date"] >= df_train["date"].max() - pd.Timedelta(days=14)]
        _mr_by_player = _recent.groupby("player_id").agg(
            hit_hs=("match_hs_flag", "any"),
            hit_sprint=("match_sprint_flag", "any"),
        ).reset_index()
        n_hs = int(_mr_by_player["hit_hs"].sum())
        n_sp = int(_mr_by_player["hit_sprint"].sum())
        n_pl = len(_mr_by_player)
        mr1, mr2, mr3 = st.columns(3)
        mr1.metric("Players Assessed", f"{n_pl}")
        mr2.metric("Hit HS Threshold (14d)", f"{n_hs}/{n_pl}", delta=f"{100*n_hs/max(n_pl,1):.0f}%")
        mr3.metric("Hit Sprint Threshold (14d)", f"{n_sp}/{n_pl}", delta=f"{100*n_sp/max(n_pl,1):.0f}%")

    # Data Quality Summary
    st.markdown("---")
    st.markdown("### 📋 Data Quality at a Glance")
    total_rows = len(df_train)
    dq_cols = {
        "total_player_load": "Player Load",
        "maximum_velocity": "Max Velocity",
        "acwr_ewma_7_28": "ACWR",
        "wellness_total": "Wellness",
        "injury_status_actual": "Injury Status",
        "anomaly_score": "Anomaly Score",
    }
    dq_rows = []
    for col, label in dq_cols.items():
        if col in df_train.columns:
            non_null = df_train[col].notna().sum()
            pct = non_null / total_rows * 100
            dq_rows.append({"Metric": label, "Available": f"{non_null:,}", "Coverage": f"{pct:.1f}%"})
        else:
            dq_rows.append({"Metric": label, "Available": "—", "Coverage": "0%"})
    st.dataframe(pd.DataFrame(dq_rows), use_container_width=True, hide_index=True)
    st.caption(
        f"Total training rows: **{total_rows:,}** • Date range: "
        f"**{df_train['date'].min().strftime('%m/%d/%Y')}** to **{df_train['date'].max().strftime('%m/%d/%Y')}**"
    )

    # ── TEAM LOAD CALENDAR HEATMAP ──
    st.markdown("---")
    st.markdown("### 🗓 Training Load Calendar")
    st.caption("Daily team average load shown as a heatmap. Darker = higher intensity.")
    if "total_player_load" in df_train.columns:
        _cal = df_train.groupby("date")["total_player_load"].mean().reset_index()
        _cal.columns = ["date", "avg_load"]
        _cal["date"] = pd.to_datetime(_cal["date"])
        _cal["week"] = _cal["date"].dt.isocalendar().week.astype(int)
        _cal["dow"] = _cal["date"].dt.dayofweek
        _cal["dow_name"] = _cal["date"].dt.strftime("%a")
        _cal["year_month"] = _cal["date"].dt.to_period("M").astype(str)

        _cal_pivot = _cal.pivot_table(index="dow", columns="year_month", values="avg_load", aggfunc="mean")
        _dow_labels = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
        _cal_pivot.index = [_dow_labels[i] if i < 7 else str(i) for i in _cal_pivot.index]

        fig_cal = px.imshow(
            _cal_pivot.values, x=_cal_pivot.columns.tolist(), y=_cal_pivot.index.tolist(),
            color_continuous_scale="YlOrRd", aspect="auto",
            title="Avg Team Load by Day-of-Week × Month",
        )
        fig_cal.update_layout(height=300, margin=dict(l=50, r=20, t=50, b=30),
                               xaxis_title="Month", yaxis_title="")
        st.plotly_chart(fig_cal, use_container_width=True)

    # Weekly report export
    st.markdown("---")
    st.markdown("### 🗞 Weekly Coach Report (Export)")
    st.caption("Generates a report from the currently filtered data.")

    def _make_weekly_report_md(df_: pd.DataFrame) -> str:
        if df_.empty:
            return "No data available under current filters."
        latest_date = pd.to_datetime(df_["date"]).max()
        week = int(latest_date.isocalendar().week)
        year = int(latest_date.year)

        latest = df_.sort_values("date").groupby("player_id", as_index=False).tail(1).copy()
        lines = []
        lines.append(f"# CU WSOC Weekly Report — Week {week}, {year}")
        lines.append(f"Filtered date range: {df_['date'].min().strftime('%m/%d/%Y')} → {df_['date'].max().strftime('%m/%d/%Y')}")
        lines.append("")
        lines.append("## Squad snapshot")
        if "total_player_load" in df_.columns:
            lines.append(f"- Avg load: **{df_['total_player_load'].mean():.0f}**")
        if "acwr_ewma_7_28" in df_.columns:
            hi = int((df_["acwr_ewma_7_28"] > 1.3).sum())
            lines.append(f"- Sessions with ACWR > 1.3: **{hi}**")
        if "anomaly_flag" in df_.columns:
            lines.append(f"- Anomaly rate: **{100*df_['anomaly_flag'].mean():.1f}%**")
        if "wellness_total" in df_.columns:
            w = df_.dropna(subset=["wellness_total"])
            if len(w):
                lines.append(f"- Avg wellness: **{w['wellness_total'].mean():.1f}/20**")
        lines.append("")

        lines.append("## Players to watch")
        # Low readiness
        if "readiness" in latest.columns:
            low = latest.dropna(subset=["readiness"]).sort_values("readiness").head(8)
            if len(low):
                lines.append("### Low readiness")
                for _, r in low.iterrows():
                    lines.append(f"- Player **{r['player_id']}** readiness **{r['readiness']:.0f}/100**")
                lines.append("")
        # High fatigue
        if "fatigue_idx" in latest.columns:
            hf = latest.dropna(subset=["fatigue_idx"]).sort_values("fatigue_idx", ascending=False).head(8)
            if len(hf):
                lines.append("### High fatigue index")
                for _, r in hf.iterrows():
                    lines.append(f"- Player **{r['player_id']}** fatigue idx **{r['fatigue_idx']:.2f}**")
                lines.append("")
        # Injury status
        if "injury_status_actual" in latest.columns:
            outp = latest[latest["injury_status_actual"].astype(str).str.lower().isin(["out", "limited", "as tolerated"])]
            if len(outp):
                lines.append("### Injury status (non-available)")
                for _, r in outp.sort_values("injury_status_actual").iterrows():
                    lines.append(f"- Player **{r['player_id']}** status **{r['injury_status_actual']}**")
                lines.append("")
        return "\n".join(lines)

    report_md = _make_weekly_report_md(df_train)
    st.download_button(
        "📥 Download report (Markdown)",
        data=report_md.encode("utf-8"),
        file_name="weekly_coach_report.md",
        mime="text/markdown",
    )

    if DOCX_AVAILABLE:
        if st.button("Generate DOCX report", key="make_docx"):
            try:
                doc = Document()
                for line in report_md.splitlines():
                    if line.startswith("# "):
                        doc.add_heading(line.replace("# ", ""), level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line.replace("## ", ""), level=2)
                    elif line.startswith("### "):
                        doc.add_heading(line.replace("### ", ""), level=3)
                    elif line.startswith("- "):
                        doc.add_paragraph(line.replace("- ", ""), style="List Bullet")
                    else:
                        doc.add_paragraph(line)
                import io
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button(
                    "📥 Download report (DOCX)",
                    data=buf.getvalue(),
                    file_name="weekly_coach_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.warning("DOCX generation failed.")
                st.exception(e)
    else:
        st.caption("DOCX export optional: install `python-docx` to enable.")


# ==============================================================================
# TAB: TEAM TRAINING
# ==============================================================================
with tab_team:
    st.subheader("Team Training Overview")
    coach_guide("Team", [
        "Squad-level **load, distance, speed, ACWR** trends across the season.",
        "Use to answer: *is the team getting fitter? did last week bury us?*",
        "Compare to Microcycle when planning the next training block.",
    ])

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Sessions", int(len(df_train)))
    c2.metric("Players", int(df_train["player_id"].nunique()))
    if "total_player_load" in df_train.columns:
        c3.metric("Avg Load", f"{df_train['total_player_load'].mean():.1f}")
    if "player_load_per_min_est" in df_train.columns:
        c4.metric("Avg Load/min", f"{df_train['player_load_per_min_est'].mean():.2f}")

    _tt_c1, _tt_c2 = st.columns(2)
    with _tt_c1:
        metric_tooltip("total_player_load")
    with _tt_c2:
        metric_tooltip("player_load_per_min")

    daily = df_train.groupby("date", as_index=False).agg(
        players=("player_id", "nunique"),
        load=("total_player_load", "mean") if "total_player_load" in df_train.columns else ("player_id", "count"),
        hsd=("hsd_m", "mean") if "hsd_m" in df_train.columns else ("player_id", "count"),
        sprint=("sprint_m", "mean") if "sprint_m" in df_train.columns else ("player_id", "count"),
        acwr=("acwr_ewma_7_28", "mean") if "acwr_ewma_7_28" in df_train.columns else ("player_id", "count"),
    )

    colA, colB = st.columns(2)
    with colA:
        st.plotly_chart(px.line(daily, x="date", y="load", markers=True, title="Avg Player Load (Daily)",
                               labels={"date": "Date", "load": "Load"}).update_layout(xaxis=dict(tickformat="%m/%d/%Y")),
                        use_container_width=True)
    with colB:
        _fig_acwr_daily = px.line(daily, x="date", y="acwr", markers=True, title="Avg ACWR (Daily)",
                                   labels={"date": "Date", "acwr": "ACWR"})
        _fig_acwr_daily.add_hline(y=0.8, line_dash="dash", line_color="#CFB87C", annotation_text="Min (0.8)")
        _fig_acwr_daily.add_hline(y=1.3, line_dash="dash", line_color="#ef4444", annotation_text="Max (1.3)")
        _fig_acwr_daily.update_layout(xaxis=dict(tickformat="%m/%d/%Y"))
        st.plotly_chart(_fig_acwr_daily, use_container_width=True)

    if "position_name" in df_train.columns and "total_player_load" in df_train.columns:
        st.plotly_chart(px.box(df_train, x="position_name", y="total_player_load", points="outliers",
                               title="Load Distribution by Position",
                               labels={"position_name": pretty("position_name"), "total_player_load": pretty("total_player_load")}),
                        use_container_width=True)

    st.download_button(
        "⬇️ Download filtered training sessions (CSV)",
        data=df_train.to_csv(index=False).encode("utf-8"),
        file_name="training_filtered.csv",
        mime="text/csv"
    )

# ==============================================================================
# TAB: PLAYER TRAINING
# ==============================================================================
with tab_player:
    st.subheader("Player Training Explorer")
    coach_guide("Player", [
        "**Single-athlete deep dive**: load curve, ACWR, readiness, fatigue, wellness, injury history.",
        "The **“What’s impacting this player most right now?”** panel ranks the current drivers in plain English.",
        "Open the **Injury log details** expander to see every documented row for this athlete.",
    ])

    if player_sel == "(All Players)":
        st.info("Pick a player from the sidebar to view detailed trends.")
    else:
        p = df_train[df_train["player_id"] == player_sel].sort_values("date").copy()

        a, b, c, d = st.columns(4)
        a.metric("Sessions", int(len(p)))
        b.metric("Anomalies", int(p["anomaly_flag"].sum()))
        if "total_player_load" in p.columns:
            c.metric("Avg Load", f"{p['total_player_load'].mean():.1f}")
        if "acwr_ewma_7_28" in p.columns:
            d.metric("Avg ACWR", f"{p['acwr_ewma_7_28'].mean():.2f}")
        if "readiness" in p.columns and p["readiness"].notna().any():
            st.metric("Readiness (latest)", f"{p['readiness'].dropna().iloc[-1]:.0f}/100")
        if "fatigue_idx" in p.columns and p["fatigue_idx"].notna().any():
            st.metric("Fatigue idx (latest)", f"{p['fatigue_idx'].dropna().iloc[-1]:.2f}")
        if "vmax_zone" in p.columns and p["vmax_zone"].notna().any():
            st.metric("Vmax zone (latest)", str(p["vmax_zone"].iloc[-1]))

        st.markdown("### 🧭 What’s impacting this player most right now?")
        _drivers = summarize_player_drivers(p)
        if not _drivers:
            st.info("Not enough data to generate drivers yet (need recent sessions and/or wellness/injury fields).")
        else:
            _top = _drivers[0]
            st.success(f"Top driver: **{_top['title']}** — {_top['detail']}")
            for dd in _drivers[1:4]:
                st.write(f"- **{dd['title']}** — {dd['detail']}")
            if len(_drivers) > 4:
                with st.expander("Show all drivers"):
                    for dd in _drivers[4:]:
                        st.write(f"- **{dd['title']}** — {dd['detail']}")

        with st.expander("🏥 Injury log details for this player", expanded=False):
            if "injuries_all" not in globals() or injuries_all is None or len(injuries_all) == 0:
                st.info("No injury log loaded.")
            else:
                _inj_p = injuries_all[injuries_all["player_id"].astype(str) == str(player_sel)].copy()
                if _inj_p.empty:
                    st.info("No injury rows found for this player in the loaded injury files.")
                else:
                    # sort newest-first by best available date column
                    for _dc in ["status_start", "injury_date", "report_date"]:
                        if _dc in _inj_p.columns:
                            _inj_p[_dc] = pd.to_datetime(_inj_p[_dc], errors="coerce")
                    _sort_col = "status_start" if "status_start" in _inj_p.columns else ("injury_date" if "injury_date" in _inj_p.columns else ("report_date" if "report_date" in _inj_p.columns else None))
                    if _sort_col:
                        _inj_p = _inj_p.sort_values(_sort_col, ascending=False)

                    # show key fields first, then any extra columns from the original file
                    base_cols = [c for c in ["injury_status", "status_start", "injury_date", "report_date", "days_in_status", "year"] if c in _inj_p.columns]
                    extra_cols = [c for c in _inj_p.columns if c not in (["player_id"] + base_cols)]
                    show_cols = ["player_id"] + base_cols + extra_cols

                    _show = _inj_p[show_cols].head(30).copy()
                    for _dc in ["status_start", "injury_date", "report_date"]:
                        if _dc in _show.columns:
                            _show[_dc] = pd.to_datetime(_show[_dc], errors="coerce").dt.strftime("%m/%d/%Y")
                    st.dataframe(pretty_cols(_show), use_container_width=True, hide_index=True)

        _pt_c1, _pt_c2, _pt_c3 = st.columns(3)
        with _pt_c1:
            metric_tooltip("readiness")
        with _pt_c2:
            metric_tooltip("fatigue_idx")
        with _pt_c3:
            metric_tooltip("acwr")

        col1, col2 = st.columns(2)
        with col1:
            if "total_player_load" in p.columns:
                st.plotly_chart(px.line(p, x="date", y="total_player_load", markers=True, title=pretty("total_player_load"),
                                       labels={"date": "Date", "total_player_load": pretty("total_player_load")}).update_layout(xaxis=dict(tickformat="%m/%d/%Y")),
                                use_container_width=True)
        with col2:
            if "player_load_per_min_est" in p.columns:
                st.plotly_chart(px.line(p, x="date", y="player_load_per_min_est", markers=True, title="Load/Min (Intensity)").update_layout(xaxis=dict(tickformat="%m/%d/%Y")),
                                use_container_width=True)

        if "acwr_ewma_7_28" in p.columns:
            fig = px.line(p, x="date", y="acwr_ewma_7_28", markers=True, title="EWMA ACWR (7:28)",
                          labels={"date": "Date", "acwr_ewma_7_28": pretty("acwr_ewma_7_28")})
            fig.add_hline(y=0.8, line_dash="dash", line_color="#CFB87C", annotation_text="Min (0.8)")
            fig.add_hline(y=1.3, line_dash="dash", line_color="#ef4444", annotation_text="Max (1.3)")
            fig.update_layout(xaxis=dict(tickformat="%m/%d/%Y"))
            st.plotly_chart(fig, use_container_width=True)

        # Readiness + fatigue trends (ported from dash.py)
        r1, r2 = st.columns(2)
        with r1:
            if "readiness" in p.columns and p["readiness"].notna().any():
                _latest_readiness = float(p["readiness"].dropna().iloc[-1])
                _readiness_color = "#ef4444" if _latest_readiness < 40 else "#f59e0b" if _latest_readiness < 60 else "#22c55e" if _latest_readiness < 80 else CU_GOLD
                fig_ready_gauge = go.Figure(go.Indicator(
                    mode="gauge+number+delta",
                    value=_latest_readiness,
                    delta={"reference": float(p["readiness"].dropna().mean()), "suffix": " vs avg"},
                    title={"text": "Current Readiness"},
                    gauge=dict(
                        axis=dict(range=[0, 100]),
                        bar=dict(color=_readiness_color),
                        steps=[
                            dict(range=[0, 40], color="rgba(239,68,68,0.15)"),
                            dict(range=[40, 60], color="rgba(245,158,11,0.15)"),
                            dict(range=[60, 80], color="rgba(34,197,94,0.15)"),
                            dict(range=[80, 100], color="rgba(99,102,241,0.15)"),
                        ],
                        threshold=dict(line=dict(color="#f1f5f9", width=3), thickness=0.8, value=_latest_readiness),
                    ),
                    number=dict(suffix="/100"),
                ))
                fig_ready_gauge.update_layout(height=280, margin=dict(l=20, r=20, t=50, b=10))
                st.plotly_chart(fig_ready_gauge, use_container_width=True)

                _ready_hist = p["readiness"].dropna()
                fig_ready_dist = go.Figure()
                fig_ready_dist.add_trace(go.Histogram(
                    x=_ready_hist, nbinsx=20,
                    marker_color=CU_GOLD, opacity=0.8,
                    name="Sessions",
                ))
                fig_ready_dist.add_vline(x=_latest_readiness, line_dash="dash", line_color="#f59e0b",
                                          annotation_text=f"Latest: {_latest_readiness:.0f}")
                fig_ready_dist.add_vrect(x0=0, x1=40, fillcolor="rgba(239,68,68,0.06)", line_width=0)
                fig_ready_dist.add_vrect(x0=40, x1=60, fillcolor="rgba(245,158,11,0.06)", line_width=0)
                fig_ready_dist.add_vrect(x0=60, x1=80, fillcolor="rgba(34,197,94,0.06)", line_width=0)
                fig_ready_dist.add_vrect(x0=80, x1=100, fillcolor="rgba(99,102,241,0.06)", line_width=0)
                fig_ready_dist.update_layout(
                    title="Readiness Distribution (all sessions)",
                    xaxis_title="Readiness Score", yaxis_title="Sessions",
                    height=250, margin=dict(l=40, r=20, t=50, b=30),
                    showlegend=False,
                )
                st.plotly_chart(fig_ready_dist, use_container_width=True)

        with r2:
            if "fatigue_idx" in p.columns and p["fatigue_idx"].notna().any():
                st.plotly_chart(px.line(p, x="date", y="fatigue_idx", markers=True, title="Fatigue Index (7/28)",
                                       labels={"date": "Date", "fatigue_idx": pretty("fatigue_idx")}).update_layout(xaxis=dict(tickformat="%m/%d/%Y")),
                                use_container_width=True)

        # ── Multi-player comparison ──
        st.markdown("---")
        st.markdown("### 🆚 Compare Players")
        st.caption("Pick a position group first — the player list updates automatically. 'All' shows everyone.")

        def _pos_group_of(x: str) -> str:
            xl = str(x).lower()
            if any(k in xl for k in ["gk", "keeper", "goal"]): return "GK"
            if any(k in xl for k in ["def", "back", "cb", "rb", "lb", "fb"]): return "DEF"
            if any(k in xl for k in ["mid", "cm", "cdm", "cam", "dm", "am"]): return "MID"
            if any(k in xl for k in ["fwd", "forward", "striker", "wing", "att", "st", "cf", "lw", "rw"]): return "FWD"
            return "OTHER"

        _pos_choices = {"All Players": None, "Forwards": "FWD", "Midfielders": "MID", "Defenders": "DEF", "Goalkeepers": "GK"}
        _pos_filter = st.radio("Position group", list(_pos_choices.keys()), horizontal=True, key="cmp_pos_filter")

        all_pids = sorted(df_train["player_id"].dropna().unique().tolist())
        if _pos_choices[_pos_filter] is not None and "position_name" in df_train.columns:
            _grp_tag = _pos_choices[_pos_filter]
            filtered_pids = sorted(
                df_train[df_train["position_name"].apply(_pos_group_of) == _grp_tag]["player_id"].unique().tolist()
            )
        else:
            filtered_pids = all_pids

        cmp_players = st.multiselect(
            "Players to compare (max 6)", filtered_pids,
            default=filtered_pids[:min(4, len(filtered_pids))],
            max_selections=6, key="cmp_multi",
        )

        radar_metrics = [
            ("total_player_load", "Load"),
            ("maximum_velocity", "Top Speed"),
            ("acwr_ewma_7_28", "ACWR"),
            ("wellness_total", "Wellness"),
            ("readiness", "Readiness"),
            ("fatigue_idx", "Fatigue"),
        ]
        radar_metrics = [(c, lbl) for c, lbl in radar_metrics if c in df_train.columns]

        _cmp_colors = [CU_GOLD, "#22c55e", "#f59e0b", "#ef4444", CU_DARK_GRAY, "#06b6d4"]

        if cmp_players and len(radar_metrics) >= 3:
            metric_tooltip("radar_chart")
            # Build latest row per selected player
            _latest_all = df_train[df_train["player_id"].isin(cmp_players)].sort_values("date").groupby("player_id").tail(1)

            # Radar chart
            fig_rad = go.Figure()
            for i, pid in enumerate(cmp_players):
                prow = _latest_all[_latest_all["player_id"] == pid]
                if prow.empty:
                    continue
                prow = prow.iloc[0]
                vals, labels = [], []
                for col, lbl in radar_metrics:
                    squad_col = pd.to_numeric(df_train[col], errors="coerce").dropna()
                    if len(squad_col) < 20:
                        continue
                    vmin, vmax_ = float(squad_col.quantile(0.10)), float(squad_col.quantile(0.90))
                    if vmax_ <= vmin:
                        continue
                    v = float(pd.to_numeric(prow.get(col), errors="coerce")) if pd.notna(prow.get(col)) else np.nan
                    if not np.isfinite(v):
                        continue
                    vals.append(float(np.clip(100 * (v - vmin) / (vmax_ - vmin), 0, 100)))
                    labels.append(lbl)
                if len(vals) >= 3:
                    fig_rad.add_trace(go.Scatterpolar(
                        r=vals, theta=labels, fill="toself",
                        name=str(pid), line=dict(color=_cmp_colors[i % len(_cmp_colors)]),
                    ))

            fig_rad.update_layout(
                title="Player Radar (normalized 0–100 vs squad)",
                polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                showlegend=True, height=520,
            )
            st.plotly_chart(fig_rad, use_container_width=True)

            # Side-by-side metric bars
            st.markdown("### Metric Comparison (latest session)")
            bar_metric = st.selectbox(
                "Metric",
                [lbl for _, lbl in radar_metrics],
                key="cmp_bar_metric",
            )
            bar_col = next((c for c, l in radar_metrics if l == bar_metric), None)
            if bar_col:
                bar_data = _latest_all[["player_id", bar_col]].dropna()
                bar_data[bar_col] = pd.to_numeric(bar_data[bar_col], errors="coerce")
                fig_bar = px.bar(
                    bar_data.sort_values(bar_col, ascending=False),
                    x="player_id", y=bar_col, color="player_id",
                    color_discrete_sequence=_cmp_colors,
                    title=f"{bar_metric} — selected players",
                )
                fig_bar.update_layout(showlegend=False, height=380)
                st.plotly_chart(fig_bar, use_container_width=True)

            # Trend overlay
            st.markdown("### Trend Over Time")
            trend_metric = st.selectbox(
                "Trend metric",
                [c for c in ["total_player_load", "maximum_velocity", "acwr_ewma_7_28", "readiness", "fatigue_idx", "wellness_total"] if c in df_train.columns],
                key="cmp_trend",
            )
            cmp_df = df_train[df_train["player_id"].isin(cmp_players)].copy()
            fig_trend = px.line(
                cmp_df.sort_values("date"), x="date", y=trend_metric,
                color="player_id", markers=False,
                color_discrete_sequence=_cmp_colors,
                title=f"{pretty(trend_metric)} trend — selected players",
            )
            fig_trend.update_layout(height=400, xaxis=dict(tickformat="%m/%d/%Y"))
            st.plotly_chart(fig_trend, use_container_width=True)

            # Summary table
            st.markdown("### Summary Table")
            sum_cols = [c for c, _ in radar_metrics]
            sum_df = _latest_all[["player_id"] + [c for c in sum_cols if c in _latest_all.columns]].copy()
            for sc in sum_cols:
                if sc in sum_df.columns:
                    sum_df[sc] = pd.to_numeric(sum_df[sc], errors="coerce")
            sum_df = sum_df.set_index("player_id").round(2)
            sum_df.columns = [lbl for _, lbl in radar_metrics if _ in sum_df.columns]
            st.dataframe(sum_df, use_container_width=True)
        elif not cmp_players:
            st.info("Select players above to compare.")
        else:
            st.info("Not enough metrics with data for a comparison.")

        show_cols = [c for c in [
            "date", "md_label", "Opponent_clean", "venue", "result", "position_name",
            "total_player_load", "player_load_per_min_est", "hsd_m", "sprint_m",
            "total_acceleration_load", "explosive_efforts", "maximum_velocity",
            "acwr_ewma_7_28", "pct_vs_baseline", "anomaly_score", "anomaly_flag",
            "readiness", "fatigue_idx", "vmax_zone",
            "session_classification"
        ] if c in p.columns]
        _display_p = p[show_cols].sort_values("date", ascending=False).head(40).copy()
        if "date" in _display_p.columns:
            _display_p["date"] = pd.to_datetime(_display_p["date"]).dt.strftime("%m/%d/%Y")
        st.dataframe(pretty_cols(_display_p), use_container_width=True)

# ==============================================================================
# TAB: MATCH CENTER (TRAINING)
# ==============================================================================
with tab_match:
    st.subheader("🏟 Match Center (Training context by opponent/date)")
    coach_guide("Match Center", [
        "Every match lined up against the **training load that preceded it**.",
        "Great for Mondays: *how did the week into the game actually look?*",
        "Pair with **Microcycle** to compare planned taper vs what really happened.",
    ])
    with st.expander("ℹ️ How is **Match Center** data compiled?", expanded=False):
        st.markdown(
            "Training sessions are matched to schedule dates by date and year. "
            "Team totals show the sum of all player loads on each match date from Catapult GPS data.\n\n"
            "**Team Load** = sum of all players' total_player_load on that date.\n"
            "**Mean ACWR** = average ACWR across all players on that date.\n\n"
            "**Example:** On 2024-09-15 vs Baylor, if 18 players trained with loads "
            "[350, 400, 280, ...], team load = sum of all = ~6,200. "
            "This helps identify which opponents correlated with heavier training days."
        )

    years2 = sorted(team_match["year"].dropna().unique().tolist())
    ysel = st.selectbox("Year", ["All"] + years2, key="mc_year")

    view = team_match.copy()
    if ysel != "All":
        view = view[view["year"] == ysel].copy()

    opps = ["All"] + sorted(view["Opponent_clean"].dropna().unique().tolist())
    opp_sel = st.selectbox("Opponent", opps, key="mc_opp")

    if opp_sel != "All":
        view = view[view["Opponent_clean"] == opp_sel].copy()

    st.markdown("### Team totals (training sessions on match date)")
    _view_display = view.sort_values("date", ascending=False).copy()
    if "date" in _view_display.columns:
        _view_display["date"] = pd.to_datetime(_view_display["date"]).dt.strftime("%m/%d/%Y")
    st.dataframe(pretty_cols(_view_display), use_container_width=True)

    colA, colB = st.columns(2)
    with colA:
        st.plotly_chart(px.line(view.sort_values("date"), x="date", y="team_load", markers=True,
                                title="Team Load on Match Dates (from training file)",
                                labels={"date": "Date", "team_load": pretty("team_load")}).update_layout(xaxis=dict(tickformat="%m/%d/%Y")),
                        use_container_width=True)
    with colB:
        _fig_mc_acwr = px.line(view.sort_values("date"), x="date", y="mean_acwr", markers=True,
                               title="Mean ACWR on Match Dates",
                               labels={"date": "Date", "mean_acwr": pretty("mean_acwr")})
        _fig_mc_acwr.add_hline(y=0.8, line_dash="dash", line_color="#CFB87C", annotation_text="Min (0.8)")
        _fig_mc_acwr.add_hline(y=1.3, line_dash="dash", line_color="#ef4444", annotation_text="Max (1.3)")
        _fig_mc_acwr.update_layout(xaxis=dict(tickformat="%m/%d/%Y"))
        st.plotly_chart(_fig_mc_acwr, use_container_width=True)

    st.markdown("### Player breakdown for a match date")
    dates = sorted(player_match["date"].dropna().unique().tolist())
    if len(dates) > 0:
        dsel = st.selectbox("Pick match date", dates[::-1], key="mc_date")
        pm = player_match[player_match["date"] == dsel].sort_values("load", ascending=False)
        _pm_display = pm.copy()
        if "date" in _pm_display.columns:
            _pm_display["date"] = pd.to_datetime(_pm_display["date"]).dt.strftime("%m/%d/%Y")
        st.dataframe(pretty_cols(_pm_display), use_container_width=True)
        st.plotly_chart(px.bar(pm, x="player_id", y="load", title=f"Player Load — {dsel.date()}",
                               labels={"player_id": pretty("player_id"), "load": pretty("load")}),
                        use_container_width=True)
    else:
        st.info("No match dates found (schedule join may be missing).")

# ==============================================================================
# TAB: GAME STATS (CSV folders)
# ==============================================================================
with tab_game:
    st.subheader("📄 Game Stats (from per-match CSV files)")
    coach_guide("Game Stats", [
        "Season **box scores** from the per-match CSVs: goals, shots, possession, minutes, discipline.",
        "Pick any metric to chart it across the season or per opponent.",
        "Read alongside Match Center to tie performance back to the training week.",
    ])
    with st.expander("ℹ️ How are **Game Stats** collected?", expanded=False):
        st.markdown(
            "Game statistics come from individual CSV files per match, containing per-player "
            "box-score data (minutes, goals, assists, shots, saves, etc.).\n\n"
            "**Columns:**\n"
            "- **MIN** — Minutes played\n"
            "- **G** — Goals scored\n"
            "- **A** — Assists\n"
            "- **Sh** — Total shots taken\n"
            "- **SOG** — Shots on goal\n"
            "- **GA** — Goals against (goalkeepers)\n"
            "- **Saves** — Saves made (goalkeepers)\n\n"
            "**Example:** Player PID_abc with MIN=90, G=1, A=2, Sh=5, SOG=3 "
            "played the full match, scored once, assisted twice, took 5 shots (3 on target)."
        )

    if game_all is None or len(game_all) == 0:
        st.warning("No game stat files loaded. Check your folders and file extensions.")
    else:
        gy = st.selectbox("Year", sorted(game_all["year"].dropna().unique().tolist()), key="gs_year")
        g = game_all[game_all["year"] == gy].copy()

        opps = ["All"] + sorted(g["Opponent_clean"].dropna().unique().tolist())
        opp = st.selectbox("Opponent", opps, key="gs_opp")

        if opp != "All":
            g = g[g["Opponent_clean"] == opp].copy()

        md_list = sorted(g["match_date"].dropna().unique().tolist())
        if len(md_list) > 0:
            md_sel = st.selectbox("Match date", md_list[::-1], key="gs_date")
            one = g[g["match_date"] == md_sel].copy()

            st.markdown("### Player match stats")
            show = [c for c in ["player_id", "Type", "No", "Pos", "MIN", "G", "A", "Sh", "SOG", "GA", "Saves",
                                "Opponent_clean", "venue", "result", "file_name"] if c in one.columns]
            st.dataframe(pretty_cols(one[show].sort_values("MIN", ascending=False)), use_container_width=True)

            team_tot = {}
            _stat_labels = {"G": "Goals", "A": "Assists", "Sh": "Shots", "SOG": "Shots on Goal",
                            "GA": "Goals Against", "Saves": "Saves", "MIN": "Minutes"}
            for c in ["G", "A", "Sh", "SOG", "GA", "Saves", "MIN"]:
                if c in one.columns:
                    team_tot[c] = int(pd.to_numeric(one[c], errors="coerce").fillna(0).sum())
            st.markdown("### Team Totals")
            _tt_cols = st.columns(min(len(team_tot), 7))
            for idx, (stat, val) in enumerate(team_tot.items()):
                _tt_cols[idx % len(_tt_cols)].metric(_stat_labels.get(stat, stat), val)

            _avail_stats = [c for c in ["Sh", "SOG", "G", "A", "MIN", "GA", "Saves"] if c in one.columns]
            if _avail_stats:
                _chart_stat = st.selectbox("Select metric to chart", _avail_stats,
                                           format_func=lambda x: _stat_labels.get(x, x), key="gs_chart_stat")
                fig = px.bar(one.sort_values(_chart_stat, ascending=False), x="player_id", y=_chart_stat,
                             title=f"{_stat_labels.get(_chart_stat, _chart_stat)} by Player",
                             color_discrete_sequence=[CU_GOLD])
                fig.update_layout(xaxis_title="Player", yaxis_title=_stat_labels.get(_chart_stat, _chart_stat))
                st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No match_date mapped for this selection (schedule token join issue).")

        st.download_button(
            "⬇️ Download filtered game stats",
            data=g.to_csv(index=False).encode("utf-8"),
            file_name="game_stats_filtered.csv",
            mime="text/csv"
        )

# ==============================================================================
# TAB: MICROCYCLE — Periodization, Day-of-Week, Velocity Zones, Match Readiness
# ==============================================================================
with tab_md:
    st.subheader("📅 Microcycle & Periodization")
    coach_guide("Microcycle", [
        "The **match-week planner**: MD-5 → MD+3 load profile; you want a **clean taper into match day**.",
        "**Monotony & Strain** (Foster) — monotony > 2.0 across a week flags overly uniform training.",
        "**Match-readiness ratio** (Gabbett) — training vs match load; < 0.5 = under-prepared.",
    ])

    # ── SECTION A: Matchday Window Profile (cleaned) ──
    st.markdown("### MD Window Profile (MD-5 → MD+3)")
    st.caption("Training sessions only (match sessions excluded). 'Non-MD Window' and unknown sessions are also excluded.")
    metric_tooltip("md_labels")
    if "md_label" not in df_train.columns:
        st.warning("No md_label available.")
    else:
        md_keep = ["MD-5", "MD-4", "MD-3", "MD-2", "MD-1", "MD", "MD+1", "MD+2", "MD+3"]
        df_md = df_train[df_train["md_label"].astype(str).isin(md_keep)].copy()
        # Exclude actual match sessions and match-level loads from training microcycle
        # Match sessions typically have load > 1000; training is usually 300-800
        if "total_player_load" in df_md.columns:
            _load_cap = df_md["total_player_load"].quantile(0.90)  # 90th percentile as cap
            df_md = df_md[df_md["total_player_load"] <= max(900, _load_cap)].copy()
        df_md["md_label"] = pd.Categorical(df_md["md_label"].astype(str), categories=md_keep, ordered=True)

        metric = st.selectbox(
            "Metric",
            [c for c in ["total_player_load", "player_load_per_min_est", "hsd_m", "sprint_m",
                         "total_acceleration_load", "maximum_velocity", "acwr_ewma_7_28"] if c in df_md.columns],
            format_func=pretty,
            index=0,
            key="md_metric",
        )

        _md_mean = df_md.groupby("md_label", as_index=False)[metric].mean(numeric_only=True)
        _md_count = df_md.groupby("md_label", as_index=False)[metric].count()
        _md_count.columns = ["md_label", "n_sessions"]
        md_profile = _md_mean.merge(_md_count, on="md_label")
        md_profile = md_profile.rename(columns={metric: "mean_val"})
        md_profile = md_profile.sort_values("md_label")
        fig_md = go.Figure()
        fig_md.add_trace(go.Bar(
            x=md_profile["md_label"], y=md_profile["mean_val"],
            marker_color=[
                "#ef4444" if lbl == "MD" else "#f59e0b" if lbl in ("MD-1", "MD+1") else CU_GOLD
                for lbl in md_profile["md_label"]
            ],
            text=[f"{v:.0f}" for v in md_profile["mean_val"]],
            textposition="outside",
            hovertemplate="<b>%{x}</b><br>" + pretty(metric) + ": %{y:.1f}<br>Sessions: %{customdata}<extra></extra>",
            customdata=md_profile["n_sessions"].values,
        ))
        fig_md.update_layout(
            title=f"Team Avg {pretty(metric)} by Matchday Window",
            xaxis_title="Matchday Window", yaxis_title=pretty(metric),
            height=420,
        )
        st.plotly_chart(fig_md, use_container_width=True)
        st.caption(f"Based on {int(md_profile['n_sessions'].sum())} total sessions across all matchday windows.")

        if player_sel != "(All Players)":
            p = df_md[df_md["player_id"] == player_sel]
            pprof = p.groupby("md_label", as_index=False)[metric].mean(numeric_only=True)
            st.plotly_chart(px.line(pprof.sort_values("md_label"), x="md_label", y=metric, markers=True,
                                    title=f"{player_sel} — {pretty(metric)} by Matchday Window"),
                            use_container_width=True)

    # ── SECTION B: Day-of-Week Profiling ──
    st.markdown("---")
    st.markdown("### 📆 Day-of-Week Load Profile")
    st.caption(
        "What does a typical Tuesday vs Thursday look like? "
        "Red flags appear when a session's load is **>1.5 SD above** that day's historical average."
    )
    if "total_player_load" in df_train.columns:
        _dow = df_train.copy()
        _dow["day_name"] = pd.to_datetime(_dow["date"]).dt.day_name()
        dow_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        _dow["day_name"] = pd.Categorical(_dow["day_name"], categories=dow_order, ordered=True)

        dow_metric = st.selectbox(
            "Day-of-week metric",
            [c for c in ["total_player_load", "player_load_per_min_est", "maximum_velocity",
                         "acwr_ewma_7_28", "hsd_m", "sprint_m"] if c in _dow.columns],
            index=0, key="dow_metric",
        )

        dow_stats = _dow.groupby("day_name")[dow_metric].agg(["mean", "std", "count"]).reset_index()
        dow_stats.columns = ["day_name", "mean", "std", "count"]

        col_dow1, col_dow2 = st.columns(2)
        with col_dow1:
            fig_dow = go.Figure()
            fig_dow.add_trace(go.Bar(
                x=dow_stats["day_name"], y=dow_stats["mean"],
                error_y=dict(type="data", array=dow_stats["std"].fillna(0).tolist(), visible=True),
                marker_color=CU_GOLD, name="Mean ± SD",
            ))
            fig_dow.update_layout(title=f"Typical {pretty(dow_metric)} by Day of Week",
                                  yaxis_title=pretty(dow_metric), height=380)
            st.plotly_chart(fig_dow, use_container_width=True)

        with col_dow2:
            fig_dow_box = px.box(_dow, x="day_name", y=dow_metric, points="outliers",
                                 title=f"Distribution of {pretty(dow_metric)} by Day")
            fig_dow_box.update_layout(height=380)
            st.plotly_chart(fig_dow_box, use_container_width=True)

        # Flag sessions that deviate from the day-of-week norm
        st.markdown("#### 🚩 Sessions above normal for that day (>1.5 SD)")
        _dow_merged = _dow.merge(dow_stats[["day_name", "mean", "std"]], on="day_name", how="left")
        _dow_merged["z_score"] = (_dow_merged[dow_metric] - _dow_merged["mean"]) / _dow_merged["std"].replace(0, np.nan)
        flagged = _dow_merged[_dow_merged["z_score"] > 1.5].copy()
        if flagged.empty:
            st.success("No sessions significantly above day-of-week norms.")
        else:
            show_flag = [c for c in ["player_id", "date", "day_name", dow_metric, "z_score", "md_label"] if c in flagged.columns]
            st.dataframe(
                pretty_cols(flagged[show_flag].sort_values("z_score", ascending=False).head(30)),
                use_container_width=True, hide_index=True,
            )
            st.caption(f"z-score = how many SDs above the typical {pretty(dow_metric)} for that day of week.")

    # ── SECTION C: Velocity Zone Optimization ──
    st.markdown("---")
    st.markdown("### 🏃 Velocity Zone Optimization (CU Boulder)")
    st.caption(
        "Comprehensive velocity analysis: **Catapult speed bands**, **% of personal best zones**, "
        "**match-based position-specific thresholds**, and **gap/misclassification analysis**. "
        "All speeds in MPH. Values ≥ 22 mph removed as GPS errors."
    )
    _vz_info_c1, _vz_info_c2 = st.columns(2)
    with _vz_info_c1:
        metric_tooltip("speed_bands")
    with _vz_info_c2:
        metric_tooltip("maximum_velocity")

    # Match-based position-specific thresholds from Capstone analysis (67 validated matches, 2023-2025)
    _MATCH_THRESHOLDS = {
        "DEF": {"hs": 17.7, "sprint": 18.2},
        "FWD": {"hs": 17.7, "sprint": 19.0},
        "MID": {"hs": 17.8, "sprint": 18.4},
        "GK":  {"hs": 15.0, "sprint": 15.9},
    }

    def _pos_group_vz(p: str) -> str:
        pl = str(p).lower()
        if any(k in pl for k in ["gk", "keeper", "goal"]):
            return "GK"
        if any(k in pl for k in ["def", "back", "cb", "rb", "lb", "fb"]):
            return "DEF"
        if any(k in pl for k in ["mid", "cm", "cdm", "cam", "dm", "am"]):
            return "MID"
        if any(k in pl for k in ["fwd", "forward", "striker", "wing", "att", "st", "cf", "lw", "rw"]):
            return "FWD"
        return "OTHER"

    if "maximum_velocity" in df_train.columns and "vmax_pct_pb" in df_train.columns:
        _vz = df_train[df_train["vmax_pct_pb"].notna()].copy()
        if "position_name" in _vz.columns:
            _vz["_pos_grp"] = _vz["position_name"].apply(_pos_group_vz)
        else:
            _vz["_pos_grp"] = "OTHER"

        _vz_sub = st.radio(
            "Velocity section", ["Speed Bands", "% Personal Best Zones",
                                  "Match-Based Position Thresholds", "Gap & Misclassification",
                                  "Training vs Match Exposure"],
            horizontal=True, key="vz_sub",
        )

        # ─────────────────────────────────────────────────────────
        # SUB A: Absolute speed band distribution (Data Dictionary)
        # ─────────────────────────────────────────────────────────
        if _vz_sub == "Speed Bands":
            st.markdown("#### Catapult Speed Bands (per Data Dictionary)")
            st.caption(
                "**Band 1:** 0–9.6 mph (jogging) | **Band 2:** 9.6–12 mph (running) | "
                "**Band 3:** 12–14.4 mph (high-speed running) | **Band 4:** 14.4+ mph (sprinting)"
            )
            if "speed_band" in _vz.columns:
                band_counts = _vz.groupby(["player_id", "speed_band"]).size().reset_index(name="sessions")
                fig_bands = px.bar(
                    band_counts, x="player_id", y="sessions", color="speed_band",
                    title="Speed Band Distribution per Player (mph)",
                    color_discrete_map={
                        "Band 1 (0–9.6 mph)": "#22c55e", "Band 2 (9.6–12 mph)": "#84cc16",
                        "Band 3 (12–14.4 mph)": "#f59e0b", "Band 4 (14.4+ mph)": "#ef4444",
                    },
                )
                fig_bands.update_layout(xaxis_tickangle=-45, height=420)
                st.plotly_chart(fig_bands, use_container_width=True)

                # Position-level band usage (% of total distance from Set 2 columns)
                _band_dist_cols = [c for c in [
                    "velocity_band_1_total_distance_set_2", "velocity_band_2_total_distance_set_2",
                    "velocity_band_3_total_distance_set_2", "velocity_band_4_total_distance_set_2",
                ] if c in _vz.columns]
                if len(_band_dist_cols) == 4 and _vz["_pos_grp"].nunique() > 1:
                    st.markdown("#### Velocity Band Distance by Position (meters)")
                    _bd = _vz.groupby("_pos_grp")[_band_dist_cols].mean().reset_index()
                    _bd.columns = ["Position", "Band 1 (0–9.6)", "Band 2 (9.6–12)", "Band 3 (12–14.4)", "Band 4 (14.4+)"]
                    _bd_total = _bd[["Band 1 (0–9.6)", "Band 2 (9.6–12)", "Band 3 (12–14.4)", "Band 4 (14.4+)"]].sum(axis=1).replace(0, np.nan)
                    for bc in ["Band 1 (0–9.6)", "Band 2 (9.6–12)", "Band 3 (12–14.4)", "Band 4 (14.4+)"]:
                        _bd[bc + " %"] = (100 * _bd[bc] / _bd_total).round(1)
                    _bd_pct = _bd[["Position", "Band 1 (0–9.6) %", "Band 2 (9.6–12) %", "Band 3 (12–14.4) %", "Band 4 (14.4+) %"]].copy()
                    _bd_long = _bd_pct.melt(id_vars="Position", var_name="Band", value_name="% of Distance")
                    fig_bd_pct = px.bar(
                        _bd_long, x="% of Distance", y="Position", color="Band", orientation="h",
                        title="Velocity Band Usage by Position (% of Total Distance)",
                        color_discrete_map={
                            "Band 1 (0–9.6) %": "#22c55e", "Band 2 (9.6–12) %": "#84cc16",
                            "Band 3 (12–14.4) %": "#f59e0b", "Band 4 (14.4+) %": "#ef4444",
                        },
                        text="% of Distance",
                    )
                    fig_bd_pct.update_layout(height=350, xaxis_title="% of Total Distance")
                    fig_bd_pct.update_traces(texttemplate="%{text:.1f}%", textposition="inside")
                    st.plotly_chart(fig_bd_pct, use_container_width=True)
                    st.caption(
                        "Training load is heavily skewed toward low-intensity work. "
                        "Typically ~97% of total distance is below 12 mph (Bands 1+2)."
                    )
            else:
                st.info("Speed band data not computed.")

        # ─────────────────────────────────────────────────────────
        # SUB B: % of Personal Best zones
        # ─────────────────────────────────────────────────────────
        elif _vz_sub == "% Personal Best Zones":
            st.markdown("#### % of Personal Best (Arbitrary Threshold Method)")
            with st.expander("⚠️ Limitations of the arbitrary 70%/90% approach", expanded=False):
                st.markdown("""
The arbitrary method defines High-Speed as ≥ 70% of a player's all-time max velocity,
and Sprint as ≥ 90%. While this *appears* individualized, **every player can reach 70% and
90% of their own maximum** — meaning this method classifies 100% of athletes as high-speed
and sprint capable, regardless of their actual speed relative to positional peers.

This produces a **~28% misclassification rate for high-speed** and **~50% for sprint** work
compared to match-based position-specific benchmarks.
                """)
            col_vz1, col_vz2 = st.columns(2)
            with col_vz1:
                fig_vz_hist = px.histogram(
                    _vz, x="vmax_pct_pb", nbins=50,
                    title="Session Vmax as % of Personal Best",
                    color_discrete_sequence=["#8b5cf6"],
                )
                fig_vz_hist.add_vline(x=70, line_dash="dash", line_color="#f59e0b", annotation_text="High Speed (70%)")
                fig_vz_hist.add_vline(x=90, line_dash="dash", line_color="#ef4444", annotation_text="Sprint (90%)")
                fig_vz_hist.update_layout(xaxis_title="% of Personal Best Velocity", height=400)
                st.plotly_chart(fig_vz_hist, use_container_width=True)

            with col_vz2:
                player_thresholds = _vz.groupby("player_id")["vmax_pct_pb"].agg(
                    P25="median", P70=lambda x: x.quantile(0.70),
                    P90=lambda x: x.quantile(0.90), P95=lambda x: x.quantile(0.95),
                ).reset_index().round(1)
                player_thresholds.columns = ["Player", "Median %PB", "P70 (High Speed)", "P90 (Sprint)", "P95 (Max Effort)"]
                st.markdown("**Per-player arbitrary thresholds**")
                st.dataframe(player_thresholds.sort_values("P90 (Sprint)", ascending=False),
                             use_container_width=True, hide_index=True)

            if "vmax_zone" in _vz.columns:
                zone_counts = _vz.groupby(["player_id", "vmax_zone"]).size().reset_index(name="sessions")
                fig_zone = px.bar(zone_counts, x="player_id", y="sessions", color="vmax_zone",
                                  title="Velocity Zone Distribution per Player",
                                  color_discrete_map={
                                      "Z1 (<70%)": "#22c55e", "Z2 (70–80%)": "#84cc16",
                                      "Z3 (80–90%)": "#f59e0b", "Z4 (90–95%)": "#f97316",
                                      "Z5 (95–100%)": "#ef4444",
                                  })
                fig_zone.update_layout(xaxis_tickangle=-45, height=420)
                st.plotly_chart(fig_zone, use_container_width=True)

        # ─────────────────────────────────────────────────────────
        # SUB C: Match-Based Position-Specific Thresholds
        # ─────────────────────────────────────────────────────────
        elif _vz_sub == "Match-Based Position Thresholds":
            st.markdown("#### Match-Based Position-Specific Velocity Thresholds")
            st.caption(
                "Derived from 67 validated matches (2023–2025). Uses the **70th percentile** "
                "of match max velocities for High-Speed and the **90th percentile** for Sprint, per position."
            )
            with st.expander("ℹ️ Why match-based thresholds?", expanded=False):
                st.markdown("""
**1. Ecological Validity:** Training sessions inherently limit max velocity (small-sided games,
drills, recovery). Match data captures true competitive capability.

**2. Competition-Focused Benchmarking:** Training prepares athletes for competition — thresholds
should reflect real game demands. The 70th and 90th match percentiles identify the top 30%
(high-speed) and top 10% (sprint) of competitive performances per position.

**3. Critical for Low-Intensity Programs:** ~97% of training distance occurs below 12 mph.
With so little game-speed work, **correct classification of what "match-intensity" means**
is critical for the limited high-speed exposure that does occur.
                """)

            # Threshold table
            _thresh_df = pd.DataFrame([
                {"Position": k, "High-Speed (mph)": v["hs"], "Sprint (mph)": v["sprint"]}
                for k, v in _MATCH_THRESHOLDS.items()
            ])
            st.dataframe(_thresh_df, use_container_width=True, hide_index=True)

            # Player classification against match thresholds
            _player_max = _vz.groupby(["player_id", "_pos_grp"]).agg(
                all_time_max=("maximum_velocity", "max"),
                n_sessions=("maximum_velocity", "count"),
            ).reset_index()
            _player_max = _player_max[_player_max["_pos_grp"].isin(_MATCH_THRESHOLDS.keys())].copy()

            _player_max["hs_threshold"] = _player_max["_pos_grp"].map(lambda g: _MATCH_THRESHOLDS[g]["hs"])
            _player_max["sprint_threshold"] = _player_max["_pos_grp"].map(lambda g: _MATCH_THRESHOLDS[g]["sprint"])
            _player_max["arbitrary_hs"] = _player_max["all_time_max"] * 0.70
            _player_max["arbitrary_sprint"] = _player_max["all_time_max"] * 0.90

            _player_max["meets_hs_match"] = _player_max["all_time_max"] >= _player_max["hs_threshold"]
            _player_max["meets_sprint_match"] = _player_max["all_time_max"] >= _player_max["sprint_threshold"]
            _player_max["meets_hs_arb"] = True  # by definition, everyone hits 70% of their own max
            _player_max["meets_sprint_arb"] = True  # by definition, everyone hits 90% of their own max

            st.markdown("#### Player Classification: Match-Based vs Arbitrary")
            col_cls1, col_cls2 = st.columns(2)
            with col_cls1:
                n_hs = int(_player_max["meets_hs_match"].sum())
                n_tot = len(_player_max)
                st.metric("Meet High-Speed (match-based)", f"{n_hs} / {n_tot}",
                          delta=f"{100 * n_hs / max(n_tot, 1):.0f}%")
            with col_cls2:
                n_sp = int(_player_max["meets_sprint_match"].sum())
                st.metric("Meet Sprint (match-based)", f"{n_sp} / {n_tot}",
                          delta=f"{100 * n_sp / max(n_tot, 1):.0f}%")

            # Grouped bar chart: one group per position, player max vs thresholds
            _pos_colors = {"DEF": "#6366f1", "MID": "#22c55e", "FWD": "#ef4444", "GK": "#f59e0b"}
            for grp in ["DEF", "MID", "FWD", "GK"]:
                gdf = _player_max[_player_max["_pos_grp"] == grp].sort_values("all_time_max", ascending=True).copy()
                if gdf.empty:
                    continue
                hs_val = _MATCH_THRESHOLDS[grp]["hs"]
                sp_val = _MATCH_THRESHOLDS[grp]["sprint"]

                fig_grp = go.Figure()
                bar_colors = [
                    "#22c55e" if v >= sp_val else "#f59e0b" if v >= hs_val else "#ef4444"
                    for v in gdf["all_time_max"]
                ]
                fig_grp.add_trace(go.Bar(
                    y=gdf["player_id"], x=gdf["all_time_max"], orientation="h",
                    marker_color=bar_colors, name="All-Time Max",
                    text=[f"{v:.1f}" for v in gdf["all_time_max"]],
                    textposition="outside",
                    hovertemplate="%{y}<br>Max: %{x:.1f} mph<extra></extra>",
                ))
                fig_grp.add_vline(x=hs_val, line_dash="dash", line_color="#f59e0b", line_width=2,
                                  annotation_text=f"HS ({hs_val})", annotation_position="top right")
                fig_grp.add_vline(x=sp_val, line_dash="dash", line_color="#ef4444", line_width=2,
                                  annotation_text=f"Sprint ({sp_val})", annotation_position="top right")
                fig_grp.update_layout(
                    title=f"{grp} — Player Max Velocity vs Thresholds",
                    xaxis_title="Max Velocity (mph)",
                    height=max(280, 32 * len(gdf) + 80),
                    margin=dict(l=10, r=40, t=50, b=40),
                    showlegend=False,
                )
                st.plotly_chart(fig_grp, use_container_width=True)
            st.caption("🟢 Green = meets sprint threshold · 🟡 Yellow = meets high-speed only · 🔴 Red = below high-speed threshold")

            # Detailed table
            _show_cls = _player_max[["player_id", "_pos_grp", "all_time_max", "hs_threshold",
                                      "sprint_threshold", "meets_hs_match", "meets_sprint_match"]].copy()
            _show_cls.columns = ["Player", "Position", "All-Time Max (mph)", "HS Threshold",
                                  "Sprint Threshold", "Meets HS?", "Meets Sprint?"]
            _show_cls = _show_cls.sort_values(["Position", "All-Time Max (mph)"], ascending=[True, False])
            st.dataframe(_show_cls, use_container_width=True, hide_index=True)

        # ─────────────────────────────────────────────────────────
        # SUB D: Gap & Misclassification Analysis
        # ─────────────────────────────────────────────────────────
        elif _vz_sub == "Gap & Misclassification":
            st.markdown("#### Arbitrary vs Match-Based: Gap Analysis")
            st.caption(
                "Compares the arbitrary 70%/90%-of-PB thresholds to match-based position-specific "
                "thresholds. Shows how many players would be **misclassified** by the arbitrary method."
            )

            _player_max2 = _vz.groupby(["player_id", "_pos_grp"]).agg(
                all_time_max=("maximum_velocity", "max"),
            ).reset_index()
            _player_max2 = _player_max2[_player_max2["_pos_grp"].isin(_MATCH_THRESHOLDS.keys())].copy()
            if not _player_max2.empty:
                _player_max2["hs_match"] = _player_max2["_pos_grp"].map(lambda g: _MATCH_THRESHOLDS[g]["hs"])
                _player_max2["sp_match"] = _player_max2["_pos_grp"].map(lambda g: _MATCH_THRESHOLDS[g]["sprint"])
                _player_max2["arb_hs"] = _player_max2["all_time_max"] * 0.70
                _player_max2["arb_sp"] = _player_max2["all_time_max"] * 0.90

                # Gap in mph
                _player_max2["hs_gap_mph"] = _player_max2["hs_match"] - _player_max2["arb_hs"]
                _player_max2["sp_gap_mph"] = _player_max2["sp_match"] - _player_max2["arb_sp"]

                # Misclassification: arbitrary says "yes" but match-based says "no"
                _player_max2["hs_misclass"] = (_player_max2["all_time_max"] < _player_max2["hs_match"])
                _player_max2["sp_misclass"] = (_player_max2["all_time_max"] < _player_max2["sp_match"])

                col_g1, col_g2 = st.columns(2)
                with col_g1:
                    # Gap by position
                    _gap_pos = _player_max2.groupby("_pos_grp").agg(
                        n_players=("player_id", "count"),
                        hs_gap_avg=("hs_gap_mph", "mean"),
                        sp_gap_avg=("sp_gap_mph", "mean"),
                    ).reset_index()
                    _gap_pos.columns = ["Position", "Players", "Avg HS Gap (mph)", "Avg Sprint Gap (mph)"]
                    _gap_long = _gap_pos.melt(id_vars=["Position", "Players"],
                                              var_name="Metric", value_name="Gap (mph)")
                    fig_gap = px.bar(
                        _gap_long, x="Position", y="Gap (mph)", color="Metric", barmode="group",
                        title="How Much Higher Are Match-Based Thresholds? (mph)",
                        color_discrete_sequence=["#f59e0b", "#ef4444"],
                    )
                    fig_gap.update_layout(height=380)
                    st.plotly_chart(fig_gap, use_container_width=True)

                with col_g2:
                    # Misclassification rates by position
                    _mis_pos = _player_max2.groupby("_pos_grp").agg(
                        n_players=("player_id", "count"),
                        hs_misclass=("hs_misclass", "sum"),
                        sp_misclass=("sp_misclass", "sum"),
                    ).reset_index()
                    _mis_pos["HS Error Rate (%)"] = (100 * _mis_pos["hs_misclass"] / _mis_pos["n_players"]).round(1)
                    _mis_pos["Sprint Error Rate (%)"] = (100 * _mis_pos["sp_misclass"] / _mis_pos["n_players"]).round(1)
                    _mis_pos.columns = ["Position", "Players", "HS Misclass", "Sprint Misclass",
                                        "HS Error Rate (%)", "Sprint Error Rate (%)"]
                    _mis_long = _mis_pos.melt(id_vars=["Position", "Players"],
                                              value_vars=["HS Error Rate (%)", "Sprint Error Rate (%)"],
                                              var_name="Type", value_name="Error Rate (%)")
                    fig_mis = px.bar(
                        _mis_long, x="Position", y="Error Rate (%)", color="Type", barmode="group",
                        title="Misclassification Rate: Arbitrary vs Match-Based",
                        color_discrete_sequence=["#f59e0b", "#ef4444"],
                        text="Error Rate (%)",
                    )
                    fig_mis.update_traces(texttemplate="%{text:.0f}%", textposition="outside")
                    fig_mis.update_layout(height=380)
                    st.plotly_chart(fig_mis, use_container_width=True)

                # Overall misclassification summary
                total_p = len(_player_max2)
                hs_mis = int(_player_max2["hs_misclass"].sum())
                sp_mis = int(_player_max2["sp_misclass"].sum())
                st.warning(
                    f"**Arbitrary method misclassifies {hs_mis}/{total_p} ({100*hs_mis/max(total_p,1):.0f}%) "
                    f"for high-speed** and **{sp_mis}/{total_p} ({100*sp_mis/max(total_p,1):.0f}%) for sprint** "
                    f"compared to match-based position-specific benchmarks."
                )

                # Threshold comparison table
                _comp_df = _player_max2.groupby("_pos_grp").agg(
                    n=("player_id", "count"),
                    arb_hs_avg=("arb_hs", "mean"),
                    match_hs=("hs_match", "first"),
                    arb_sp_avg=("arb_sp", "mean"),
                    match_sp=("sp_match", "first"),
                ).reset_index()
                _comp_df.columns = ["Position", "Players", "Arbitrary HS (avg)", "Match HS",
                                     "Arbitrary Sprint (avg)", "Match Sprint"]
                st.markdown("#### Threshold Comparison (mph)")
                st.dataframe(_comp_df.round(1), use_container_width=True, hide_index=True)

        # ─────────────────────────────────────────────────────────
        # SUB E: Training vs Match Velocity Exposure
        # ─────────────────────────────────────────────────────────
        elif _vz_sub == "Training vs Match Exposure":
            st.markdown("#### Training vs Match-Day Velocity Exposure")
            st.caption(
                "Shows how much (or how little) high-speed and sprint work athletes get "
                "in training compared to match days. ~97% of training distance is below 12 mph."
            )

            if "md_label" in _vz.columns:
                _vz_exp = _vz.copy()
                _vz_exp["_session_type"] = np.where(
                    _vz_exp["md_label"].astype(str) == "MD", "Match Day", "Training"
                )

                col_e1, col_e2 = st.columns(2)
                with col_e1:
                    fig_exp = px.histogram(
                        _vz_exp, x="maximum_velocity", color="_session_type",
                        nbins=40, barmode="overlay", opacity=0.7,
                        title="Max Velocity Distribution: Training vs Match",
                        color_discrete_map={"Training": CU_GOLD, "Match Day": "#ef4444"},
                    )
                    fig_exp.update_layout(xaxis_title="Max Velocity (mph)", height=400)
                    st.plotly_chart(fig_exp, use_container_width=True)

                with col_e2:
                    _exp_stats = _vz_exp.groupby(["_pos_grp", "_session_type"]).agg(
                        mean_vmax=("maximum_velocity", "mean"),
                        p90_vmax=("maximum_velocity", lambda x: x.quantile(0.90)),
                        sessions=("maximum_velocity", "count"),
                    ).reset_index()
                    if not _exp_stats.empty:
                        fig_exp_pos = px.bar(
                            _exp_stats, x="_pos_grp", y="mean_vmax", color="_session_type",
                            barmode="group",
                            title="Mean Max Velocity by Position: Training vs Match",
                            color_discrete_map={"Training": CU_GOLD, "Match Day": "#ef4444"},
                            text=_exp_stats["mean_vmax"].round(1),
                        )
                        fig_exp_pos.update_traces(textposition="outside")
                        fig_exp_pos.update_layout(xaxis_title="Position", yaxis_title="Avg Max Velocity (mph)", height=400)
                        st.plotly_chart(fig_exp_pos, use_container_width=True)

                # High-intensity exposure analysis
                _band_dist_cols = [c for c in [
                    "velocity_band_1_total_distance_set_2", "velocity_band_2_total_distance_set_2",
                    "velocity_band_3_total_distance_set_2", "velocity_band_4_total_distance_set_2",
                ] if c in _vz_exp.columns]
                if len(_band_dist_cols) == 4:
                    _vz_exp["total_band_dist"] = _vz_exp[_band_dist_cols].sum(axis=1)
                    _vz_exp["hi_speed_dist"] = _vz_exp[_band_dist_cols[2]] + _vz_exp[_band_dist_cols[3]]
                    _vz_exp["hi_speed_pct"] = 100 * _vz_exp["hi_speed_dist"] / _vz_exp["total_band_dist"].replace(0, np.nan)

                    _hi_by_type = _vz_exp.groupby("_session_type").agg(
                        avg_hi_pct=("hi_speed_pct", "mean"),
                        avg_hi_m=("hi_speed_dist", "mean"),
                    ).reset_index()

                    st.markdown("#### High-Intensity Distance (≥12 mph) — Training vs Match")
                    col_hi1, col_hi2 = st.columns(2)
                    with col_hi1:
                        for _, r in _hi_by_type.iterrows():
                            st.metric(
                                f"{r['_session_type']} — Avg high-speed distance",
                                f"{r['avg_hi_m']:.0f} m/session",
                                delta=f"{r['avg_hi_pct']:.1f}% of total",
                            )
                    with col_hi2:
                        _hi_pos = _vz_exp.groupby(["_pos_grp", "_session_type"]).agg(
                            avg_hi_pct=("hi_speed_pct", "mean"),
                        ).reset_index()
                        if not _hi_pos.empty:
                            fig_hi = px.bar(
                                _hi_pos, x="_pos_grp", y="avg_hi_pct", color="_session_type",
                                barmode="group",
                                title="High-Speed Distance % by Position",
                                color_discrete_map={"Training": CU_GOLD, "Match Day": "#ef4444"},
                                text=_hi_pos["avg_hi_pct"].round(1),
                            )
                            fig_hi.update_traces(texttemplate="%{text:.1f}%", textposition="outside")
                            fig_hi.update_layout(yaxis_title="% of Total Distance ≥12 mph", height=400)
                            st.plotly_chart(fig_hi, use_container_width=True)

                st.info(
                    "**Key finding:** High-intensity training exposure is very low across all positions "
                    "(typically 0.9–2.7% of total distance > 12 mph). With ~97% of training below 12 mph, "
                    "correct velocity zone classification is critical for monitoring the limited "
                    "game-speed work that does occur."
                )
            else:
                st.info("Match-day labels not available for training vs match comparison.")

    else:
        st.info("Velocity zone data not available.")

    # ── SECTION D: Match Readiness Model ──
    st.markdown("---")
    st.markdown("### 🎯 Match Readiness — Are We Preparing for Game Day?")
    st.caption(
        "Based on Gabbett's training–competition match principle (BJSM, 2016): "
        "athletes must be exposed to training loads that approximate match demands. "
        "Under-exposure (ratio < 0.5) leaves athletes physically unprepared for peak game scenarios."
    )
    metric_tooltip("match_readiness_ratio")
    if "md_label" in df_train.columns and "total_player_load" in df_train.columns:
        _mr = df_train.copy()
        _mr["is_matchday"] = _mr["md_label"].astype(str) == "MD"
        _mr["is_training"] = ~_mr["is_matchday"] & ~_mr["md_label"].astype(str).isin(["UNK", "Non-MD Window"])

        mr_metrics = [c for c in ["total_player_load", "player_load_per_min_est", "maximum_velocity",
                                   "hsd_m", "sprint_m", "total_acceleration_load"] if c in _mr.columns]
        mr_metric = st.selectbox("Readiness metric", mr_metrics, index=0, key="mr_metric")

        md_vals = _mr[_mr["is_matchday"]][mr_metric].dropna()
        tr_vals = _mr[_mr["is_training"]][mr_metric].dropna()

        if len(md_vals) > 5 and len(tr_vals) > 5:
            col_mr1, col_mr2 = st.columns(2)

            with col_mr1:
                fig_mr_comp = go.Figure()
                fig_mr_comp.add_trace(go.Histogram(x=tr_vals, name="Training days", marker_color=CU_GOLD, opacity=0.7))
                fig_mr_comp.add_trace(go.Histogram(x=md_vals, name="Match days", marker_color="#ef4444", opacity=0.7))
                fig_mr_comp.update_layout(
                    barmode="overlay", title=f"{pretty(mr_metric)}: Training vs Match Day",
                    xaxis_title=pretty(mr_metric), height=400,
                )
                st.plotly_chart(fig_mr_comp, use_container_width=True)

            with col_mr2:
                md_mean = float(md_vals.mean())
                md_p75 = float(md_vals.quantile(0.75))
                md_p90 = float(md_vals.quantile(0.90))
                tr_mean = float(tr_vals.mean())
                tr_p90 = float(tr_vals.quantile(0.90))

                st.metric("Match Day avg", f"{md_mean:.1f}")
                st.metric("Match Day P90 (worst-case 10%)", f"{md_p90:.1f}")
                st.metric("Training avg", f"{tr_mean:.1f}")
                st.metric("Training P90", f"{tr_p90:.1f}")
                gap = md_p90 - tr_p90
                if gap > 0:
                    st.error(f"⚠️ Training P90 is **{gap:.1f} below** match-day P90 — athletes may not be prepared for peak game demands.")
                else:
                    st.success(f"✅ Training P90 **meets or exceeds** match-day P90 by {abs(gap):.1f}.")

            # Per-player readiness ratio
            st.markdown("#### Per-Player: Training vs Game-Day Readiness Ratio")
            st.caption(
                "Ratio = player's avg training load ÷ avg match-day load (Gabbett, 2016). "
                "Only ratios below 0.5 are flagged — these athletes have trained at less than "
                "half of match intensity and may be under-prepared for competition demands."
            )
            _p_md = _mr[_mr["is_matchday"]].groupby("player_id")[mr_metric].mean().rename("md_avg")
            _p_tr = _mr[_mr["is_training"]].groupby("player_id")[mr_metric].mean().rename("tr_avg")
            _p_ratio = pd.concat([_p_tr, _p_md], axis=1).dropna()
            _p_ratio["readiness_ratio"] = (_p_ratio["tr_avg"] / _p_ratio["md_avg"].replace(0, np.nan)).round(2)
            _p_ratio["flag"] = np.where(_p_ratio["readiness_ratio"] < 0.5, "🔴 Under-prepared", "🟢 Good")
            _p_ratio = _p_ratio.reset_index().sort_values("readiness_ratio")
            st.dataframe(pretty_cols(_p_ratio), use_container_width=True, hide_index=True)

            fig_ratio = px.bar(
                _p_ratio.sort_values("readiness_ratio"), x="player_id", y="readiness_ratio",
                color="flag",
                color_discrete_map={"🔴 Under-prepared": "#ef4444", "🟢 Good": "#22c55e"},
                title="Training / Match-Day Load Ratio per Player",
            )
            fig_ratio.add_hline(y=0.5, line_dash="dash", line_color="#ef4444", annotation_text="Under-prepared (<0.5)")
            fig_ratio.update_layout(height=420, xaxis_tickangle=-45, showlegend=True)
            st.plotly_chart(fig_ratio, use_container_width=True)
        else:
            st.info("Need at least 5 match-day and 5 training sessions to compute readiness comparison.")
    else:
        st.info("Match readiness requires md_label and load columns.")

    # ── SECTION E: Weekly Monotony & Strain ──
    st.markdown("---")
    st.markdown("### Weekly Monotony & Strain (Foster, 1998)")
    st.caption(
        "**Monotony** = mean daily load ÷ SD of daily loads (Foster, 1998). Values > 2.0 indicate "
        "insufficient training variation. **Strain** = weekly load × monotony. High strain weeks "
        "are linked to illness/injury in the following 7–14 days. Weeks with < 3 sessions excluded. "
        "Monotony capped at 10.0 to prevent blow-ups from near-zero SD."
    )
    metric_tooltip("monotony_strain")

    if weekly_team.empty:
        st.info("Not enough data to compute weekly monotony/strain.")
    else:
        wt = weekly_team.copy()
        wt["week_label"] = wt["year"].astype(str) + "-W" + wt["week"].astype(str).str.zfill(2)
        wt = wt.sort_values(["year", "week"])

        # Flag dangerous weeks
        mono_thresh = 2.0
        strain_p90 = float(wt["strain"].quantile(0.90)) if wt["strain"].notna().any() else 999
        wt["mono_flag"] = wt["monotony"].fillna(0) >= mono_thresh
        wt["strain_flag"] = wt["strain"].fillna(0) >= strain_p90

        n_mono_high = int(wt["mono_flag"].sum())
        n_strain_high = int(wt["strain_flag"].sum())

        mc1, mc2, mc3 = st.columns(3)
        mc1.metric("Total weeks", len(wt))
        mc2.metric("🔴 High monotony weeks (≥2.0)", n_mono_high)
        mc3.metric("🔴 High strain weeks (≥P90)", n_strain_high)

        col_mono, col_strain = st.columns(2)
        with col_mono:
            fig_m = go.Figure()
            fig_m.add_trace(go.Bar(
                x=wt["week_label"], y=wt["monotony"],
                marker_color=[
                    "#ef4444" if f else CU_GOLD for f in wt["mono_flag"]
                ],
                name="Monotony",
            ))
            fig_m.add_hline(y=mono_thresh, line_dash="dash", line_color="#f59e0b",
                            annotation_text="Monotony = 2.0 (danger zone)")
            fig_m.update_layout(title="Weekly Monotony", height=380,
                                margin=dict(l=40, r=10, t=50, b=60),
                                xaxis_tickangle=-45)
            st.plotly_chart(fig_m, use_container_width=True)

        with col_strain:
            fig_s = go.Figure()
            fig_s.add_trace(go.Bar(
                x=wt["week_label"], y=wt["strain"],
                marker_color=[
                    "#ef4444" if f else "#22c55e" for f in wt["strain_flag"]
                ],
                name="Strain",
            ))
            fig_s.add_hline(y=strain_p90, line_dash="dash", line_color="#f59e0b",
                            annotation_text=f"P90 strain = {strain_p90:,.0f}")
            fig_s.update_layout(title="Weekly Strain", height=380,
                                margin=dict(l=40, r=10, t=50, b=60),
                                xaxis_tickangle=-45)
            st.plotly_chart(fig_s, use_container_width=True)

        # Weekly load trend
        fig_wl = go.Figure()
        fig_wl.add_trace(go.Scatter(
            x=wt["week_label"], y=wt["week_load"], mode="lines+markers",
            name="Weekly load", line=dict(color=CU_GOLD, width=3),
            fill="tozeroy", fillcolor="rgba(207,184,124,0.08)",
        ))
        fig_wl.update_layout(title="Total Weekly Load Over Time", height=340,
                             margin=dict(l=40, r=10, t=50, b=60),
                             xaxis_tickangle=-45)
        st.plotly_chart(fig_wl, use_container_width=True)

        with st.expander("📋 Raw weekly data table"):
            show_wt = wt[["year", "week", "week_load", "mean_daily", "sd_daily",
                          "monotony", "strain", "days", "mono_flag", "strain_flag"]].copy()
            st.dataframe(pretty_cols(show_wt.sort_values(["year", "week"], ascending=False)), use_container_width=True, hide_index=True)
            st.download_button("📥 Download weekly data", show_wt.to_csv(index=False).encode(),
                               "weekly_monotony_strain.csv", "text/csv")

# ==============================================================================
# TAB: OPPONENT COMPARISON
# ==============================================================================
with tab_vs:
    st.subheader("🆚 Opponent Comparison (Training + Game Output)")
    coach_guide("Opponent", [
        "Compare **our physical output** against each opponent: load, high-speed, sprint distance.",
        "Useful for scouting: *did we match their intensity? out-sprint them?*",
        "Pair with Microcycle so you can plan the lead-in week to meet what they’ll bring.",
    ])
    with st.expander("ℹ️ How is **Opponent Comparison** calculated?", expanded=False):
        st.markdown(
            "Training metrics (load, HSD, sprint, acceleration, ACWR) are grouped by match date, "
            "then averaged per opponent. This shows how the team physically performed in training "
            "sessions on match days against each opponent.\n\n"
            "**Example:** If avg team load vs Baylor = 12,000 and vs TCU = 8,500, the team was "
            "working significantly harder in training around the Baylor match. "
            "Game stats (goals, assists, shots) come from per-match CSV files."
        )

    tm = team_match.copy()
    _opp_metric_labels = {
        "team_load": "Team Player Load", "team_hsd": "Team HSD (m)",
        "team_sprint": "Team Sprint (m)", "team_accel": "Team Accel Load",
        "mean_acwr": "Mean ACWR",
    }
    metric_tm = st.selectbox(
        "Training metric (match dates)",
        list(_opp_metric_labels.keys()),
        format_func=lambda x: _opp_metric_labels.get(x, x),
        key="opp_tm_metric",
    )
    opp_avg = tm.groupby("Opponent_clean", as_index=False)[metric_tm].mean()
    st.plotly_chart(px.bar(opp_avg.sort_values(metric_tm, ascending=False), x="Opponent_clean", y=metric_tm,
                           title=f"Average {_opp_metric_labels.get(metric_tm, metric_tm)} vs Opponent (match dates)",
                           labels={"Opponent_clean": "Opponent", metric_tm: _opp_metric_labels.get(metric_tm, metric_tm)}),
                    use_container_width=True)

    if game_all is not None and len(game_all) > 0:
        st.markdown("---")
        st.markdown("### Game output vs Opponent (Team totals)")

        gy = st.selectbox("Year (game output)", sorted(game_all["year"].dropna().unique().tolist()), key="opp_gs_year")
        g = game_all[game_all["year"] == gy].copy()

        grp_cols = ["match_date", "Opponent_clean", "venue", "result"]
        avail = [c for c in ["G", "A", "Sh", "SOG", "GA", "Saves"] if c in g.columns]
        gt = g.groupby(grp_cols, as_index=False)[avail].sum(numeric_only=True)

        metric_g = st.selectbox("Game metric", avail, key="opp_g_metric")
        opp_g = gt.groupby("Opponent_clean", as_index=False)[metric_g].mean()
        st.plotly_chart(px.bar(opp_g.sort_values(metric_g, ascending=False), x="Opponent_clean", y=metric_g,
                               title=f"Avg {metric_g} vs Opponent (Game totals)"),
                        use_container_width=True)

# ==============================================================================
# TAB: HEALTH & WELLNESS (merged: Injury + Return-to-Play + Wellness)
# ==============================================================================
with tab_inj:
    st.subheader("🏥 Health & Wellness Center")
    coach_guide("Health & Wellness", [
        "**Injury status** comes straight from the athletic-trainer logs (Available / Out / Limited / As Tolerated).",
        "**Wellness = Mental + Physical + Sleep + Soreness** (each 1–5, total 4–20). Below 12/20 = check-in.",
        "Use this with **Coach ML** — staff always has the final call on medical decisions.",
    ])
    _hw_info_c1, _hw_info_c2 = st.columns(2)
    with _hw_info_c1:
        metric_tooltip("injury_status")
    with _hw_info_c2:
        metric_tooltip("wellness")
    _hw_section = st.radio("Section", ["Injury Status", "Return-to-Play", "Wellness"], horizontal=True, key="hw_section")

    if _hw_section == "Injury Status":
      if injuries_all.empty:
        st.warning("No injury data loaded.")
      else:
        _inj_all = injuries_all[
            injuries_all["injury_status"].notna()
            & (injuries_all["injury_status"].str.strip() != "")
            & (injuries_all["injury_status"].str.strip().str.lower() != "nan")
        ].copy()
        st.markdown(f"**{len(_inj_all)} injury records** across {_inj_all['player_id'].nunique()} players.")

        st.markdown("### Current Injury Status (latest per player)")
        latest_inj = _inj_all.sort_values("status_start").groupby("player_id").last().reset_index()
        status_order = {"Out": 0, "Limited": 1, "As Tolerated": 2, "Full Go": 3}
        latest_inj["_sort"] = latest_inj["injury_status"].map(status_order).fillna(4)
        show_inj = [c for c in ["player_id", "injury_status", "injury_date", "status_start", "days_in_status", "year"] if c in latest_inj.columns]
        _inj_display = latest_inj.sort_values("_sort")[show_inj].copy()
        for _dc in ["injury_date", "status_start"]:
            if _dc in _inj_display.columns:
                _inj_display[_dc] = pd.to_datetime(_inj_display[_dc], errors="coerce").dt.strftime("%m/%d/%Y")
        st.dataframe(pretty_cols(_inj_display), use_container_width=True)

        st.markdown("### Status Distribution (all records)")
        status_counts = _inj_all["injury_status"].value_counts().reset_index()
        status_counts.columns = ["status", "count"]
        st.plotly_chart(px.bar(status_counts, x="status", y="count", color="status",
                               color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e"},
                               title="Injury status distribution (all records)"), use_container_width=True)

        if "injury_date" in _inj_all.columns:
            st.markdown("### Injury Timeline")
            timeline = _inj_all.dropna(subset=["injury_date"]).copy()
            timeline["injury_date"] = pd.to_datetime(timeline["injury_date"], errors="coerce")
            timeline = timeline.dropna(subset=["injury_date"])
            if not timeline.empty:
                timeline["month"] = timeline["injury_date"].dt.to_period("M").astype(str)
                _inj_monthly = timeline.groupby(["month", "injury_status"]).size().reset_index(name="count")
                fig_inj_tl = px.bar(
                    _inj_monthly, x="count", y="month", color="injury_status", orientation="h",
                    color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e"},
                    title="Injury Events per Month by Status",
                    labels={"count": "Number of Events", "month": "Month", "injury_status": "Status"},
                    barmode="group",
                )
                fig_inj_tl.update_layout(
                    height=max(350, len(_inj_monthly["month"].unique()) * 30 + 100),
                    yaxis=dict(categoryorder="category ascending"),
                )
                st.plotly_chart(fig_inj_tl, use_container_width=True)

        if "days_in_status" in _inj_all.columns:
            st.markdown("### Days in Status — how long do players stay in each status?")
            dis = _inj_all[_inj_all["days_in_status"].notna() & (_inj_all["days_in_status"] > 0)]
            if not dis.empty:
                # Focus on the part coaches care about most: 0–60 days
                dis_short = dis[dis["days_in_status"] <= 60].copy()
                c_left, c_right = st.columns(2)
                with c_left:
                    fig_hist = px.histogram(
                        dis_short,
                        x="days_in_status",
                        color="injury_status",
                        nbins=30,
                        color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e"},
                        title="0–60 days in status (most cases)",
                    )
                    st.plotly_chart(fig_hist, use_container_width=True)
                with c_right:
                    fig_box = px.box(
                        dis,
                        x="injury_status",
                        y="days_in_status",
                        points="outliers",
                        color="injury_status",
                        color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e"},
                        title="Days in status by category (with outliers)",
                    )
                    fig_box.update_yaxes(type="log")
                    st.plotly_chart(fig_box, use_container_width=True)

        if "injury_status_actual" in df_train.columns and "total_player_load" in df_train.columns:
            st.markdown("### How does training differ by status?")
            st.caption("These show how much load players carry in each injury category — useful for spotting if 'Limited' players are still overloading.")

            _inj_rename = {"Available": "Available", "Out": "Out", "Limited": "Limited", "As Tolerated": "Modified"}
            _inj_colors = {"Available": "#22c55e", "Out": "#ef4444", "Limited": CU_GOLD, "Modified": CU_LIGHT_GRAY}

            _inj_box = df_train[["injury_status_actual", "total_player_load", "maximum_velocity"]].copy() if "maximum_velocity" in df_train.columns else df_train[["injury_status_actual", "total_player_load"]].copy()
            _inj_box = _inj_box[~_inj_box["injury_status_actual"].astype(str).str.strip().str.lower().isin(["nan", ""])].copy()
            _inj_box["Status"] = _inj_box["injury_status_actual"].replace(_inj_rename)

            col_ib1, col_ib2 = st.columns(2)
            with col_ib1:
                fig_load_box = px.violin(
                    _inj_box, x="Status", y="total_player_load", color="Status",
                    color_discrete_map=_inj_colors, box=True, points="outliers",
                    title="Training Load by Player Status",
                )
                fig_load_box.update_yaxes(title="Total Player Load")
                st.plotly_chart(fig_load_box, use_container_width=True)

            with col_ib2:
                if "maximum_velocity" in _inj_box.columns:
                    vel_clean = _inj_box[_inj_box["maximum_velocity"].between(0, 22.0)].copy()
                    fig_vel_box = px.violin(
                        vel_clean, x="Status", y="maximum_velocity", color="Status",
                        color_discrete_map=_inj_colors, box=True, points="outliers",
                        title="Max Velocity by Player Status",
                    )
                    fig_vel_box.update_yaxes(title="Max Velocity (mph)")
                    st.plotly_chart(fig_vel_box, use_container_width=True)

        st.download_button("Download Injury Data", data=_inj_all.to_csv(index=False).encode(),
                           file_name="injuries_all.csv", mime="text/csv")

    elif _hw_section == "Return-to-Play":
      st.markdown("#### 🔄 Return-to-Play Tracker")
      st.caption(
        "Tracks each player's journey through injury statuses: "
        "**Out → Limited → As Tolerated → Full Go / Available**. "
        "Helps coaches see who is progressing and who is stuck."
      )

      if injuries_all.empty:
        st.warning("No injury data loaded.")
      else:
        rtp_status_order = {"Out": 0, "Limited": 1, "As Tolerated": 2, "Full Go": 3, "Available": 4}
        _rtp_all = injuries_all[
            injuries_all["injury_status"].notna()
            & (injuries_all["injury_status"].str.strip() != "")
            & (injuries_all["injury_status"].str.strip().str.lower() != "nan")
        ].copy()
        rtp = _rtp_all.copy()
        if "status_start" in rtp.columns:
            rtp["status_start"] = pd.to_datetime(rtp["status_start"], errors="coerce")
            rtp = rtp.dropna(subset=["status_start"]).sort_values(["player_id", "status_start"])
        elif "report_date" in rtp.columns:
            rtp["status_start"] = pd.to_datetime(rtp["report_date"], errors="coerce")
            rtp = rtp.dropna(subset=["status_start"]).sort_values(["player_id", "status_start"])
        else:
            st.warning("Cannot determine status dates — need 'status_start' or 'report_date'.")
            rtp = pd.DataFrame()

        if not rtp.empty and "injury_status" in rtp.columns:
            rtp["status_rank"] = rtp["injury_status"].map(rtp_status_order).fillna(2.5)
            latest_rtp = rtp.sort_values("status_start").groupby("player_id").tail(1).copy()
            non_avail = latest_rtp[~latest_rtp["injury_status"].str.lower().isin(["full go", "available"])]

            if non_avail.empty:
                st.success("All players are currently Full Go / Available.")
            else:
                st.markdown("### Currently non-available players")
                rc1, rc2, rc3 = st.columns(3)
                rc1.metric("🔴 Out", int((non_avail["injury_status"] == "Out").sum()))
                rc2.metric("🟡 Limited", int((non_avail["injury_status"] == "Limited").sum()))
                rc3.metric("🟠 As Tolerated", int((non_avail["injury_status"] == "As Tolerated").sum()))
                show_rtp = [c for c in ["player_id", "injury_status", "status_start", "days_in_status", "injury_date"] if c in non_avail.columns]
                _rtp_display = non_avail.sort_values("status_rank")[show_rtp].copy()
                for _dc in ["status_start", "injury_date"]:
                    if _dc in _rtp_display.columns:
                        _rtp_display[_dc] = pd.to_datetime(_rtp_display[_dc], errors="coerce").dt.strftime("%m/%d/%Y")
                st.dataframe(pretty_cols(_rtp_display), use_container_width=True, hide_index=True)

            st.markdown("### Player Status Timeline")
            rtp_players = sorted(rtp["player_id"].unique().tolist())
            rtp_sel = st.multiselect(
                "Select players to view", rtp_players,
                default=non_avail["player_id"].tolist() if not non_avail.empty else rtp_players[:5],
                key="rtp_sel",
            )
            if rtp_sel:
                rtp_view = rtp[rtp["player_id"].isin(rtp_sel)].copy()
                rtp_view["status_end"] = rtp_view.groupby("player_id")["status_start"].shift(-1)
                rtp_view["status_end"] = rtp_view["status_end"].fillna(pd.Timestamp.today())
                fig_gantt = px.timeline(
                    rtp_view, x_start="status_start", x_end="status_end", y="player_id",
                    color="injury_status",
                    color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e", "Available": "#22c55e"},
                    title="Status progression (Gantt)",
                    hover_data=[c for c in ["days_in_status", "injury_date"] if c in rtp_view.columns],
                )
                fig_gantt.update_yaxes(autorange="reversed")
                fig_gantt.update_layout(height=max(300, 60 * len(rtp_sel)), margin=dict(l=10, r=10, t=50, b=30), xaxis=dict(tickformat="%m/%d/%Y"))
                st.plotly_chart(fig_gantt, use_container_width=True)

            st.markdown("### Status Transitions")
            st.caption("How often does each status transition to the next? (row → column)")
            rtp_sorted = rtp.sort_values(["player_id", "status_start"]).copy()
            rtp_sorted["next_status"] = rtp_sorted.groupby("player_id")["injury_status"].shift(-1)
            trans = rtp_sorted.dropna(subset=["next_status"])
            if not trans.empty:
                matrix = pd.crosstab(trans["injury_status"], trans["next_status"])
                st.dataframe(matrix, use_container_width=True)
                trans_counts = trans.groupby(["injury_status", "next_status"]).size().reset_index(name="count")
                fig_sun = px.sunburst(
                    trans_counts, path=["injury_status", "next_status"], values="count",
                    title="Status transition flow", color="injury_status",
                    color_discrete_map={"Out": "#ef4444", "Limited": CU_GOLD, "As Tolerated": CU_LIGHT_GRAY, "Full Go": "#22c55e", "Available": "#22c55e"},
                )
                fig_sun.update_layout(height=450)
                st.plotly_chart(fig_sun, use_container_width=True)

            st.markdown("### Average Days Spent in Each Status")
            if "days_in_status" in rtp.columns:
                avg_days = rtp.groupby("injury_status")["days_in_status"].agg(["mean", "median", "max"]).reset_index()
                avg_days.columns = ["Status", "Mean days", "Median days", "Max days"]
                avg_days = avg_days.sort_values("Mean days", ascending=False)
                st.dataframe(avg_days, use_container_width=True, hide_index=True)
                fig_avg = px.bar(avg_days, x="Status", y=["Mean days", "Median days"],
                                 barmode="group", title="Days in status by category")
                st.plotly_chart(fig_avg, use_container_width=True)

    else:  # Wellness section
      st.markdown("#### 💚 Wellness / Internal Load (2024–2025)")
      if wellness_all.empty:
        st.warning("No wellness data loaded.")
      else:
        st.markdown(f"**{len(wellness_all)} wellness reports** across {wellness_all['player_id'].nunique()} players.")

        if "wellness_total" in df_train.columns:
            c1, c2 = st.columns(2)
            coverage = float(df_train["wellness_total"].notna().mean()) * 100.0
            c1.metric("Wellness coverage", f"{coverage:.1f}%")
            if "year" in df_train.columns:
                by_year = (
                    df_train.groupby("year", as_index=False)
                    .agg(rows=("player_id", "count"), with_wellness=("wellness_total", lambda x: x.notna().sum()))
                )
                by_year["coverage_pct"] = 100 * by_year["with_wellness"] / by_year["rows"].replace(0, np.nan)
                c2.metric("Rows without wellness", f"{100 - coverage:.1f}%")
                st.dataframe(pretty_cols(by_year), use_container_width=True)

        if "wellness_date" in wellness_all.columns and "wellness_total" in wellness_all.columns:
            st.markdown("### Squad Wellness Over Time")
            daily_well = wellness_all.groupby("wellness_date", as_index=False).agg(
                avg_total=("wellness_total", "mean"),
                avg_mental=("mental_score", "mean"),
                avg_physical=("physical_score", "mean"),
                avg_sleep=("sleep_score", "mean"),
                avg_soreness=("soreness_score", "mean"),
                reports=("player_id", "count"),
            )
            st.plotly_chart(px.line(daily_well, x="wellness_date", y="avg_total", markers=True,
                                    title="Squad Avg Overall Wellness (out of 20)").update_layout(xaxis=dict(tickformat="%m/%d/%Y")), use_container_width=True)
            st.markdown("### Wellness Sub-Scores Over Time")
            st.plotly_chart(px.line(daily_well, x="wellness_date",
                                    y=["avg_mental", "avg_physical", "avg_sleep", "avg_soreness"],
                                    markers=True, title="Squad Avg Sub-Scores (each out of 5)").update_layout(xaxis=dict(tickformat="%m/%d/%Y")), use_container_width=True)

        st.markdown("### Latest Wellness per Player")
        latest_well = wellness_all.sort_values("wellness_date").groupby("player_id").last().reset_index()
        score_cols = [c for c in ["player_id", "wellness_date", "wellness_total", "mental_score",
                                  "physical_score", "sleep_score", "soreness_score"] if c in latest_well.columns]
        _well_display = latest_well.sort_values("wellness_total")[score_cols].copy()
        if "wellness_date" in _well_display.columns:
            _well_display["wellness_date"] = pd.to_datetime(_well_display["wellness_date"], errors="coerce").dt.strftime("%m/%d/%Y")
        st.dataframe(pretty_cols(_well_display), use_container_width=True)

        sub_scores = [c for c in ["mental_score", "physical_score", "sleep_score", "soreness_score"] if c in latest_well.columns]
        if sub_scores:
            heat = latest_well.set_index("player_id")[sub_scores].dropna()
            if len(heat) > 0:
                st.markdown("### Player Wellness Heatmap (latest report)")
                st.plotly_chart(px.imshow(heat, aspect="auto", color_continuous_scale="RdYlGn",
                                          title="Wellness sub-scores by player (5 = best)"), use_container_width=True)

        if "wellness_total" in df_train.columns and df_train["wellness_total"].notna().sum() > 10:
            st.markdown("### Wellness vs Training Load")
            wl = df_train.dropna(subset=["wellness_total", "total_player_load"]) if "total_player_load" in df_train.columns else pd.DataFrame()
            if len(wl) > 10:
                st.plotly_chart(px.scatter(wl, x="wellness_total", y="total_player_load",
                                           color="player_id" if wl["player_id"].nunique() <= 15 else None,
                                           trendline="ols", title="Wellness Total vs Player Load",
                                           labels={"wellness_total": pretty("wellness_total"), "total_player_load": pretty("total_player_load")}),
                                use_container_width=True)
            if "acwr_ewma_7_28" in df_train.columns:
                wa = df_train.dropna(subset=["wellness_total", "acwr_ewma_7_28"])
                if len(wa) > 10:
                    _fig_wa = px.scatter(wa, x="wellness_total", y="acwr_ewma_7_28",
                                         trendline="ols", title="Wellness Total vs ACWR")
                    _fig_wa.add_hline(y=0.8, line_dash="dash", line_color="#CFB87C", annotation_text="Min (0.8)")
                    _fig_wa.add_hline(y=1.3, line_dash="dash", line_color="#ef4444", annotation_text="Max (1.3)")
                    st.plotly_chart(_fig_wa, use_container_width=True)

        if "injury_status_actual" in df_train.columns and "wellness_total" in df_train.columns:
            wis = df_train.dropna(subset=["wellness_total"])
            wis = wis[~wis["injury_status_actual"].astype(str).str.strip().str.lower().isin(["nan", ""])].copy()
            if len(wis) > 10:
                st.markdown("### Wellness by Injury Status")
                st.plotly_chart(px.box(wis, x="injury_status_actual", y="wellness_total",
                                       title="Overall Wellness by Injury Status",
                                       labels={"injury_status_actual": pretty("injury_status_actual"), "wellness_total": pretty("wellness_total")}), use_container_width=True)

        st.download_button("Download Wellness Data", data=wellness_all.to_csv(index=False).encode(),
                           file_name="wellness_all.csv", mime="text/csv")

# ==============================================================================
# TAB: POSITION GROUPS
# ==============================================================================
with tab_pos:
    st.subheader("👥 Position-Group Analysis")
    st.caption("Compare physical output, readiness, and fatigue across positional lines.")
    coach_guide("Position Groups", [
        "Splits the squad by **DEF / MID / FWD / GK** for load, readiness, fatigue, and speed profile.",
        "Spot when **one line is carrying more** and rebalance drills or who sits out a finishing block.",
        "Thresholds are position-specific because demands differ (e.g. forwards sprint more; defenders decelerate more).",
    ])
    metric_tooltip("position_groups")

    if "position_name" not in df_train.columns or df_train["position_name"].dropna().nunique() < 2:
        st.warning("Not enough position data to build group comparisons.")
    else:
        pos_df = df_train.copy()
        raw_pos = pos_df["position_name"].astype(str).str.strip()

        # Map raw positions to groups (customize as needed for your team's naming)
        def _pos_group(p: str) -> str:
            pl = p.lower()
            if any(k in pl for k in ["gk", "keeper", "goal"]):
                return "GK"
            if any(k in pl for k in ["def", "back", "cb", "rb", "lb", "fb"]):
                return "DEF"
            if any(k in pl for k in ["mid", "cm", "cdm", "cam", "dm", "am"]):
                return "MID"
            if any(k in pl for k in ["fwd", "forward", "striker", "wing", "att", "st", "cf", "lw", "rw"]):
                return "FWD"
            return "OTHER"

        pos_df["pos_group"] = raw_pos.apply(_pos_group)
        grp_order = ["GK", "DEF", "MID", "FWD", "OTHER"]
        pos_df["pos_group"] = pd.Categorical(pos_df["pos_group"], categories=grp_order, ordered=True)

        grp_counts = pos_df.groupby("pos_group")["player_id"].nunique().reset_index(name="players")
        gc1, gc2, gc3, gc4 = st.columns(4)
        for i, g in enumerate(["GK", "DEF", "MID", "FWD"]):
            n = int(grp_counts.loc[grp_counts["pos_group"] == g, "players"].values[0]) if g in grp_counts["pos_group"].values else 0
            [gc1, gc2, gc3, gc4][i].metric(g, f"{n} players")

        pos_metrics = [c for c in [
            "total_player_load", "player_load_per_min_est", "maximum_velocity",
            "hsd_m", "sprint_m", "total_acceleration_load",
            "acwr_ewma_7_28", "readiness", "fatigue_idx", "wellness_total",
        ] if c in pos_df.columns]

        # Box plots per metric
        st.markdown("### Load & Performance by Position Group")
        pos_metric_sel = st.selectbox("Metric", pos_metrics, index=0, key="pos_metric")
        fig_pbox = px.box(
            pos_df.dropna(subset=[pos_metric_sel]),
            x="pos_group", y=pos_metric_sel, color="pos_group",
            points="outliers",
            color_discrete_map={"GK": "#f59e0b", "DEF": "#6366f1", "MID": "#22c55e", "FWD": "#ef4444", "OTHER": "#9ca3af"},
            title=f"{pretty(pos_metric_sel)} distribution by position group",
            labels={"pos_group": pretty("pos_group"), pos_metric_sel: pretty(pos_metric_sel)},
        )
        st.plotly_chart(fig_pbox, use_container_width=True)

        # Radar comparison of position groups
        st.markdown("### Position-Group Radar")
        radar_pos_metrics = [c for c in [
            "total_player_load", "maximum_velocity", "acwr_ewma_7_28",
            "readiness", "fatigue_idx", "wellness_total",
        ] if c in pos_df.columns]

        if len(radar_pos_metrics) >= 3:
            fig_pos_radar = go.Figure()
            pos_colors = {"GK": "#f59e0b", "DEF": "#6366f1", "MID": "#22c55e", "FWD": "#ef4444"}
            for grp in ["GK", "DEF", "MID", "FWD"]:
                gdata = pos_df[pos_df["pos_group"] == grp]
                if gdata.empty:
                    continue
                vals = []
                for rm in radar_pos_metrics:
                    col_data = pd.to_numeric(pos_df[rm], errors="coerce").dropna()
                    vmin = float(col_data.quantile(0.10))
                    vmax = float(col_data.quantile(0.90))
                    grp_mean = float(pd.to_numeric(gdata[rm], errors="coerce").mean())
                    if vmax > vmin:
                        vals.append(float(np.clip(100 * (grp_mean - vmin) / (vmax - vmin), 0, 100)))
                    else:
                        vals.append(50.0)

                fig_pos_radar.add_trace(go.Scatterpolar(
                    r=vals,
                    theta=[c.replace("_", " ").title() for c in radar_pos_metrics],
                    fill="toself",
                    name=grp,
                    line=dict(color=pos_colors.get(grp, "#9ca3af")),
                ))

            fig_pos_radar.update_layout(
                title="Position-Group Comparison (normalized to squad percentiles)",
                polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                showlegend=True, height=500,
            )
            st.plotly_chart(fig_pos_radar, use_container_width=True)
        else:
            st.info("Not enough shared metrics for a radar chart.")

        # Readiness / fatigue trends by position group over time
        st.markdown("### Readiness / Fatigue Trends by Position Group")
        if "readiness" in pos_df.columns:
            _tmp = pos_df[["date", "pos_group", "readiness"]].copy()
            _tmp["pos_group"] = _tmp["pos_group"].astype(str)
            if "fatigue_idx" in pos_df.columns:
                _tmp["fatigue_idx"] = pos_df["fatigue_idx"].values
            _tmp["week"] = pd.to_datetime(_tmp["date"]).dt.to_period("W").dt.start_time
            agg_dict = {"avg_readiness": ("readiness", "mean")}
            if "fatigue_idx" in _tmp.columns:
                agg_dict["avg_fatigue"] = ("fatigue_idx", "mean")
            daily_pos = _tmp.groupby(["week", "pos_group"], as_index=False).agg(**agg_dict)

            fig_pos_time = px.line(
                daily_pos.dropna(subset=["avg_readiness"]),
                x="week", y="avg_readiness", color="pos_group",
                markers=True,
                color_discrete_map={"GK": "#f59e0b", "DEF": "#6366f1", "MID": "#22c55e", "FWD": "#ef4444", "OTHER": "#9ca3af"},
                title="Weekly avg readiness by position group",
                labels={"week": "Week", "avg_readiness": "Avg Readiness", "pos_group": pretty("pos_group")},
            )
            fig_pos_time.update_layout(xaxis=dict(tickformat="%m/%d/%Y"))
            st.plotly_chart(fig_pos_time, use_container_width=True)

            if "avg_fatigue" in daily_pos.columns:
                fig_fat_time = px.line(
                    daily_pos.dropna(subset=["avg_fatigue"]),
                    x="week", y="avg_fatigue", color="pos_group",
                    markers=True,
                    color_discrete_map={"GK": "#f59e0b", "DEF": "#6366f1", "MID": "#22c55e", "FWD": "#ef4444", "OTHER": "#9ca3af"},
                    title="Weekly avg fatigue index by position group",
                    labels={"week": "Week", "avg_fatigue": "Avg Fatigue", "pos_group": pretty("pos_group")},
                )
                fig_fat_time.update_layout(xaxis=dict(tickformat="%m/%d/%Y"))
                st.plotly_chart(fig_fat_time, use_container_width=True)

        # Summary table
        st.markdown("### Position-Group Summary (latest data)")
        latest_pos = pos_df.sort_values("date").groupby("player_id", as_index=False).tail(1)
        summary_metrics = [c for c in ["total_player_load", "maximum_velocity", "acwr_ewma_7_28",
                                        "readiness", "fatigue_idx", "wellness_total"] if c in latest_pos.columns]
        pos_summary = latest_pos.groupby("pos_group", as_index=False)[summary_metrics].mean(numeric_only=True)
        pos_summary = pos_summary.round(2)
        st.dataframe(pretty_cols(pos_summary), use_container_width=True, hide_index=True)

# ==============================================================================
# TAB: ANOMALIES
# ==============================================================================
with tab_anom:
    st.subheader("⚠️ Unusual Sessions")
    st.caption("Sessions that stand out from normal training patterns — potential overload, under-load, or data issues.")
    coach_guide("Anomalies", [
        "Auto-flags sessions that look **unusual across load, speed, distance, accel, and ACWR at once**.",
        "The **sensitivity slider** (sidebar) decides what % of sessions are flagged (default ~4%).",
        "Expected flags: match days often look ‘unusual’ vs training — that’s normal, not a problem.",
    ])
    metric_tooltip("anomaly_score")

    with st.expander("ℹ️ What does this mean? (click to read)", expanded=False):
        st.markdown("""
**Think of it like this:** if most training sessions look "normal," these are the ones that don't.

The system checks each session across **load, speed, distance, acceleration, and ACWR all at once**.
If a session is unusual in multiple ways simultaneously, it gets flagged.

**What to look for:**
- 🔴 **Session way above normal** → Was this player supposed to train that hard? Potential overload.
- 🔵 **Session way below normal** → Did the player leave early? Resting? Or a GPS tracking issue?
- **Match days** will often show up because games are naturally different from training — that's expected.

**The sensitivity slider** in the sidebar controls how strict this is (default ~4% flagged).
        """)

    _anom = df_train.copy()
    if "anomaly_score" not in _anom.columns or _anom["anomaly_score"].isna().all():
        st.warning("Anomaly scores not available. Check that enough data is loaded.")
    else:
        n_flagged = int(_anom["anomaly_flag"].sum())
        n_total = len(_anom)
        n_players_flagged = _anom.loc[_anom["anomaly_flag"], "player_id"].nunique() if n_flagged else 0

        ac1, ac2, ac3, ac4 = st.columns(4)
        ac1.metric("Total Sessions", f"{n_total:,}")
        ac2.metric("Flagged Anomalies", n_flagged)
        ac3.metric("Anomaly Rate", f"{100*n_flagged/max(n_total,1):.1f}%")
        ac4.metric("Players Affected", n_players_flagged)

        # ── Per-player anomaly counts ──
        st.markdown("### 🧑 Which Players Have the Most Unusual Sessions?")
        _player_anom = _anom[_anom["anomaly_flag"]].groupby("player_id").agg(
            count=("anomaly_flag", "sum"),
            avg_score=("anomaly_score", "mean"),
            sessions=("date", "count"),
        ).reset_index().sort_values("count", ascending=False)
        if not _player_anom.empty:
            _player_anom["pct_flagged"] = (100 * _player_anom["count"] / _anom.groupby("player_id").size().reindex(_player_anom["player_id"]).values).round(1)
            fig_pa = px.bar(
                _player_anom.head(20), x="player_id", y="count",
                hover_data=["avg_score", "pct_flagged"],
                title="Top 20 Players by Anomaly Count",
                color="avg_score", color_continuous_scale="YlOrRd",
                labels={"count": "Anomalies", "player_id": "Player", "avg_score": "Avg Score"},
            )
            fig_pa.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", height=380)
            st.plotly_chart(fig_pa, use_container_width=True)

        # ── Anomaly timeline ──
        st.markdown("### 📅 When Did Unusual Sessions Happen?")
        _daily_anom = _anom.groupby("date").agg(
            total=("anomaly_flag", "size"), flagged=("anomaly_flag", "sum")
        ).reset_index()
        _daily_anom["rate"] = 100 * _daily_anom["flagged"] / _daily_anom["total"]
        fig_timeline = go.Figure()
        fig_timeline.add_trace(go.Bar(
            x=_daily_anom["date"], y=_daily_anom["flagged"], name="Flagged",
            marker_color="#ef4444", opacity=0.8,
        ))
        fig_timeline.add_trace(go.Scatter(
            x=_daily_anom["date"], y=_daily_anom["rate"], name="Anomaly %",
            yaxis="y2", mode="lines", line=dict(color="#f59e0b", width=2),
        ))
        fig_timeline.update_layout(
            title="Daily Anomaly Count & Rate",
            xaxis=dict(tickformat="%m/%d/%Y"),
            yaxis=dict(title="Flagged Sessions"),
            yaxis2=dict(title="Anomaly %", overlaying="y", side="right", showgrid=False),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=350, legend=dict(orientation="h", y=1.12),
        )
        st.plotly_chart(fig_timeline, use_container_width=True)

        # ── Anomaly by MD window ──
        if "md_label" in _anom.columns:
            st.markdown("### 📌 Are Unusual Sessions Linked to Game Days?")
            _md_anom = _anom.groupby("md_label").agg(
                total=("anomaly_flag", "size"), flagged=("anomaly_flag", "sum")
            ).reset_index()
            _md_anom["rate"] = (100 * _md_anom["flagged"] / _md_anom["total"]).round(1)
            _md_anom = _md_anom[_md_anom["total"] >= 5].sort_values("rate", ascending=False)
            if not _md_anom.empty:
                fig_md_anom = px.bar(
                    _md_anom, x="md_label", y="rate",
                    title="Anomaly Rate (%) by MD Window",
                    color="rate", color_continuous_scale="RdYlGn_r",
                    labels={"rate": "Anomaly %", "md_label": "MD Window"},
                    hover_data=["flagged", "total"],
                )
                fig_md_anom.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", height=350)
                st.plotly_chart(fig_md_anom, use_container_width=True)
                st.caption("Higher anomaly rates on match days (MD 0) and MD+1 are expected — game demands differ from training norms.")

        # ── Scatter: Load vs Velocity ──
        st.markdown("### 🔬 Training Load vs Speed")
        if "total_player_load" in _anom.columns and "maximum_velocity" in _anom.columns:
            _anom["flag_label"] = _anom["anomaly_flag"].map({True: "⚠️ Anomaly", False: "Normal"})
            fig_sc = px.scatter(
                _anom, x="total_player_load", y="maximum_velocity",
                color="flag_label",
                color_discrete_map={"⚠️ Anomaly": "#ef4444", "Normal": "#3b82f6"},
                opacity=0.6,
                hover_data=[c for c in ["player_id", "date", "md_label", "acwr_ewma_7_28", "anomaly_score"] if c in _anom.columns],
                title="Each dot = one training session",
                labels={"total_player_load": "Total Player Load", "maximum_velocity": "Max Velocity (mph)", "flag_label": "Status"},
            )
            fig_sc.update_traces(marker_size=5, selector=dict(name="Normal"))
            fig_sc.update_traces(marker_size=9, selector=dict(name="⚠️ Anomaly"))
            fig_sc.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", height=450)
            st.plotly_chart(fig_sc, use_container_width=True)

        # ── Normal vs flagged comparison ──
        st.markdown("### 📊 How Do Flagged Sessions Differ From Normal?")
        st.caption("Pick a metric to compare its distribution across normal vs. flagged sessions.")
        _compare_cols = [c for c in [
            "total_player_load", "maximum_velocity", "total_distance",
            "total_acceleration_load", "explosive_efforts", "acwr_ewma_7_28",
        ] if c in _anom.columns]
        if _compare_cols and n_flagged > 0:
            anom_compare_metric = st.selectbox(
                "Compare metric",
                _compare_cols,
                format_func=lambda x: pretty(x),
                key="anom_compare_metric",
            )
            _anom["_status"] = _anom["anomaly_flag"].map({True: "Flagged", False: "Normal"})
            fig_comp = px.box(
                _anom, x="_status", y=anom_compare_metric, color="_status",
                color_discrete_map={"Normal": "#3b82f6", "Flagged": "#ef4444"},
                title=f"{pretty(anom_compare_metric)}: Normal vs Flagged",
                points="outliers",
                labels={"_status": "Session Type", anom_compare_metric: pretty(anom_compare_metric)},
            )
            fig_comp.update_layout(
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                height=420, showlegend=False,
            )
            st.plotly_chart(fig_comp, use_container_width=True)

            _norm_avg = _anom.loc[~_anom["anomaly_flag"], anom_compare_metric].mean()
            _flag_avg = _anom.loc[_anom["anomaly_flag"], anom_compare_metric].mean()
            _diff_pct = 100 * (_flag_avg - _norm_avg) / _norm_avg if _norm_avg else 0
            st.markdown(f"**Normal avg:** {_norm_avg:,.1f} · **Flagged avg:** {_flag_avg:,.1f} · "
                        f"**Difference:** {_diff_pct:+.1f}%")

        # ── Flagged sessions table ──
        st.markdown("### 📋 Full List of Flagged Sessions")
        show_cols = [c for c in [
            "date", "player_id", "position_name", "md_label",
            "total_player_load", "maximum_velocity", "acwr_ewma_7_28",
            "anomaly_score",
        ] if c in _anom.columns]
        flagged_df = _anom[_anom["anomaly_flag"]].sort_values("anomaly_score", ascending=False)[show_cols]
        _flagged_display = flagged_df.head(60).copy()
        if "date" in _flagged_display.columns:
            _flagged_display["date"] = pd.to_datetime(_flagged_display["date"], errors="coerce").dt.strftime("%m/%d/%Y")

        def _highlight_flagged(row):
            return ["background-color: rgba(239, 68, 68, 0.15)"] * len(row)

        st.dataframe(
            pretty_cols(_flagged_display).style.apply(_highlight_flagged, axis=1),
            use_container_width=True, hide_index=True,
        )
        st.download_button(
            "📥 Download all flagged sessions",
            data=flagged_df.to_csv(index=False).encode("utf-8"),
            file_name="anomaly_flagged_sessions.csv", mime="text/csv",
        )

# ==============================================================================
# TAB: FORECAST — Next-session load / ACWR / vmax prediction
# ==============================================================================
with tab_forecast:
    st.subheader("📈 What Should Tomorrow Look Like?")
    st.caption("Predicts each player's next-session metric based on their recent training history.")
    coach_guide("Forecast", [
        "Projects each player’s **next-session value** for a chosen metric (e.g. load).",
        "Uses lagged values, 3- & 7-session rolling means, std dev, and % change.",
        "Tests **Ridge / RF / GBDT / MLP** and keeps the model with the **lowest MAE** for your data.",
    ])
    metric_tooltip("forecast")

    with st.expander("ℹ️ How does this work? (click to read)", expanded=False):
        st.markdown("""
The model looks at each player's **last few sessions** and predicts what their next session
should look like. Then we compare the prediction to what actually happened.

**Why it's useful:**
- 🔴 **Actual much higher than predicted** = unexpected spike — was this planned?
- 🟢 **Actual much lower than predicted** = player did less than expected — rest day or pulled early?

The model only uses past data to predict — no cheating with future information.
        """)


    _fc_labels = {"total_player_load": "Training Load", "acwr_ewma_7_28": "ACWR (Workload Ratio)", "maximum_velocity": "Top Speed"}
    _fc_options = list(_fc_labels.keys())
    forecast_target = st.selectbox(
        "What to predict", _fc_options,
        format_func=lambda x: _fc_labels[x],
        index=0, key="fc_target",
    )
    fc_lags = st.slider("How many past sessions to use", 3, 10, 5, key="fc_lags")

    fc_data, fc_feats = build_forecast_dataset(df_train, forecast_target, lags=fc_lags)
    if len(fc_data) < 40:
        st.warning("Not enough data for forecasting after lag construction.")
    else:
        fc_data = fc_data.sort_values("date")
        split_idx = int(len(fc_data) * 0.8)
        train_fc = fc_data.iloc[:split_idx]
        test_fc = fc_data.iloc[split_idx:]

        X_tr_fc, y_tr_fc = train_fc[fc_feats].values, train_fc[forecast_target].values
        X_te_fc, y_te_fc = test_fc[fc_feats].values, test_fc[forecast_target].values

        with st.spinner("Training forecast models..."):
            leader_fc, best_name_fc, best_mdl_fc = run_forecast_benchmark(X_tr_fc, y_tr_fc, X_te_fc, y_te_fc)

        with st.expander("🔬 Model comparison (4 algorithms tested)", expanded=False):
            st.dataframe(leader_fc, use_container_width=True, hide_index=True)

        c1, c2, c3 = st.columns(3)
        best_row = leader_fc.iloc[0]
        c1.metric("Best Algorithm", best_name_fc)
        c2.metric("Avg Error (MAE)", f"{best_row['MAE']:.1f}")
        c3.metric("Fit Quality (R²)", f"{best_row['R2']:.2f}")

        preds_fc = best_mdl_fc.predict(X_te_fc)
        test_fc = test_fc.copy()
        test_fc["predicted"] = preds_fc
        test_fc["error"] = test_fc[forecast_target] - test_fc["predicted"]

        # Forecast vs Actual — aggregate by date for clean visualization
        _fva_daily = test_fc.groupby("date", as_index=False).agg(
            actual=(forecast_target, "mean"),
            predicted=("predicted", "mean"),
            n_players=("player_id", "nunique"),
        ).sort_values("date")

        fig_fva = go.Figure()
        fig_fva.add_trace(go.Scatter(
            x=_fva_daily["date"], y=_fva_daily["actual"], mode="lines+markers",
            name="Actual (daily avg)", line=dict(color="#6366f1", width=2.5), marker=dict(size=5),
        ))
        fig_fva.add_trace(go.Scatter(
            x=_fva_daily["date"], y=_fva_daily["predicted"], mode="lines+markers",
            name="Predicted (daily avg)", line=dict(color="#22c55e", width=2.5, dash="dash"), marker=dict(size=5),
        ))
        fig_fva.update_layout(
            title=f"Forecast vs Actual — {pretty(forecast_target)} (daily team average)",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)", height=420,
            margin=dict(l=40, r=20, t=50, b=30),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
            xaxis=dict(tickformat="%m/%d/%Y"),
            yaxis_title=pretty(forecast_target),
        )
        if forecast_target == "acwr_ewma_7_28":
            fig_fva.add_hline(y=0.8, line_dash="dash", line_color="#CFB87C", annotation_text="Min (0.8)")
            fig_fva.add_hline(y=1.3, line_dash="dash", line_color="#ef4444", annotation_text="Max (1.3)")
        st.plotly_chart(fig_fva, use_container_width=True)

        # Error distribution
        col_e1, col_e2 = st.columns(2)
        with col_e1:
            fig_err = px.histogram(
                test_fc, x="error", nbins=40,
                title="Prediction Error Distribution",
                color_discrete_sequence=["#8b5cf6"],
                template="plotly_dark",
            )
            fig_err.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_err, use_container_width=True)

        with col_e2:
            fig_err_time = go.Figure()
            fig_err_time.add_trace(go.Bar(
                x=test_fc["date"], y=test_fc["error"],
                marker_color=[
                    "#ef4444" if e > 0 else "#22c55e" for e in test_fc["error"]
                ],
            ))
            fig_err_time.add_hline(y=0, line_dash="dash", line_color="white", opacity=0.5)
            fig_err_time.update_layout(
                title="Error Over Time (red = under-predicted spike)",
                template="plotly_dark", paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)", height=350,
                xaxis=dict(tickformat="%m/%d/%Y"),
            )
            st.plotly_chart(fig_err_time, use_container_width=True)

        # Per-player forecast scatter
        st.markdown("### How Accurate Is the Forecast per Player?")
        player_fc = test_fc.groupby("player_id").agg(
            mean_actual=(forecast_target, "mean"),
            mean_predicted=("predicted", "mean"),
            mae=("error", lambda x: np.abs(x).mean()),
        ).reset_index().sort_values("mae")
        fig_pp = px.scatter(
            player_fc, x="mean_actual", y="mean_predicted", size="mae",
            hover_data=["player_id", "mae"],
            title="Player Avg Actual vs Predicted",
            template="plotly_dark", color="mae",
            color_continuous_scale="RdYlGn_r",
        )
        fig_pp.add_trace(go.Scatter(
            x=[player_fc["mean_actual"].min(), player_fc["mean_actual"].max()],
            y=[player_fc["mean_actual"].min(), player_fc["mean_actual"].max()],
            mode="lines", line=dict(color="white", dash="dash"), name="Perfect",
        ))
        fig_pp.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", height=420,
        )
        st.plotly_chart(fig_pp, use_container_width=True)

        # Spike detection table
        st.markdown("### ⚡ Unexpected Sessions")
        st.caption("Sessions where the actual value was significantly different from what we expected.")
        err_std = test_fc["error"].std()
        alerts = test_fc[test_fc["error"].abs() > 1.5 * err_std][["player_id", "date", forecast_target, "predicted", "error"]].copy()
        alerts["alert_type"] = np.where(alerts["error"] > 0, "🔴 SPIKE (over-exertion)", "🟢 TAPER (under-load)")
        alerts = alerts.sort_values("error", ascending=False, key=abs)
        if alerts.empty:
            st.success("No significant spikes or tapers detected in holdout window.")
        else:
            _alerts_display = alerts.copy()
            if "date" in _alerts_display.columns:
                _alerts_display["date"] = pd.to_datetime(_alerts_display["date"], errors="coerce").dt.strftime("%m/%d/%Y")
            st.dataframe(pretty_cols(_alerts_display), use_container_width=True, hide_index=True)

        st.download_button(
            "📥 Download forecast results",
            test_fc.to_csv(index=False).encode(),
            "forecast_results.csv", "text/csv",
        )


# ==============================================================================
# TAB: COACH ML
# ==============================================================================
with tab_ml:
    st.subheader("🧠 Player Availability Predictor")
    st.caption("Predicts which players may be unavailable for the next session and tells you what to do about it.")
    coach_guide("Coach ML", [
        "Each player gets a traffic-light: 🔴 **RED ≥ 60%**, 🟡 **AMBER 35–60%**, 🟢 **GREEN < 35%** probability of being unavailable next session.",
        "The **Why?** column on the Risk Board explains *in plain English* what pushed each player into their color.",
        "This is **decision support** — sports medicine always has the final call on medical decisions.",
    ])
    metric_tooltip("risk_board")
    with st.expander("ℹ️ How does this work? (click to read)", expanded=False):
        st.markdown("""
**What it does:** Looks at each player's recent workload, wellness, injury history,
and ACWR to predict whether they'll be **available for the next session**.

**Traffic-light system:**
- 🔴 **RED** = High risk of being unavailable → reduce load, increase recovery
- 🟡 **AMBER** = Watch closely → hold or slightly reduce load
- 🟢 **GREEN** = Good to go → continue normal training

**Why is it so accurate?**
The model sees a player's current injury status — someone who is "Out" today is almost
certainly "Out" tomorrow. The most **useful** predictions are for currently-available
players who show warning signs in their workload or wellness data.

**How we validate:** The model trains on older data and predicts the most recent sessions,
just like it would in real life. No "peeking" at future data.
        """)


    ml_df, feat_cols = build_availability_risk_dataset(df_train)
    if ml_df.empty or len(feat_cols) < 5:
        st.warning("Not enough labeled data to train availability model with current filters.")
    else:
        y = ml_df["target_not_available"].astype(int)
        class_counts = y.value_counts().to_dict()
        if len(class_counts) < 2:
            st.warning("Only one target class present after filtering. Broaden filters to train the model.")
        else:
            X = ml_df[feat_cols].copy()

            # Time-aware split (last 20% as holdout)
            ml_sorted = ml_df.sort_values("date")
            cut = max(30, int(len(ml_sorted) * 0.8))
            train_idx = ml_sorted.index[:cut]
            test_idx = ml_sorted.index[cut:]
            if len(test_idx) < 10:
                train_idx = ml_sorted.index[: max(1, len(ml_sorted) - 10)]
                test_idx = ml_sorted.index[max(1, len(ml_sorted) - 10):]

            # Impute with train-only medians to prevent data leakage
            train_medians = X.loc[train_idx].median(numeric_only=True)
            X = X.fillna(train_medians)

            X_train, y_train = X.loc[train_idx].copy(), y.loc[train_idx].copy()
            X_test, y_test = X.loc[test_idx].copy(), y.loc[test_idx].copy()

            import hashlib as _hl
            _ml_hash = _hl.md5(f"{len(X_train)}_{len(X_test)}_{list(feat_cols)}".encode()).hexdigest()
            leaderboard, best_name, models, _ml_errors = _train_availability_models(
                _ml_hash, X_train, y_train, X_test, y_test, feat_cols
            )

            board_df = pd.DataFrame(leaderboard).sort_values(
                ["cv_roc_auc", "holdout_roc_auc", "holdout_accuracy"], ascending=False, na_position="last"
            ).reset_index(drop=True)
            board_df_display = board_df.rename(columns={
                "model": "Algorithm", "cv_roc_auc": "Cross-Val Score",
                "holdout_roc_auc": "Test Score", "holdout_accuracy": "Accuracy"
            })
            with st.expander(f"🔬 Model comparison ({len(models)} algorithms tested)", expanded=False):
                st.dataframe(board_df_display, use_container_width=True, hide_index=True)

            # ---- GBDT diagnostics (confusion matrix, ROC, full feature importance) — paper / reports ----
            st.markdown("---")
            st.markdown("### 📈 Gradient Boosting (GBDT) — diagnostics")
            st.caption(
                "Same temporal train/test split as above (80% earliest → train, 20% latest → test). "
                "Binary rule: P(unavailable) ≥ 0.5 → predict unavailable."
            )
            try:
                gbdt_diag = GradientBoostingClassifier(
                    random_state=42, max_depth=4, learning_rate=0.08, n_estimators=200,
                )
                gbdt_diag.fit(X_train, y_train)
                p_gb = gbdt_diag.predict_proba(X_test)[:, 1]
                y_hat_gb = (p_gb >= 0.5).astype(int)
                cm = confusion_matrix(y_test, y_hat_gb, labels=[0, 1])
                labels_txt = ["Available (next)", "Unavailable (next)"]
                fig_cm = go.Figure(
                    data=go.Heatmap(
                        z=cm,
                        x=labels_txt,
                        y=labels_txt,
                        colorscale="Blues",
                        text=cm,
                        texttemplate="%{text}",
                        colorbar=dict(title="Count"),
                    )
                )
                fig_cm.update_layout(
                    title="Confusion matrix — GBDT (holdout)",
                    xaxis_title="Predicted",
                    yaxis_title="Actual",
                    height=400,
                    margin=dict(l=60, r=20, t=60, b=60),
                )
                st.plotly_chart(fig_cm, use_container_width=True)

                fpr, tpr, _ = roc_curve(y_test, p_gb)
                roc_auc = auc(fpr, tpr) if len(np.unique(y_test)) > 1 else float("nan")
                fig_roc = go.Figure()
                fig_roc.add_trace(go.Scatter(x=fpr, y=tpr, mode="lines", name=f"GBDT (AUC = {roc_auc:.3f})", line=dict(color=CU_GOLD, width=3)))
                fig_roc.add_trace(go.Scatter(x=[0, 1], y=[0, 1], mode="lines", name="Chance", line=dict(color="#888", dash="dash")))
                fig_roc.update_layout(
                    title="ROC curve — GBDT (holdout)",
                    xaxis_title="False positive rate",
                    yaxis_title="True positive rate",
                    height=420,
                    margin=dict(l=50, r=20, t=50, b=50),
                )
                st.plotly_chart(fig_roc, use_container_width=True)

                imp_all = (
                    pd.DataFrame({"feature": feat_cols, "importance": gbdt_diag.feature_importances_})
                    .sort_values("importance", ascending=False)
                    .reset_index(drop=True)
                )
                imp_all["rank"] = imp_all.index + 1
                imp_all = imp_all[["rank", "feature", "importance"]]
                st.markdown("#### GBDT feature importance (full table)")
                st.dataframe(imp_all, use_container_width=True, hide_index=True)
                fig_imp_gb = px.bar(
                    imp_all.sort_values("importance", ascending=True).tail(20),
                    x="importance",
                    y="feature",
                    orientation="h",
                    title="Top features — GBDT (gini importance)",
                    color="importance",
                    color_continuous_scale="Viridis",
                )
                fig_imp_gb.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)",
                    plot_bgcolor="rgba(0,0,0,0)",
                    height=max(400, 24 * min(20, len(imp_all))),
                    coloraxis_showscale=False,
                    yaxis_title="",
                )
                st.plotly_chart(fig_imp_gb, use_container_width=True)

                with st.expander("Classification report (precision / recall / F1)", expanded=False):
                    st.code(
                        classification_report(
                            y_test,
                            y_hat_gb,
                            target_names=["Available next", "Unavailable next"],
                            digits=3,
                        ),
                        language="text",
                    )

                st.download_button(
                    "📥 Download GBDT feature importance (CSV)",
                    data=imp_all.to_csv(index=False).encode("utf-8"),
                    file_name="gbdt_feature_importance.csv",
                    mime="text/csv",
                    key="dl_gbdt_imp",
                )
            except Exception as _e_gb:
                st.warning(f"GBDT diagnostics could not be computed: {_e_gb}")

            if not best_name:
                st.warning("Model benchmarking failed under current filters. Try broadening date/year filters.")
            else:
                model = models[best_name]
                model.fit(X_train, y_train)
                p_test = model.predict_proba(X_test)[:, 1]
                y_hat = (p_test >= 0.50).astype(int)
                acc = accuracy_score(y_test, y_hat)
                best_auc = roc_auc_score(y_test, p_test) if y_test.nunique() > 1 else np.nan

                a, b, c, d = st.columns(4)
                a.metric("Sessions Analyzed", f"{len(ml_df):,}")
                b.metric("Times Unavailable", int(y.sum()))
                c.metric("Best Algorithm", best_name)
                d.metric("Prediction Accuracy", f"{acc:.0%}")

                if _ml_errors:
                    with st.expander(f"⚠️ {len(_ml_errors)} model(s) had issues", expanded=False):
                        for _err in _ml_errors:
                            st.caption(_err)

                # Rolling backtest
                st.markdown("### 🧪 Model Reliability Over Time")
                st.caption("Tests whether predictions stay accurate week by week — not just on one snapshot.")
                run_bt = st.checkbox("Run reliability check", value=False, help="Takes 10–30 seconds.")
                if run_bt:
                    with st.spinner("Running rolling backtest..."):
                        from sklearn.base import clone as _sklearn_clone
                        def _factory():
                            return _sklearn_clone(models[best_name])

                        bt = rolling_backtest_availability(ml_df, feat_cols, model_factory=_factory, train_days=60, test_days=7, step_days=7)
                    if bt.empty:
                        st.warning("Backtest could not run (not enough data density under current filters).")
                    else:
                        _bt_display = bt.sort_values("window_end", ascending=False).copy()
                        if "window_end" in _bt_display.columns:
                            _bt_display["window_end"] = pd.to_datetime(_bt_display["window_end"], errors="coerce").dt.strftime("%m/%d/%Y")
                        st.dataframe(pretty_cols(_bt_display), use_container_width=True, hide_index=True)
                        fig_bt = go.Figure()
                        fig_bt.add_trace(go.Scatter(x=bt["window_end"], y=bt["roc_auc"], mode="lines+markers", name="ROC-AUC", line=dict(color=CU_GOLD, width=3)))
                        fig_bt.add_trace(go.Scatter(x=bt["window_end"], y=bt["accuracy"], mode="lines+markers", name="Accuracy", line=dict(color="#22c55e", width=2, dash="dash")))
                        fig_bt.update_layout(title="Rolling performance over time", height=380, margin=dict(l=40, r=20, t=50, b=30), xaxis=dict(tickformat="%m/%d/%Y"))
                        st.plotly_chart(fig_bt, use_container_width=True)
                        st.metric("Avg rolling ROC-AUC", f"{bt['roc_auc'].mean():.2f}")

                # Feature importance/explainability fallback by model type
                importances = None
                try:
                    if hasattr(model, "feature_importances_"):
                        importances = model.feature_importances_
                    elif hasattr(model, "named_steps") and "model" in model.named_steps:
                        base = model.named_steps["model"]
                        if hasattr(base, "feature_importances_"):
                            importances = base.feature_importances_
                        elif hasattr(base, "coef_"):
                            coef = np.ravel(base.coef_)
                            importances = np.abs(coef)
                    elif hasattr(model, "coef_"):
                        coef = np.ravel(model.coef_)
                        importances = np.abs(coef)
                except Exception:
                    importances = None

                if importances is not None and len(importances) == len(feat_cols):
                    _feat_labels = {
                        "total_player_load": "Training Load", "acwr_ewma_7_28": "ACWR",
                        "maximum_velocity": "Top Speed", "wellness_total": "Wellness Score",
                        "fatigue_idx": "Fatigue Index", "readiness": "Readiness Score",
                        "total_distance": "Distance (mi)", "explosive_efforts": "Explosive Efforts",
                        "total_acceleration_load": "Acceleration Load",
                        "unavail_lag1": "Unavail Prev Session", "unavail_lag2": "Unavail 2 Sessions Ago",
                        "unavail_lag3": "Unavail 3 Sessions Ago", "unavail_roll7": "Unavail Count (7 sess)",
                        "unavail_roll14": "Unavail Count (14 sess)", "unavail_pct28": "Unavail Rate (28 sess)",
                        "acwr_delta": "ACWR Change", "acwr_lag1": "Previous ACWR",
                        "fatigue_lag1": "Previous Fatigue", "readiness_lag1": "Previous Readiness",
                        "player_load_per_min_est": "Load/Min",
                    }
                    imp = pd.DataFrame({"feature": feat_cols, "importance": importances})
                    imp["label"] = imp["feature"].map(lambda f: _feat_labels.get(f, f.replace("_", " ").title()))
                    imp = imp.sort_values("importance", ascending=False).head(10)
                    fig_imp = px.bar(
                        imp, x="importance", y="label", orientation="h",
                        title="What matters most for predicting availability",
                        color="importance", color_continuous_scale="Viridis",
                    )
                    fig_imp.update_layout(
                        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                        height=380, coloraxis_showscale=False, yaxis_title="",
                    )
                    st.plotly_chart(fig_imp, use_container_width=True)

                st.markdown("---")
                # Player-level actionable risk board — use ml_df's last row per player
                # (already has all engineered features from build_availability_risk_dataset)
                latest = ml_df.sort_values("date").groupby("player_id", as_index=False).tail(1).copy()
                for c0 in feat_cols:
                    if c0 not in latest.columns:
                        latest[c0] = np.nan
                X_latest = latest[feat_cols].copy().fillna(train_medians)
                latest["risk_next_non_available"] = model.predict_proba(X_latest)[:, 1]
                latest["coach_flag"] = np.where(
                    latest["risk_next_non_available"] >= 0.60, "RED",
                    np.where(latest["risk_next_non_available"] >= 0.35, "AMBER", "GREEN")
                )
                latest["coach_action"] = np.where(
                    latest["coach_flag"] == "RED",
                    "Reduce load 20-30%, boost recovery/wellness follow-up, monitor daily",
                    np.where(
                        latest["coach_flag"] == "AMBER",
                        "Hold or reduce load 10-15%, monitor soreness/sleep before full intensity",
                        "Normal progression; maintain planned load"
                    )
                )
                latest["why"] = latest.apply(explain_risk_row, axis=1)
                show_cols = [c for c in [
                    "player_id", "date", "injury_status_actual", "wellness_total",
                    "acwr_ewma_7_28", "total_player_load", "risk_next_non_available",
                    "coach_flag", "why", "coach_action"
                ] if c in latest.columns]
                st.markdown("### 🎯 Player Risk Board")
                with st.expander("🧑‍🏫 How to read the Risk Board (RED / AMBER / GREEN)", expanded=False):
                    st.markdown(
                        "**Color = model's next-session unavailability probability, thresholded:**\n\n"
                        "- 🔴 **RED** (≥ 60%) — act today: reduce load 20–30%, prioritize recovery, coordinate with AT.\n"
                        "- 🟡 **AMBER** (35–60%) — monitor: hold or reduce 10–15%, check soreness/sleep before full intensity.\n"
                        "- 🟢 **GREEN** (< 35%) — normal progression.\n\n"
                        "**What drives the color** (used in the **Why?** column):\n"
                        "- **Injury status** from the log (Out / Limited / As Tolerated) — strongest signal.\n"
                        "- **ACWR** — danger > 1.30, under-load < 0.80.\n"
                        "- **Fatigue index** — high > 1.50, rising > 1.20.\n"
                        "- **Wellness total** — low < 12/20, moderate < 14/20.\n"
                        "- **Sleep / Mental** sub-scores below 3/5.\n"
                        "- **Anomaly flag** — current session looks unusual across load/speed/distance/ACWR.\n\n"
                        "The **Why?** column shows the actual values that pushed each player into their color."
                    )
                risk_sorted = latest.sort_values("risk_next_non_available", ascending=False)[show_cols]
                n_red = (latest["coach_flag"] == "RED").sum()
                n_amber = (latest["coach_flag"] == "AMBER").sum()
                n_green = (latest["coach_flag"] == "GREEN").sum()
                rc1, rc2, rc3 = st.columns(3)
                rc1.metric("🔴 High Risk", n_red)
                rc2.metric("🟡 Medium Risk", n_amber)
                rc3.metric("🟢 Low Risk", n_green)

                fig_risk = px.bar(
                    latest.groupby("coach_flag").size().reset_index(name="count"),
                    x="coach_flag", y="count", color="coach_flag",
                    color_discrete_map={"RED": "#ef4444", "AMBER": "#f59e0b", "GREEN": "#22c55e"},
                    title="Risk Distribution",
                    template="plotly_dark",
                )
                fig_risk.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    showlegend=False, height=300,
                )
                st.plotly_chart(fig_risk, use_container_width=True)

                _risk_display = risk_sorted.copy()
                if "date" in _risk_display.columns:
                    _risk_display["date"] = pd.to_datetime(_risk_display["date"], errors="coerce").dt.strftime("%m/%d/%Y")
                st.dataframe(pretty_cols(_risk_display), use_container_width=True, hide_index=True)

                st.download_button(
                    "📥 Download Player Risk Board",
                    data=risk_sorted.to_csv(index=False).encode("utf-8"),
                    file_name="coach_ml_player_risk_board.csv",
                    mime="text/csv",
                )

                # SHAP explanations
                if SHAP_AVAILABLE:
                    st.markdown("### 🔎 Why is this player flagged?")
                    st.caption("Select a player to see which factors are pushing their risk up or down.")
                    try:
                        _shap_players = sorted(latest["player_id"].astype(str).unique().tolist())
                        if not _shap_players:
                            st.info("No players available for SHAP analysis.")
                        else:
                            shap_pid = st.selectbox(
                                "Select player",
                                _shap_players,
                                key="shap_pid",
                            )
                            row = latest[latest["player_id"].astype(str) == str(shap_pid)].head(1).copy()
                            if len(row) == 0:
                                st.info("No data for that player.")
                            else:
                                x_row = row[feat_cols].copy().fillna(train_medians)
                                base_model = model
                                if hasattr(model, "named_steps") and "model" in model.named_steps:
                                    base_model = model.named_steps["model"]
                                    try:
                                        x_explain = model.named_steps["scaler"].transform(x_row.values)
                                        x_explain_df = pd.DataFrame(x_explain, columns=feat_cols)
                                    except Exception:
                                        x_explain_df = x_row
                                else:
                                    x_explain_df = x_row
                                bg_data = ml_df[feat_cols].copy().fillna(train_medians).sample(
                                    n=min(50, len(ml_df)), random_state=42
                                )
                                explainer = shap.Explainer(base_model, bg_data)
                                sv = explainer(x_explain_df)
                                import matplotlib.pyplot as plt
                                plt.figure()
                                sv0 = sv[0]
                                if sv0.values.ndim > 1:
                                    sv0 = sv0[:, 1]
                                shap.plots.waterfall(sv0, max_display=10, show=False)
                                st.pyplot(plt.gcf(), clear_figure=True)
                                st.caption("Red bars = factors increasing risk. Blue bars = factors reducing risk.")
                    except Exception as _shap_err:
                        st.warning(f"Could not generate SHAP explanation: {_shap_err}")

# ==============================================================================
# TAB: CHAT WITH DATA (runs as a fragment — no full-page reloads)
# ==============================================================================
with tab_rag:
    coach_guide("Chat with Data", [
        "Type a plain-English question: *‘Who has low readiness this week?’* or *‘ACWR by position?’*",
        "Answers are generated **offline** from your loaded squad data — nothing leaves the machine.",
        "Good for the moments when you don’t want to click through tabs.",
    ])
    _rag_sys = get_rag_system()
    _rag_sys._last_df = df_train
    _rag_sys._last_game_all = game_all
    _rag_sys._last_team_match = team_match
    if len(_rag_sys.knowledge_base) == 0:
        with st.spinner("Building knowledge base (one-time)…"):
            _rag_sys.create_knowledge_base(df_train, game_all, team_match, player_match)

    @st.fragment
    def _chat_fragment():
        """Isolated fragment — interactions here do NOT rerun the rest of the app."""
        rag_system = get_rag_system()

        st.subheader("💬 Chat with Your Data")
        st.caption("Ask anything about your squad — like texting your sports scientist.")

        # ── Header bar ──
        _ch_c1, _ch_c2, _ch_c3 = st.columns([3, 1, 1])
        with _ch_c1:
            st.caption(f"📚 {len(rag_system.knowledge_base)} data points · "
                       f"{'⚡ Cached' if rag_system._HASH_FILE.exists() else '🔨 Fresh build'}")
        with _ch_c2:
            if st.button("🔄 Refresh", key="rag_rebuild", help="Rebuild knowledge base from current data"):
                with st.spinner("Rebuilding…"):
                    rag_system.create_knowledge_base(
                        rag_system._last_df, rag_system._last_game_all,
                        rag_system._last_team_match, player_match, force_rebuild=True)
                st.rerun()
        with _ch_c3:
            if st.button("🗑️ Clear", key="rag_clear_chat", help="Clear conversation"):
                st.session_state["rag_chat_history"] = []
                st.rerun()

        # ── Alerts badge ──
        auto_insights = rag_system.generate_insights(rag_system._last_df)
        if auto_insights:
            with st.expander(f"🔔 {len(auto_insights)} alert(s) detected — tap to review", expanded=False):
                for ins in auto_insights:
                    st.markdown(ins)
        st.markdown("---")

        # ── Chat state ──
        if "rag_chat_history" not in st.session_state:
            st.session_state["rag_chat_history"] = []

        # ── Starter chips ──
        if len(st.session_state["rag_chat_history"]) == 0:
            st.markdown(
                "<div style='text-align:center;color:#999;margin:24px 0 8px'>"
                "What do you want to know about your squad?</div>",
                unsafe_allow_html=True,
            )
            _quick = {
                "💡 Injury risk":     "Which players are at high injury risk right now?",
                "📈 Load this week":  "How should we manage training load this week?",
                "🏃 Fastest players": "Who are the fastest players on the squad?",
                "🔍 Tired players":   "Which players have low readiness or high fatigue?",
                "🏥 Who's out?":      "Which players are out or limited?",
                "📊 ACWR check":      "What do the ACWR values look like across the squad?",
                "😴 Wellness":        "Are any players reporting low sleep or high soreness?",
                "📅 Weekly pattern":  "What does the typical weekly load pattern look like?",
            }
            _qcols = st.columns(4)
            for i, (label, q) in enumerate(_quick.items()):
                with _qcols[i % 4]:
                    if st.button(label, key=f"qk_{i}", use_container_width=True):
                        st.session_state["rag_chat_history"].append({"role": "user", "content": q})
                        answer, _ = rag_system.answer_for_coach(q, top_k=8)
                        st.session_state["rag_chat_history"].append({
                            "role": "assistant", "content": answer, "question": q,
                        })
                        st.rerun()

        # ── Render chat history ──
        _chat_container = st.container()
        with _chat_container:
            for _mi, msg in enumerate(st.session_state["rag_chat_history"]):
                if msg["role"] == "user":
                    with st.chat_message("user", avatar="⚽"):
                        st.markdown(msg["content"])
                else:
                    with st.chat_message("assistant", avatar="🧠"):
                        st.markdown(msg["content"])
                        _q_orig = msg.get("question", "")
                        if _q_orig and _mi == len(st.session_state["rag_chat_history"]) - 1:
                            suggestions = OfflineRAGSystem.follow_up_suggestions(_q_orig)
                            if suggestions:
                                st.markdown("---")
                                st.caption("Follow-up questions:")
                                _fu_cols = st.columns(len(suggestions))
                                for _fi, (_fu_col, _fu_q) in enumerate(zip(_fu_cols, suggestions)):
                                    with _fu_col:
                                        if st.button(f"💬 {_fu_q}", key=f"fu_{_mi}_{_fi}", use_container_width=True):
                                            st.session_state["rag_chat_history"].append({"role": "user", "content": _fu_q})
                                            _fu_ans, _ = rag_system.answer_for_coach(_fu_q, top_k=8)
                                            st.session_state["rag_chat_history"].append({
                                                "role": "assistant", "content": _fu_ans, "question": _fu_q,
                                            })
                                            st.rerun()

        # ── Chat input ──
        _user_input = st.chat_input(
            "Ask about injuries, load, ACWR, speed, wellness, match results…",
            key="rag_chat_input",
        )
        if _user_input:
            st.session_state["rag_chat_history"].append({"role": "user", "content": _user_input})
            with _chat_container:
                with st.chat_message("user", avatar="⚽"):
                    st.markdown(_user_input)
                with st.chat_message("assistant", avatar="🧠"):
                    with st.spinner("Thinking…"):
                        answer, results = rag_system.answer_for_coach(_user_input, top_k=8)
                    if answer and len(answer.strip()) > 20:
                        st.markdown(answer)
                        st.session_state["rag_chat_history"].append({
                            "role": "assistant", "content": answer, "question": _user_input,
                        })
                        suggestions = OfflineRAGSystem.follow_up_suggestions(_user_input)
                        if suggestions:
                            st.markdown("---")
                            st.caption("Follow-up questions:")
                            _fu_cols2 = st.columns(len(suggestions))
                            for _fi2, (_fu_c2, _fu_q2) in enumerate(zip(_fu_cols2, suggestions)):
                                with _fu_c2:
                                    if st.button(f"💬 {_fu_q2}", key=f"fu_new_{_fi2}", use_container_width=True):
                                        st.session_state["rag_chat_history"].append({"role": "user", "content": _fu_q2})
                                        _fu_a2, _ = rag_system.answer_for_coach(_fu_q2, top_k=8)
                                        st.session_state["rag_chat_history"].append({
                                            "role": "assistant", "content": _fu_a2, "question": _fu_q2,
                                        })
                                        st.rerun()
                    else:
                        fallback = ("I couldn't find specific data for that. Try asking about:\n"
                                    "- Player injuries or availability\n"
                                    "- Training load and ACWR\n"
                                    "- Speed and sprint zones\n"
                                    "- Wellness and recovery\n"
                                    "- Match performance")
                        st.markdown(fallback)
                        st.session_state["rag_chat_history"].append({
                            "role": "assistant", "content": fallback, "question": _user_input,
                        })

    _chat_fragment()

st.caption("✅ Uses your 2023/2024/2025 schedules + game stat folders + Catapult training CSV")